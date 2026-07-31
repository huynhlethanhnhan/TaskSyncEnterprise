from pathlib import Path

from docx import Document


report = Path(r"E:\TaskSyncEnterprise\Report")
targets = {
    "Bao_cao_Agile_Scrum_TaskSyncEnterprise.docx",
    "Bao_cao_Kien_truc_phan_mem_TaskSyncEnterprise.docx",
    "Daily_Scrum_1_TaskSyncEnterprise.docx",
    "Daily_Scrum_2_TaskSyncEnterprise.docx",
    "Thuyet_minh_So_do_TaskSyncEnterprise.docx",
    "Bao_cao_Giua_ky_Sprint_1_TaskSyncEnterprise.docx",
    "Bao_cao_dong_gop_thanh_vien_Sprint_1_TaskSyncEnterprise.docx",
}

for path in sorted((p for p in report.rglob("*.docx") if p.name in targets), key=lambda p: p.name):
    document = Document(path)
    text_parts = [paragraph.text for paragraph in document.paragraphs]
    text_parts.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    text = "\n".join(text_parts)
    print(
        path.name,
        f"paragraphs={len(document.paragraphs)}",
        f"tables={len(document.tables)}",
        f"images={len(document.inline_shapes)}",
        f"chars={len(text)}",
        f"placeholders={text.count('ẢNH CẦN BỔ SUNG')}",
        f"sprint1={text.count('Sprint 1')}",
        f"sprint2={text.count('Sprint 2')}",
        f"mojibake={any(marker in text for marker in ('Ã', 'áº', 'Æ'))}",
        f"old_name={text.count('Nguyễn Lê Thành Nhân') + text.count('Phạm Anh Tuấn')}",
        f"old_dates={text.count('02/07-30/07') + text.count('30/07-27/08')}",
        f"figure_x={text.count('Hình X.')}",
    )
