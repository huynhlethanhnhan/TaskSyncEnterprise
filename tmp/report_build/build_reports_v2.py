from __future__ import annotations

import json
import shutil
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image

import build_documents as base


ROOT = Path(r"E:\TaskSyncEnterprise")
REPORT = ROOT / "Report"
AGILE_DIR = REPORT / "Agile"
ARCH_DIR = REPORT / "Môn kiến trúc phần mềm"
DIAGRAM_DIR = REPORT / "Sơ đồ Diagram"
ARCH_DIAGRAM_DIR = DIAGRAM_DIR / "Kiến trúc phần mềm"
AGILE_DIAGRAM_DIR = DIAGRAM_DIR / "Agile"
SOURCE_DIR = REPORT / "Tham khảo"
EXTRACT_DIR = ROOT / "tmp" / "report_build" / "reference_extract"
MEDIA_DIR = EXTRACT_DIR / "Cac_bang_table_media"
THEORY_MEDIA_DIR = EXTRACT_DIR / "Co_su_Ly_Thuyet_media"

for folder in (AGILE_DIR, ARCH_DIR, ARCH_DIAGRAM_DIR, AGILE_DIAGRAM_DIR):
    folder.mkdir(parents=True, exist_ok=True)


TEAM = [
    ("A", "Phạm Tuấn Anh", "2311558672", "Product Owner / Full-stack"),
    ("B", "Nguyễn Đức Mạnh", "2200010420", "Backend Developer"),
    ("C", "Nguyễn Lê Huy Hoàng", "2311554285", "Frontend Developer"),
    ("D", "Phạm Anh Tuấn", "2311559121", "QA / Tester"),
    ("E", "Nguyễn Lê Thành Nhân", "2000004897", "Business Analyst / Documentation"),
]


TITLE_VI = {
    "Secure User Login": "Đăng nhập người dùng an toàn",
    "Automatic Session Renewal": "Tự động gia hạn phiên đăng nhập",
    "Role Definition Layout": "Thiết lập danh mục vai trò",
    "Role-Based Access Enforcement": "Kiểm soát truy cập theo vai trò",
    "Project Creation Wizard": "Trình hướng dẫn tạo dự án",
    "Project Team Assignment Table": "Phân công thành viên vào dự án",
    "Task Kanban Board": "Bảng Kanban quản lý công việc",
    "Bulk Document Downloader": "Tải nhiều tài liệu cùng lúc",
    "Employee Profile Creation": "Tạo hồ sơ nhân viên",
    "Bulk Employee Data Import": "Nhập danh sách nhân viên hàng loạt",
    "Direct Manager Assignment": "Gán quản lý trực tiếp",
    "Department Structure Setup": "Thiết lập cơ cấu phòng ban",
    "Department Head Assignment": "Bổ nhiệm trưởng phòng",
    "Department Transfer Approval Flow": "Quy trình phê duyệt chuyển phòng ban",
    "Task Dependency Diagram": "Sơ đồ phụ thuộc công việc",
    "Advanced Backlog Filtering": "Bộ lọc Backlog nâng cao",
    "Task Decomposition": "Phân rã công việc",
    "Collaborative Comments Feed": "Luồng bình luận cộng tác",
    "Task Creation Editor": "Trình soạn thảo tạo công việc",
    "User Mention Notifications": "Thông báo khi nhắc tên người dùng",
    "Unused Files Cleanup": "Dọn dẹp tệp không còn sử dụng",
    "Leave Request Submission": "Gửi đơn nghỉ phép",
    "Leave Approvals Panel": "Bảng phê duyệt đơn nghỉ phép",
    "Annual Leave Balance Accrual": "Tích lũy số dư phép năm",
    "Emergency Leave Cancellation": "Hủy đơn nghỉ khẩn cấp",
    "Overtime Request Submission": "Gửi yêu cầu làm thêm giờ",
    "Overtime Approvals Panel": "Bảng phê duyệt làm thêm giờ",
    "Overtime Policy Validation": "Kiểm tra chính sách làm thêm giờ",
    "Real-Time In-App Alerts": "Cảnh báo thời gian thực trong ứng dụng",
    "Project Milestone Tracking": "Theo dõi cột mốc dự án",
    "Task Attachment Security": "Bảo mật tệp đính kèm công việc",
    "Notification Channel Settings": "Cài đặt kênh nhận thông báo",
    "Daily Summary Email Digest": "Email tổng hợp công việc hằng ngày",
    "Failed Notification Retry": "Gửi lại thông báo thất bại",
    "Operational Performance Dashboard": "Dashboard hiệu suất vận hành",
    "Department Resource Allocation Map": "Bản đồ phân bổ nguồn lực phòng ban",
    "Custom Project Report Generator": "Trình tạo báo cáo dự án tùy chỉnh",
    "Project Burnup Predictor": "Dự báo Burnup của dự án",
    "Structural Change Auditing": "Kiểm toán thay đổi cơ cấu tổ chức",
    "Security Audit Log Explorer": "Tra cứu nhật ký kiểm toán bảo mật",
}

EPIC_VI = {
    "EP01": "Xác thực và kiểm soát truy cập",
    "EP02": "Quản lý hồ sơ nhân viên",
    "EP03": "Quản lý cơ cấu phòng ban",
    "EP04": "Quản lý dự án và phân bổ nhóm",
    "EP05": "Vòng đời công việc và Kanban",
    "EP06": "Cộng tác nhóm",
    "EP07": "Quản lý tài liệu và tệp",
    "EP08": "Quản lý nghỉ phép",
    "EP09": "Quản lý làm thêm giờ",
    "EP10": "Thông báo đa kênh",
    "EP11": "Phân tích hiệu suất và báo cáo",
    "EP12": "Kiểm toán và an toàn hệ thống",
}


def load_backlog():
    rows = json.loads((EXTRACT_DIR / "current_backlog.json").read_text(encoding="utf-8"))
    result = []
    for row in rows:
        item = {
            "id": row[0],
            "epic": row[1],
            "feature": row[2],
            "title_en": row[3],
            "title": TITLE_VI[row[3]],
            "description_en": row[4],
            "priority": {"Highest": "Cao", "Medium": "Trung bình", "Lowest": "Thấp"}.get(row[5], row[5]),
            "sp": int(row[6]),
            "value": {"High": "Cao", "Medium": "Trung bình", "Low": "Thấp"}.get(row[7], row[7]),
            "source_assignee": row[8],
            "sprint": row[9],
            "source_status": row[10],
        }
        result.append(item)
    return result


BACKLOG = load_backlog()
SPRINT_1 = [x for x in BACKLOG if x["sprint"] == "Sprint 1"]
SPRINT_2 = [x for x in BACKLOG if x["sprint"] == "Sprint 2"]


def academic_document(running_title: str) -> Document:
    doc = base.style_document(Document(), running_title)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.3
    for style_name, size, before, after in (
        ("Heading 1", 16, 14, 8),
        ("Heading 2", 14, 12, 6),
        ("Heading 3", 13, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(13)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.3
    set_page_border(section)
    return doc


def set_page_border(section):
    sect_pr = section._sectPr
    borders = sect_pr.find(qn("w:pgBorders"))
    if borders is None:
        borders = OxmlElement("w:pgBorders")
        borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "10")
        node.set(qn("w:space"), "16")
        node.set(qn("w:color"), "1F4E79")
        borders.append(node)


def add_toc_field(doc):
    doc.add_heading("MỤC LỤC", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Nhấn F9 trong Microsoft Word để cập nhật mục lục và số trang."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    doc.add_page_break()


def add_academic_cover(doc, course: str, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRƯỜNG ĐẠI HỌC NGUYỄN TẤT THÀNH")
    base.set_run_font(r, "Times New Roman", 14, "000000", True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VIỆN ĐÀO TẠO QUỐC TẾ NIIE")
    base.set_run_font(r, "Times New Roman", 14, "000000", True)
    logo = THEORY_MEDIA_DIR / "image1.jpg"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(1.15))
    for _ in range(2):
        doc.add_paragraph()
    for text, size in (
        ("ĐỒ ÁN MÔN HỌC", 15),
        (course.upper(), 15),
        (title.upper(), 19),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        base.set_run_font(r, "Times New Roman", size, "000000", True)
        p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    base.set_run_font(r, "Times New Roman", 13, "444444", False, True)
    p.paragraph_format.space_after = Pt(26)

    rows = [("Giảng viên hướng dẫn", "ThS. Trần Thanh Nhã")]
    rows.extend((f"Thành viên {code}", f"{mssv} - {name} ({role})") for code, name, mssv, role in TEAM)
    base.add_table(doc, ["Thông tin", "Nội dung"], rows, [2500, 6860], font_size=9.2, header_fill="E7E6E6")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("TP. Hồ Chí Minh, tháng 7 năm 2026")
    base.set_run_font(r, "Times New Roman", 13, "000000", False, True)
    doc.add_page_break()


def add_front_matter(doc, document_kind: str, figures: list[str], tables: list[str]):
    doc.add_heading("LỜI CẢM ƠN", level=1)
    base.add_body(
        doc,
        "Nhóm xin chân thành cảm ơn ThS. Trần Thanh Nhã đã hướng dẫn phương pháp tiếp cận, "
        "góp ý cách tổ chức hồ sơ và yêu cầu chuẩn hóa hai Sprint cho đề tài TaskSyncEnterprise. "
        "Quá trình thực hiện giúp nhóm vận dụng kiến thức về phân tích yêu cầu, kiến trúc phần mềm, "
        "Scrum, kiểm thử và phối hợp nhóm vào một sản phẩm quản lý công việc doanh nghiệp."
    )
    base.add_body(
        doc,
        "Tài liệu được tổng hợp từ báo cáo mẫu 106 trang, cơ sở lý thuyết, bộ bảng cơ sở dữ liệu, "
        "các workbook Agile và mã nguồn hiện có. Những mốc thời gian sau ngày 29/07/2026 hoặc "
        "những hình ảnh chưa có bằng chứng được ghi rõ là kế hoạch/ảnh cần bổ sung để nhóm cập nhật trước khi nộp."
    )
    doc.add_page_break()

    doc.add_heading("DANH MỤC TÀI LIỆU LIÊN QUAN", level=1)
    sources = [
        ("1", "Mã nguồn", r"E:\TaskSyncEnterprise"),
        ("2", "Báo cáo Agile mẫu", "LTHN_Agile_Final.pdf - 106 trang"),
        ("3", "Cơ sở lý thuyết", "Co_su_Ly_Thuyet.docx"),
        ("4", "Thiết kế cơ sở dữ liệu và giao diện", "Cac_bang_table.docx"),
        ("5", "Dữ liệu quản lý Agile", "Agile_Project_Management.xlsx; Agile_Tracking.xlsx"),
        ("6", "Daily Scrum tham khảo", "Google Docs: SPRINT 1 DAILY Meeting Minutes"),
    ]
    base.add_table(doc, ["STT", "Loại", "Tài liệu"], sources, [800, 2300, 6260], font_size=9.5)
    doc.add_page_break()

    doc.add_heading("DANH MỤC HÌNH ẢNH", level=1)
    base.add_table(
        doc,
        ["STT", "Tên hình"],
        [(i, name) for i, name in enumerate(figures, 1)],
        [900, 8460],
        font_size=8.3 if len(figures) > 18 else 9.5,
    )
    doc.add_page_break()

    doc.add_heading("DANH MỤC BẢNG BIỂU", level=1)
    base.add_table(
        doc,
        ["STT", "Tên bảng"],
        [(i, name) for i, name in enumerate(tables, 1)],
        [900, 8460],
        font_size=9.5,
    )
    doc.add_page_break()
    add_toc_field(doc)


def add_part_title(doc, number: str, title: str, first: bool = False):
    if not first:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(160)
    r = p.add_run(f"PHẦN {number}")
    base.set_run_font(r, "Times New Roman", 17, "000000", True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title.upper())
    base.set_run_font(r, "Times New Roman", 19, "000000", True)
    doc.add_page_break()


def add_placeholder(doc, filename: str, description: str):
    table = doc.add_table(rows=1, cols=1)
    base.set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    base.set_cell_shading(cell, "FFF2CC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(f"ẢNH CẦN BỔ SUNG\nTên file đề xuất: {filename}\n{description}")
    base.set_run_font(r, "Times New Roman", 11, "7F6000", True)


def image_width_for_page(path: Path, max_width: float = 4.5, max_height: float = 6.4) -> float:
    """Fit tall database screenshots inside the remaining A4 page height."""
    with Image.open(path) as image:
        aspect_ratio = image.width / image.height
    return min(max_width, max_height * aspect_ratio)
    doc.add_paragraph()


def write_mmd(path: Path, text: str):
    path.write_text(text.strip() + "\n", encoding="utf-8")


def make_flow_diagram(filename: str, title: str, subtitle: str, nodes: list[str], directory=ARCH_DIAGRAM_DIR):
    path = directory / filename

    def painter(draw):
        count = len(nodes)
        width = 240 if count <= 5 else 200
        gap = max(25, int((1500 - count * width) / max(1, count - 1)))
        x = 50
        y = 330
        centers = []
        for index, label in enumerate(nodes):
            fill = ["#E8F1FA", "#EEF8F5", "#FFF7E8", "#F6F0FF"][index % 4]
            outline = ["#2E74B5", "#0F766E", "#B78318", "#6B46A1"][index % 4]
            base.box(draw, (x, y, x + width, y + 190), label, fill, outline, fnt=base.font(19, True))
            centers.append((x, x + width))
            x += width + gap
        for i in range(len(centers) - 1):
            base.arrow(draw, (centers[i][1], y + 95), (centers[i + 1][0], y + 95))

    base.save_diagram(path, title, subtitle, painter)
    mermaid_nodes = "\n".join(f'N{i}["{label}"]' for i, label in enumerate(nodes))
    mermaid_edges = "\n".join(f"N{i} --> N{i+1}" for i in range(len(nodes) - 1))
    write_mmd(path.with_suffix(".mmd"), f"flowchart LR\n{mermaid_nodes}\n{mermaid_edges}")
    return path


def build_extended_diagrams():
    base.build_diagrams()
    copies = [
        ("01_Phan_ra_chuc_nang.png", "10_Functional_Decomposition.png"),
        ("02_Use_Case_Tong_Quat.png", "11_Use_Case_Tong_Quat.png"),
        ("03_Activity_Task_Lifecycle.png", "14_Activity_Task_Lifecycle.png"),
    ]
    for source, target in copies:
        shutil.copy2(AGILE_DIAGRAM_DIR / source, ARCH_DIAGRAM_DIR / target)
        shutil.copy2((AGILE_DIAGRAM_DIR / source).with_suffix(".mmd"), (ARCH_DIAGRAM_DIR / target).with_suffix(".mmd"))

    make_flow_diagram(
        "06_Frontend_Layers.png",
        "Kiến trúc phân lớp Frontend",
        "Tổ chức React Router, trang tính năng, component, hook và lớp gọi API",
        ["React Router", "Trang tính năng", "Component dùng chung", "TanStack Query Hooks", "Axios API"],
    )
    make_flow_diagram(
        "07_Data_Flow.png",
        "Luồng dữ liệu tổng quát",
        "Dữ liệu từ trình duyệt tới nguồn dữ liệu chuẩn và quay lại giao diện",
        ["Người dùng", "React SPA", "FastAPI Router", "Service / CRUD", "SQL Server"],
    )
    make_flow_diagram(
        "08_Authentication_Flow.png",
        "Luồng xác thực và phân quyền",
        "Đăng nhập, cấp token, kiểm tra RBAC và ghi nhật ký",
        ["Thông tin đăng nhập", "Xác thực mật khẩu", "Cấp Access/Refresh Token", "Kiểm tra RBAC", "Audit Log"],
    )
    make_flow_diagram(
        "09_Realtime_Event_Flow.png",
        "Luồng sự kiện thời gian thực",
        "Domain event chỉ phát sau khi giao dịch dữ liệu đã commit",
        ["Mutation", "SQL Commit", "Publish Domain Event", "WebSocket", "Invalidate Query Cache", "Tải lại dữ liệu"],
    )
    make_flow_diagram(
        "12_Activity_Login.png",
        "Activity Diagram - Đăng nhập",
        "Kiểm tra dữ liệu, mật khẩu, trạng thái tài khoản và phiên đăng nhập",
        ["Nhập tài khoản", "Validate", "Kiểm tra mật khẩu", "Kiểm tra trạng thái", "Tạo phiên", "Vào hệ thống"],
    )
    make_flow_diagram(
        "13_Activity_Project_Creation.png",
        "Activity Diagram - Tạo dự án",
        "Tạo Project, gán thành viên và sinh Backlog khởi tạo",
        ["Nhập thông tin dự án", "Kiểm tra quyền", "Tạo Project", "Gán thành viên", "Tạo Backlog", "Ghi Audit"],
    )
    make_flow_diagram(
        "15_Activity_Leave_Approval.png",
        "Activity Diagram - Duyệt nghỉ phép",
        "Nộp đơn, kiểm tra chính sách, phê duyệt và gửi thông báo",
        ["Nộp đơn", "Kiểm tra số dư", "Quản lý xem xét", "Phê duyệt / từ chối", "Cập nhật số dư", "Thông báo"],
    )
    make_flow_diagram(
        "16_Sequence_Create_Task.png",
        "Sequence tổng quát - Tạo công việc",
        "Các thành phần tham gia trong giao dịch tạo Task",
        ["React Form", "Task Router", "Task Service", "SQLAlchemy", "SQL Server", "WebSocket"],
    )
    make_flow_diagram(
        "17_Sequence_Notification.png",
        "Sequence tổng quát - Thông báo",
        "Từ sự kiện nghiệp vụ tới kênh trong ứng dụng và email",
        ["Domain Event", "Notification Service", "Notification DB", "WebSocket", "Email Provider", "Người nhận"],
    )
    make_flow_diagram(
        "18_ERD_Core.png",
        "ERD rút gọn",
        "Quan hệ cốt lõi Organization - Project - Task - Notification",
        ["Department / Team", "Employee", "Project", "Task", "Attachment / Comment", "Notification / Audit"],
    )
    make_flow_diagram(
        "19_Module_Dependencies.png",
        "Sơ đồ phụ thuộc module",
        "Các module nghiệp vụ dùng chung Identity, Organization và hạ tầng",
        ["Identity", "Organization", "Project", "Task", "Collaboration", "Notification", "Reporting"],
    )
    make_flow_diagram(
        "20_Backup_Recovery.png",
        "Luồng sao lưu và khôi phục",
        "Kiểm tra tính toàn vẹn dữ liệu trước và sau khôi phục",
        ["SQL Server", "Backup Job", "Kho lưu trữ", "Kiểm tra checksum", "Restore thử nghiệm", "Xác nhận RPO/RTO"],
    )
    make_flow_diagram(
        "21_Observability.png",
        "Sơ đồ giám sát vận hành",
        "Log, metrics và cảnh báo phục vụ phát hiện sự cố",
        ["FastAPI / Nginx", "Structured Logs", "Prometheus Metrics", "Grafana Dashboard", "Alert", "Nhóm vận hành"],
    )


def category_for(item):
    epic = item["epic"]
    if epic == "EP01":
        return "auth"
    if epic in ("EP02", "EP03"):
        return "organization"
    if epic == "EP04":
        return "project"
    if epic in ("EP05", "EP06"):
        return "task"
    if epic == "EP07":
        return "file"
    if epic in ("EP08", "EP09"):
        return "approval"
    if epic == "EP10":
        return "notification"
    if epic == "EP11":
        return "report"
    return "audit"


def actor_for(item):
    return {
        "auth": "Nhân viên; Quản trị viên",
        "organization": "Quản trị viên nhân sự; Quản lý phòng ban",
        "project": "Quản lý dự án; Product Owner",
        "task": "Thành viên dự án; Quản lý dự án",
        "file": "Thành viên dự án; Quản trị viên",
        "approval": "Nhân viên; Quản lý trực tiếp; Nhân sự",
        "notification": "Người dùng hệ thống; Dịch vụ thông báo",
        "report": "Quản lý; Quản trị viên",
        "audit": "Quản trị viên; Kiểm toán viên",
    }[category_for(item)]


def use_case_content(item):
    category = category_for(item)
    title = item["title"]
    actor = actor_for(item).split(";")[0]
    descriptions = {
        "auth": f"{actor} sử dụng chức năng {title.lower()} để truy cập đúng phạm vi quyền và duy trì phiên làm việc an toàn.",
        "organization": f"{actor} thực hiện {title.lower()} nhằm duy trì cơ cấu tổ chức và dữ liệu nhân sự nhất quán.",
        "project": f"{actor} thực hiện {title.lower()} để thiết lập phạm vi, nguồn lực và kế hoạch của dự án.",
        "task": f"{actor} sử dụng {title.lower()} để tổ chức, theo dõi và cộng tác trên công việc.",
        "file": f"{actor} sử dụng {title.lower()} để quản lý tài liệu có kiểm soát quyền và dấu vết.",
        "approval": f"{actor} thực hiện {title.lower()} theo chính sách hành chính và luồng phê duyệt.",
        "notification": f"{actor} sử dụng {title.lower()} để nhận thông tin đúng sự kiện, đúng kênh và đúng thời điểm.",
        "report": f"{actor} sử dụng {title.lower()} để tổng hợp dữ liệu, phân tích tiến độ và hỗ trợ quyết định.",
        "audit": f"{actor} sử dụng {title.lower()} để truy vết thay đổi và kiểm tra tuân thủ.",
    }
    basic = [
        f"Người dùng mở chức năng “{title}” từ menu phù hợp với vai trò.",
        "Hệ thống tải dữ liệu trong phạm vi quyền và hiển thị trạng thái hiện tại.",
        "Người dùng nhập hoặc lựa chọn các thông tin nghiệp vụ bắt buộc.",
        "Frontend kiểm tra định dạng; Backend kiểm tra quyền, quy tắc nghiệp vụ và tính nhất quán.",
        "Hệ thống ghi dữ liệu trong giao dịch; chỉ phát sự kiện sau khi commit thành công.",
        "Giao diện cập nhật kết quả, hiển thị thông báo và ghi Audit Log khi thao tác có tính nhạy cảm.",
    ]
    alternatives = [
        "Nếu thiếu dữ liệu bắt buộc, hệ thống giữ nguyên biểu mẫu và chỉ rõ trường cần sửa.",
        "Nếu người dùng không có quyền, Backend trả lỗi 403 và không thực hiện thay đổi dữ liệu.",
        "Nếu giao dịch thất bại, hệ thống rollback, ghi log kỹ thuật và cho phép người dùng thử lại an toàn.",
    ]
    rules = [
        "Mọi truy vấn và mutation phải áp dụng phạm vi quyền tại Backend; việc ẩn nút ở Frontend chỉ hỗ trợ trải nghiệm.",
        "Dữ liệu nghiệp vụ chỉ được xem là hoàn thành khi thỏa tiêu chí chấp nhận và có thể truy vết tới người thực hiện.",
    ]
    if category == "approval":
        rules.append("Người duyệt không được tự duyệt yêu cầu của chính mình; số dư/chính sách phải được kiểm tra trong cùng giao dịch.")
    if category == "file":
        rules.append("Tệp phải được kiểm tra loại MIME, dung lượng, phạm vi dự án và quyền tải xuống.")
    if category == "notification":
        rules.append("Lỗi gửi phải lưu trạng thái, số lần retry và phản hồi nhà cung cấp; không tạo thông báo trùng cho cùng event_id.")
    nfr = [
        "Phản hồi thao tác thông thường mục tiêu dưới 2 giây trong điều kiện tải chuẩn; tác vụ nặng chạy bất đồng bộ nếu cần.",
        "Không để lộ dữ liệu ngoài phạm vi dự án/phòng ban; thông tin nhạy cảm phải được bảo vệ khi truyền và lưu.",
    ]
    acceptance = [
        f"Thực hiện được “{title}” với dữ liệu hợp lệ và kết quả được lưu đúng.",
        "Trường hợp sai dữ liệu hoặc sai quyền bị từ chối với thông báo có thể hành động.",
        "Sau thao tác, dữ liệu hiển thị nhất quán giữa các người dùng liên quan và có dấu vết kiểm tra.",
    ]
    return descriptions[category], basic, alternatives, rules, nfr, acceptance


def add_use_case_spec(doc, item, number):
    doc.add_page_break()
    doc.add_heading(f"{number}. Đặc tả Use Case {item['id']} - {item['title']}", level=2)
    description, basic, alternatives, rules, nfr, acceptance = use_case_content(item)
    metadata = [
        ("Use Case ID", item["id"]),
        ("Tên Use Case", item["title"]),
        ("Epic", f"{item['epic']} - {EPIC_VI[item['epic']]}"),
        ("Tác nhân", actor_for(item)),
        ("Độ ưu tiên", item["priority"]),
        ("Story Points", item["sp"]),
        ("Mô tả", description),
        ("Tiền điều kiện", "Người dùng đã xác thực; dữ liệu nền và phạm vi quyền đã được cấu hình."),
        ("Hậu điều kiện", "Dữ liệu và trạng thái liên quan được cập nhật nhất quán; sự kiện/audit được ghi khi áp dụng."),
    ]
    base.add_table(doc, ["Thuộc tính", "Nội dung"], metadata, [2200, 7160], font_size=8.8)
    doc.add_heading("Luồng chính", level=3)
    base.add_bullets(doc, basic, numbered=True)
    doc.add_heading("Luồng thay thế và ngoại lệ", level=3)
    base.add_bullets(doc, alternatives)
    doc.add_heading("Quy tắc nghiệp vụ", level=3)
    base.add_bullets(doc, rules)
    doc.add_heading("Yêu cầu phi chức năng", level=3)
    base.add_bullets(doc, nfr)
    doc.add_heading("Tiêu chí chấp nhận", level=3)
    base.add_bullets(doc, acceptance)


def add_theory(doc):
    doc.add_heading("1. Agile và Scrum", level=1)
    base.add_body(
        doc,
        "Agile là định hướng phát triển nhấn mạnh khả năng thích ứng, cộng tác với khách hàng và giao phần mềm có giá trị theo từng phần nhỏ. "
        "Scrum là một framework cụ thể để hiện thực hóa các nguyên tắc này thông qua chu kỳ Sprint, Product Backlog, các sự kiện kiểm tra - thích nghi "
        "và một nhóm đa chức năng tự tổ chức."
    )
    doc.add_heading("1.1. Bốn giá trị của Tuyên ngôn Agile", level=2)
    base.add_bullets(
        doc,
        [
            "Cá nhân và tương tác quan trọng hơn quy trình và công cụ.",
            "Phần mềm hoạt động quan trọng hơn tài liệu quá mức.",
            "Cộng tác với khách hàng quan trọng hơn thương lượng hợp đồng.",
            "Phản hồi với thay đổi quan trọng hơn bám cứng vào kế hoạch.",
        ],
    )
    doc.add_heading("2. Năm giá trị Scrum", level=1)
    values = [
        ("Dũng cảm", "Thẳng thắn nêu rủi ro, thử nghiệm giải pháp mới và xử lý vấn đề khó."),
        ("Tập trung", "Ưu tiên Sprint Goal, giới hạn công việc đang làm và tránh phân tán nguồn lực."),
        ("Cam kết", "Chủ động chịu trách nhiệm với mục tiêu chung và Definition of Done."),
        ("Tôn trọng", "Công nhận năng lực, chuyên môn và quyền đóng góp của từng thành viên."),
        ("Cởi mở", "Minh bạch về tiến độ, trở ngại, chất lượng và thay đổi phạm vi."),
    ]
    base.add_table(doc, ["Giá trị", "Áp dụng trong TaskSyncEnterprise"], values, [2200, 7160], font_size=9.2)
    if (THEORY_MEDIA_DIR / "image2.png").exists():
        base.add_diagram(doc, THEORY_MEDIA_DIR / "image2.png", "Hình. Năm giá trị cốt lõi của Scrum", width=4.5)

    doc.add_heading("3. Ba trụ cột Scrum", level=1)
    pillars = [
        ("Minh bạch", "Backlog, trạng thái Task, blocker, tiêu chí Done và rủi ro phải được nhìn thấy và hiểu giống nhau."),
        ("Thanh tra", "Daily Scrum, Review, kiểm thử và dashboard được dùng để phát hiện sai lệch sớm."),
        ("Thích nghi", "Backlog, cách phân công và giải pháp kỹ thuật được điều chỉnh khi bằng chứng cho thấy kế hoạch không còn phù hợp."),
    ]
    base.add_table(doc, ["Trụ cột", "Cách vận dụng"], pillars, [2200, 7160], font_size=9.2)
    if (THEORY_MEDIA_DIR / "image5.png").exists():
        base.add_diagram(doc, THEORY_MEDIA_DIR / "image5.png", "Hình. Ba trụ cột Scrum trên nền tảng lòng tin", width=5.2)

    doc.add_heading("4. Artifact và cam kết", level=1)
    artifacts = [
        ("Product Backlog", "Danh sách có thứ tự mọi yêu cầu sản phẩm", "Product Goal"),
        ("Sprint Backlog", "Sprint Goal, các Product Backlog Item được chọn và kế hoạch thực hiện", "Sprint Goal"),
        ("Increment", "Tổng phần sản phẩm đã hoàn thành và có thể sử dụng", "Definition of Done"),
    ]
    base.add_table(doc, ["Artifact", "Ý nghĩa", "Cam kết"], artifacts, [2100, 5100, 2160], font_size=9.2)

    doc.add_heading("5. Vai trò và trách nhiệm", level=1)
    roles = [
        ("Product Owner", "Tối đa hóa giá trị, sắp xếp Product Backlog và làm rõ mục tiêu."),
        ("Scrum Master", "Bảo vệ framework, tháo gỡ trở ngại và giúp nhóm cải tiến."),
        ("Development Team", "Phân tích, thiết kế, lập trình, kiểm thử và tạo Increment đạt Done."),
    ]
    base.add_table(doc, ["Vai trò", "Trách nhiệm chính"], roles, [2400, 6960], font_size=9.4)

    doc.add_heading("6. Các sự kiện Scrum", level=1)
    events = [
        ("Sprint Planning", "Xác định vì sao Sprint có giá trị, làm gì và làm như thế nào."),
        ("Daily Scrum", "Kiểm tra tiến độ hướng tới Sprint Goal và lập kế hoạch cho 24 giờ tiếp theo."),
        ("Sprint Review", "Kiểm tra Increment cùng stakeholder và cập nhật Product Backlog."),
        ("Sprint Retrospective", "Lựa chọn cải tiến cụ thể cho con người, quy trình và công cụ."),
    ]
    base.add_table(doc, ["Sự kiện", "Mục tiêu"], events, [2600, 6760], font_size=9.4)
    if (THEORY_MEDIA_DIR / "image3.png").exists():
        base.add_diagram(doc, THEORY_MEDIA_DIR / "image3.png", "Hình. Scrum Framework", width=6.0)

    doc.add_heading("7. Ước lượng, ưu tiên và Definition of Done", level=1)
    base.add_body(
        doc,
        "Nhóm sử dụng Story Points để ước lượng tương đối độ phức tạp, rủi ro và nỗ lực. Ưu tiên được cân nhắc theo giá trị kinh doanh, "
        "phụ thuộc kỹ thuật và rủi ro. Một hạng mục chỉ được xem là Done khi mã nguồn đã được review, kiểm thử đạt, tài liệu liên quan được cập nhật, "
        "không còn lỗi nghiêm trọng mở và có bằng chứng truy vết."
    )
    doc.add_heading("8. Burndown/Burnup và kiểm soát tiến độ", level=1)
    base.add_body(
        doc,
        "Burndown theo dõi khối lượng còn lại; Burnup theo dõi khối lượng đã hoàn thành so với tổng phạm vi. "
        "Trong bộ workbook, các chỉ số được tính bằng công thức từ Product Backlog và bảng phân rã công việc để tránh tổng hợp thủ công."
    )


def build_agile_report():
    figures = [
        "Năm giá trị cốt lõi của Scrum",
        "Ba trụ cột Scrum",
        "Scrum Framework",
        "Sơ đồ phân rã chức năng TaskSyncEnterprise",
        "Use Case tổng quát",
        "Activity Diagram vòng đời Task",
        "Activity Diagram chu trình Scrum",
        "Scrum Workflow",
    ]
    tables = [
        "Danh sách thành viên Scrum Team",
        "Danh sách đối tượng người dùng",
        "Danh sách 12 Epic",
        "Product Backlog 50 User Stories",
        "Sprint Backlog 1",
        "Sprint Backlog 2",
        "Bảng đánh giá Sprint 1",
        "Bảng đánh giá Sprint 2",
    ]
    doc = academic_document("BÁO CÁO AGILE | TASKSYNCENTERPRISE")
    add_academic_cover(
        doc,
        "Mô hình phát triển phần mềm Agile",
        "Phần mềm quản lý công việc TaskSyncEnterprise",
        "Báo cáo triển khai Scrum - chuẩn hóa 2 Sprint",
    )
    add_front_matter(doc, "Agile", figures, tables)
    add_part_title(doc, "1", "Tổng quan", first=True)
    doc.add_heading("1. Lý do chọn đề tài", level=1)
    base.add_body(
        doc,
        "Doanh nghiệp cần một nền tảng thống nhất để tổ chức nhân sự, lập kế hoạch dự án, phân công công việc, theo dõi tiến độ, "
        "cộng tác và quản lý các yêu cầu hành chính. Khi thông tin bị phân tán qua nhiều công cụ, người quản lý khó nhìn thấy rủi ro, "
        "nhân viên dễ bỏ sót trách nhiệm và dữ liệu kiểm toán không đầy đủ. TaskSyncEnterprise được chọn nhằm giải quyết bài toán đó."
    )
    doc.add_heading("2. Mục tiêu nghiên cứu", level=1)
    base.add_bullets(
        doc,
        [
            "Xây dựng nền tảng Full-stack sử dụng React, FastAPI, SQL Server và Redis.",
            "Chuẩn hóa Product Backlog, Sprint Backlog, Daily Scrum, Review và Retrospective.",
            "Áp dụng phân quyền, audit, realtime và các quy tắc nhất quán dữ liệu.",
            "Tạo hồ sơ có thể tiếp tục cập nhật khi nhóm bổ sung ảnh, video và bằng chứng kiểm thử.",
        ],
    )
    doc.add_heading("3. Nhiệm vụ và phương pháp", level=1)
    base.add_body(
        doc,
        "Nhóm khảo sát yêu cầu, phân tích tài liệu, đọc mã nguồn, xây dựng backlog, phân rã User Story thành công việc kỹ thuật, "
        "thiết kế kiến trúc và cơ sở dữ liệu, triển khai theo Sprint, kiểm thử và tổng hợp bằng chứng. "
        "Báo cáo mẫu 106 trang được dùng làm khung tổ chức, không sao chép nội dung nghiệp vụ rạp phim."
    )
    doc.add_heading("4. Phạm vi đề tài", level=1)
    base.add_bullets(
        doc,
        [
            "Xác thực, phiên đăng nhập và RBAC.",
            "Tổ chức, phòng ban, nhóm và hồ sơ nhân viên.",
            "Dự án, Backlog, Sprint, Task, Kanban và cộng tác.",
            "Tệp đính kèm, thông báo đa kênh, nghỉ phép, làm thêm giờ.",
            "Dashboard, báo cáo, audit và hỗ trợ vận hành.",
        ],
    )
    doc.add_heading("5. Quy ước hai Sprint", level=1)
    base.add_table(
        doc,
        ["Sprint mới", "Sprint nguồn được gộp", "Thời gian", "Phạm vi", "Story Points"],
        [
            ("Sprint 1", "Sprint cũ 1 + 2", "02/07-23/07/2026", "Kết quả giữa kỳ", sum(x["sp"] for x in SPRINT_1)),
            ("Sprint 2", "Sprint cũ 3 + 4", "24/07-20/08/2026", "Kế hoạch/tiếp tục xác minh", sum(x["sp"] for x in SPRINT_2)),
        ],
        [1500, 1900, 1900, 2100, 1960],
        font_size=9.2,
    )
    base.add_note(
        doc,
        "Nguyên tắc kiểm soát số liệu",
        f"Product Backlog đã được ước lượng lại theo thang Fibonacci, gồm {len(BACKLOG)} User Story và "
        f"{sum(x['sp'] for x in BACKLOG)} Story Point. Trạng thái Jira được tách khỏi trạng thái đã xác minh.",
        "DDEBF7",
    )

    add_part_title(doc, "2", "Cơ sở lý luận")
    add_theory(doc)

    add_part_title(doc, "3", "Diễn biến trong quy trình Scrum")
    doc.add_heading("CHƯƠNG I. TỔ CHỨC PRODUCT BACKLOG", level=1)
    doc.add_heading("1.1. Đối tượng người dùng", level=2)
    actors = [
        ("Quản trị viên", "Cấu hình hệ thống, tài khoản, vai trò, cơ cấu tổ chức, audit."),
        ("Quản lý / Trưởng nhóm", "Lập kế hoạch dự án, phân công Task, phê duyệt yêu cầu, xem báo cáo."),
        ("Nhân viên", "Thực hiện Task, cộng tác, gửi nghỉ phép/làm thêm giờ và nhận thông báo."),
        ("Product Owner", "Ưu tiên Backlog, xác định Product Goal và chấp nhận kết quả."),
        ("Kiểm toán / Vận hành", "Theo dõi log, rủi ro, hiệu năng và khả năng phục hồi."),
    ]
    base.add_table(doc, ["Đối tượng", "Nhu cầu chính"], actors, [2500, 6860], font_size=9.4)
    base.add_diagram(doc, AGILE_DIAGRAM_DIR / "01_Phan_ra_chuc_nang.png", "Hình. Phân rã chức năng TaskSyncEnterprise")
    base.add_body(
        doc,
        "Các nhánh chức năng dùng chung cho cả hai Sprint. Báo cáo chỉ giải thích ký pháp một lần: khối cấp trên là phạm vi sản phẩm, "
        "khối cấp dưới là miền chức năng và các mũi tên thể hiện quan hệ phân rã."
    )
    base.add_diagram(doc, AGILE_DIAGRAM_DIR / "02_Use_Case_Tong_Quat.png", "Hình. Use Case tổng quát")
    doc.add_heading("1.2. Scrum Team", level=2)
    base.add_table(doc, ["Mã", "Họ tên", "MSSV", "Vai trò"], TEAM, [700, 2300, 1700, 4660], font_size=9.2)
    doc.add_heading("1.3. Danh sách Epic", level=2)
    epic_rows = []
    for epic, name in EPIC_VI.items():
        count = sum(1 for item in BACKLOG if item["epic"] == epic)
        epic_rows.append((epic, name, count))
    base.add_table(doc, ["Epic", "Tên Epic", "Số User Stories"], epic_rows, [1100, 6260, 2000], font_size=9.3)
    doc.add_heading("1.4. Product Backlog tổng hợp", level=2)
    backlog_rows = [
        (x["id"], x["epic"], x["title"], x["priority"], x["sp"], x["sprint"])
        for x in BACKLOG
    ]
    base.add_table(
        doc,
        ["US", "Epic", "Tên User Story", "Ưu tiên", "SP", "Sprint"],
        backlog_rows,
        [700, 800, 4300, 1100, 700, 1760],
        font_size=7.8,
    )
    add_placeholder(
        doc,
        "product_backlog_board.png",
        "Product Backlog 50 User Story sau khi chuẩn hóa hai Sprint và tách trạng thái Jira khỏi trạng thái đã xác minh.",
    )

    for sprint_number, items, period, goal in (
        (
            1,
            SPRINT_1,
            "02/07-23/07/2026",
            "Tạo Increment có thể trình bày giữa kỳ với Login/Auth, Project, Task/Kanban, Employee sơ khai, Notification và Calendar.",
        ),
        (
            2,
            SPRINT_2,
            "24/07-20/08/2026",
            "Ổn định các chức năng chưa xác minh, hoàn thiện tổ chức, nghỉ phép/OT, báo cáo, kiểm toán và cập nhật hệ thống.",
        ),
    ):
        doc.add_page_break()
        doc.add_heading(f"CHƯƠNG {sprint_number + 1}. DIỄN BIẾN SPRINT {sprint_number}", level=1)
        doc.add_heading(f"{sprint_number + 1}.1. Lập kế hoạch Sprint", level=2)
        base.add_table(
            doc,
            ["Nội dung", "Chi tiết"],
            [
                ("Thời gian", period),
                ("Sprint Goal", goal),
                ("Phạm vi", f"{items[0]['id']}-{items[-1]['id']}"),
                ("Quy mô", f"{len(items)} User Stories - {sum(x['sp'] for x in items)} Story Points"),
                ("Definition of Done", "Code review; kiểm thử đạt; cập nhật tài liệu; không còn lỗi nghiêm trọng; có evidence."),
            ],
            [2300, 7060],
            font_size=9.3,
        )
        sprint_rows = [(x["id"], x["title"], x["priority"], x["sp"], x["epic"]) for x in items]
        base.add_table(doc, ["US", "Tên User Story", "Ưu tiên", "SP", "Epic"], sprint_rows, [750, 4900, 1200, 700, 1810], font_size=8.3)
        add_placeholder(
            doc,
            f"sprint_{sprint_number}_planning_meeting.png",
            f"Ảnh buổi Sprint Planning {sprint_number}: thành viên, màn hình backlog và kết quả cam kết Sprint Goal.",
        )
        doc.add_heading(f"{sprint_number + 1}.2. Đặc tả yêu cầu trong Sprint", level=2)
        for index, item in enumerate(items, 1):
            add_use_case_spec(doc, item, f"{sprint_number + 1}.2.{index}")

        doc.add_page_break()
        doc.add_heading(f"{sprint_number + 1}.3. Thiết kế và triển khai", level=2)
        if sprint_number == 1:
            base.add_diagram(doc, AGILE_DIAGRAM_DIR / "03_Activity_Task_Lifecycle.png", "Hình. Activity Diagram vòng đời Task")
            base.add_body(
                doc,
                "Vòng đời Task được dùng chung cho toàn sản phẩm: tạo, kiểm tra quyền, phân công, thực hiện, review tiêu chí chấp nhận và ghi audit. "
                "Các chương sau chỉ tham chiếu lại sơ đồ này, không lặp lại phần giải thích."
            )
        else:
            base.add_diagram(doc, AGILE_DIAGRAM_DIR / "04_Activity_Sprint.png", "Hình. Activity Diagram chu trình Scrum")
            base.add_body(
                doc,
                "Sơ đồ mô tả vòng phản hồi Product Backlog - Planning - Development/Daily - Review - Retrospective - Increment. "
                "Kết quả Review và Retrospective được đưa trở lại Backlog để ưu tiên cho giai đoạn tiếp theo."
            )
        add_placeholder(
            doc,
            f"sprint_{sprint_number}_implementation_evidence.png",
            "Ảnh màn hình chức năng chính hoặc pull request/commit minh chứng kết quả triển khai của Sprint.",
        )
        doc.add_heading(f"{sprint_number + 1}.4. Daily Scrum", level=2)
        base.add_body(
            doc,
            f"Biên bản Daily Scrum đã được tách thành tài liệu Daily_Scrum_{sprint_number}_TaskSyncEnterprise.docx. "
            "Mỗi ngày có ba nội dung: đã làm, sẽ làm và trở ngại; các dòng chưa có bằng chứng được đánh dấu để nhóm xác nhận."
        )
        add_placeholder(
            doc,
            f"sprint_{sprint_number}_daily_meeting_photo.png",
            "Ảnh hoặc đường dẫn video ghi hình Daily Scrum; nên thể hiện ngày họp và danh sách người tham dự.",
        )
        doc.add_heading(f"{sprint_number + 1}.5. Sprint Review", level=2)
        review_rows = [
            ("Mục tiêu Sprint", goal),
            (
                "Kết quả đã xác minh",
                "Sprint 1 dùng Git, mã nguồn và demo làm bằng chứng; Sprint 2 được trình bày dưới dạng kế hoạch hoặc phần tiếp tục xác minh.",
            ),
            ("Demos and feedback", "Demo theo luồng nghiệp vụ chính; ghi phản hồi và chuyển yêu cầu mới về Product Backlog."),
            ("Backlog review", "Kiểm tra thứ tự ưu tiên, Story Points, phụ thuộc và tiêu chí chấp nhận."),
            ("Operational review", "Kiểm tra lỗi, hiệu năng, bảo mật, readiness và khả năng khôi phục."),
        ]
        base.add_table(doc, ["Hạng mục", "Nội dung đánh giá"], review_rows, [2500, 6860], font_size=9.2)
        add_placeholder(
            doc,
            f"sprint_{sprint_number}_review_demo.png",
            "Ảnh buổi Sprint Review hoặc màn hình demo Increment; bổ sung phản hồi của giảng viên/stakeholder.",
        )
        doc.add_heading(f"{sprint_number + 1}.6. Sprint Retrospective", level=2)
        retro_rows = [
            ("Đã làm tốt", "Backlog có mã định danh; vai trò được phân công; mã nguồn và tài liệu được lưu cùng phạm vi."),
            ("Chưa làm tốt", "Một số nguồn chưa thống nhất họ tên/MSSV, tổng Story Points và bằng chứng trạng thái Done."),
            ("Ý tưởng cải tiến", "Chuẩn hóa Definition of Done; liên kết mỗi User Story với PR, test, ảnh demo và biên bản."),
            ("Hành động tiếp theo", "Cập nhật workbook hằng ngày; chốt blocker trong Daily Scrum; không để hạng mục không có người chịu trách nhiệm."),
        ]
        base.add_table(doc, ["Nhóm nội dung", "Kết luận"], retro_rows, [2300, 7060], font_size=9.2)
        add_placeholder(
            doc,
            f"sprint_{sprint_number}_retrospective_board.png",
            "Ảnh bảng Start/Stop/Continue hoặc Went well/To improve/Action items của Sprint Retrospective.",
        )

    doc.add_page_break()
    doc.add_heading("PHỤ LỤC 1. BIÊN BẢN SPRINT PLANNING", level=1)
    for sprint_number, items in ((1, SPRINT_1), (2, SPRINT_2)):
        doc.add_heading(f"Phụ lục 1.{sprint_number}. Sprint Planning {sprint_number}", level=2)
        base.add_table(
            doc,
            ["Mục", "Nội dung"],
            [
                ("Thành phần", ", ".join(name for _, name, _, _ in TEAM)),
                ("Đầu vào", "Product Backlog đã ưu tiên; capacity; rủi ro; Definition of Done."),
                ("Đầu ra", f"Sprint Backlog {items[0]['id']}-{items[-1]['id']}; {sum(x['sp'] for x in items)} SP."),
                ("Quyết định", "Chia User Story thành các việc phân tích, Backend, Frontend, kiểm thử và tài liệu."),
                ("Bằng chứng cần bổ sung", f"sprint_{sprint_number}_planning_minutes_signed.png"),
            ],
            [2500, 6860],
            font_size=9.2,
        )
    doc.add_heading("PHỤ LỤC 2. DAILY SCRUM", level=1)
    base.add_body(
        doc,
        "Daily Scrum được bàn giao thành hai tài liệu độc lập để tránh gộp biên bản giữa hai Sprint. "
        "Các mục “Cần nhóm xác nhận” phải được thay bằng thông tin thực tế trước khi nộp bản chính thức."
    )
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    base.add_bullets(
        doc,
        [
            "Scrum Guide 2020 - Scrum.org.",
            "Atlassian Agile Coach - Scrum, Product Backlog và Sprint.",
            "LTHN_Agile_Final.pdf - cấu trúc báo cáo tham khảo.",
            "Co_su_Ly_Thuyet.docx - nội dung cơ sở lý luận do nhóm cung cấp.",
            "Mã nguồn và tài liệu kiến trúc TaskSyncEnterprise tại E:\\TaskSyncEnterprise.",
        ],
    )
    base.save_doc(doc, AGILE_DIR / "Bao_cao_Agile_Scrum_TaskSyncEnterprise.docx")


SCHEMA_NAMES = [
    "roles",
    "departments",
    "teams",
    "employees",
    "projects",
    "project_members",
    "tasks",
    "task_assignments",
    "task_comments",
    "task_checklists",
    "task_attachments",
    "notifications",
    "notification_preferences",
    "notification_logs",
    "user_sessions",
    "refresh_tokens",
    "token_blacklist",
    "vacations",
    "audit_logs",
]


def schema_tables():
    data = json.loads((EXTRACT_DIR / "Cac_bang_table_tables.json").read_text(encoding="utf-8"))
    result = []
    for index, item in enumerate(data):
        name = SCHEMA_NAMES[index]
        rows = item["rows"]
        if name == "task_attachments":
            rows = [
                ["Tên", "Mô tả", "Kiểu dữ liệu", "NULL", "Ràng buộc"],
                ["id", "Mã tệp đính kèm", "INT", "NO", "Primary key"],
                ["task_id", "Task liên quan", "INT", "YES", "FK -> tasks(id), ON DELETE CASCADE"],
                ["topic_id", "Chủ đề thảo luận liên quan", "INT", "YES", "FK -> discussion_topics(id)"],
                ["reply_id", "Phản hồi thảo luận liên quan", "INT", "YES", "FK -> discussion_replies(id)"],
                ["feedback_id", "Phản hồi người dùng liên quan", "INT", "YES", "FK -> user_feedback(id)"],
                ["file_name", "Tên tệp gốc", "VARCHAR(255)", "NO", ""],
                ["file_path", "Đường dẫn lưu trữ", "VARCHAR(500)", "NO", ""],
                ["file_size", "Dung lượng byte", "INT", "NO", ""],
                ["mime_type", "Loại MIME", "VARCHAR(100)", "NO", ""],
                ["uploaded_at", "Thời điểm tải lên", "DATETIME", "NO", "Default SYSUTCDATETIME()"],
                ["uploaded_by_id", "Người tải lên", "INT", "NO", "FK -> employees(id)"],
            ]
        result.append((name, rows))
    return result


ARCH_DIAGRAMS = [
    ("01_System_Context.png", "Sơ đồ ngữ cảnh hệ thống"),
    ("02_Container_Architecture.png", "Kiến trúc container"),
    ("03_Backend_Layers.png", "Kiến trúc phân lớp Backend"),
    ("04_Domain_Model.png", "Mô hình miền"),
    ("05_Deployment.png", "Sơ đồ triển khai"),
    ("06_Frontend_Layers.png", "Kiến trúc phân lớp Frontend"),
    ("07_Data_Flow.png", "Luồng dữ liệu tổng quát"),
    ("08_Authentication_Flow.png", "Luồng xác thực và phân quyền"),
    ("09_Realtime_Event_Flow.png", "Luồng sự kiện thời gian thực"),
    ("10_Functional_Decomposition.png", "Phân rã chức năng"),
    ("11_Use_Case_Tong_Quat.png", "Use Case tổng quát"),
    ("12_Activity_Login.png", "Activity - Đăng nhập"),
    ("13_Activity_Project_Creation.png", "Activity - Tạo dự án"),
    ("14_Activity_Task_Lifecycle.png", "Activity - Vòng đời Task"),
    ("15_Activity_Leave_Approval.png", "Activity - Duyệt nghỉ phép"),
    ("16_Sequence_Create_Task.png", "Sequence - Tạo Task"),
    ("17_Sequence_Notification.png", "Sequence - Thông báo"),
    ("18_ERD_Core.png", "ERD rút gọn"),
    ("19_Module_Dependencies.png", "Phụ thuộc module"),
    ("20_Backup_Recovery.png", "Sao lưu và khôi phục"),
    ("21_Observability.png", "Giám sát vận hành"),
]


def build_architecture_report():
    doc = academic_document("KIẾN TRÚC PHẦN MỀM | TASKSYNCENTERPRISE")
    add_academic_cover(
        doc,
        "Kiến trúc phần mềm",
        "TaskSyncEnterprise",
        "Báo cáo kiến trúc tổng hợp - không tổ chức nội dung theo Sprint",
    )
    figures = [title for _, title in ARCH_DIAGRAMS] + [
        "ERD từ tài liệu cơ sở dữ liệu",
        "Các giao diện quản trị và nhân viên",
    ]
    tables = [
        "Stakeholder và mối quan tâm",
        "Thuộc tính chất lượng",
        "Danh mục quyết định kiến trúc",
        "Danh mục module",
        "Mô tả 19 bảng cơ sở dữ liệu lõi",
        "Danh mục giao diện",
        "Rủi ro và phương án xử lý",
    ]
    add_front_matter(doc, "Kiến trúc", figures, tables)
    add_part_title(doc, "1", "Tổng quan và động lực kiến trúc", first=True)
    doc.add_heading("1. Giới thiệu hệ thống", level=1)
    base.add_body(
        doc,
        "TaskSyncEnterprise là nền tảng quản lý công việc và nhân sự doanh nghiệp. Hệ thống hợp nhất quản lý tổ chức, dự án, Backlog, "
        "Task, cộng tác, tệp, nghỉ phép, làm thêm giờ, thông báo, báo cáo và audit trong một sản phẩm web."
    )
    doc.add_heading("2. Phạm vi kiến trúc", level=1)
    base.add_body(
        doc,
        "Báo cáo mô tả kiến trúc logic, runtime, dữ liệu, bảo mật, triển khai và vận hành của toàn bộ sản phẩm. "
        "Nội dung không liệt kê Sprint; các Sprint chỉ là phương thức triển khai trong hồ sơ Agile."
    )
    doc.add_heading("3. Stakeholder và mối quan tâm", level=1)
    stakeholders = [
        ("Người dùng", "Dễ sử dụng, phản hồi nhanh, dữ liệu chính xác."),
        ("Quản lý", "Nhìn thấy tiến độ, nguồn lực, rủi ro và trách nhiệm."),
        ("Quản trị viên", "Phân quyền, cấu hình, audit và hỗ trợ sự cố."),
        ("Nhóm phát triển", "Ranh giới module rõ, dễ kiểm thử và thay đổi."),
        ("Nhóm vận hành", "Khả năng giám sát, sao lưu, khôi phục và mở rộng."),
    ]
    base.add_table(doc, ["Stakeholder", "Mối quan tâm"], stakeholders, [2400, 6960], font_size=9.4)
    doc.add_heading("4. Thuộc tính chất lượng", level=1)
    qualities = [
        ("Bảo mật", "RBAC tại Backend; token; audit; kiểm tra tệp; nguyên tắc ít quyền nhất."),
        ("Hiệu năng", "Cache-aside Redis; truy vấn có chỉ mục; phân trang; tác vụ nặng tách khỏi request."),
        ("Khả dụng", "Health/readiness; retry thông báo; giám sát; backup và restore thử nghiệm."),
        ("Khả năng bảo trì", "Modular monolith; phân lớp; schema rõ; test tự động."),
        ("Nhất quán dữ liệu", "SQL Server là source of truth; event chỉ phát sau commit."),
        ("Khả năng mở rộng", "Tách Notification/Reporting/AI khi tải tăng; object storage khi nhiều replica."),
    ]
    base.add_table(doc, ["Thuộc tính", "Chiến lược"], qualities, [2200, 7160], font_size=9.2)
    doc.add_heading("5. Phong cách và quyết định kiến trúc", level=1)
    decisions = [
        ("ADR-01", "Modular Monolith", "Đơn giản triển khai nhưng vẫn giữ ranh giới module."),
        ("ADR-02", "React SPA + REST/WebSocket", "Tách trải nghiệm người dùng và API; hỗ trợ realtime."),
        ("ADR-03", "SQL Server là nguồn dữ liệu chuẩn", "Bảo đảm transaction và quan hệ nghiệp vụ."),
        ("ADR-04", "Redis cache-aside", "Tăng tốc; không thay thế dữ liệu chuẩn."),
        ("ADR-05", "Event phát sau commit", "Tránh UI nhận sự kiện cho giao dịch đã rollback."),
        ("ADR-06", "RBAC tại Backend", "Không phụ thuộc vào việc ẩn/hiện điều khiển ở Frontend."),
    ]
    base.add_table(doc, ["Mã", "Quyết định", "Lý do"], decisions, [1100, 3000, 5260], font_size=9.0)

    add_part_title(doc, "2", "Các góc nhìn kiến trúc")
    common_explanation_done = False
    for index, (filename, title) in enumerate(ARCH_DIAGRAMS, 1):
        if index > 1:
            doc.add_page_break()
        doc.add_heading(f"{index}. {title}", level=1)
        base.add_diagram(doc, ARCH_DIAGRAM_DIR / filename, f"Hình {index}. {title}")
        if not common_explanation_done:
            base.add_note(
                doc,
                "Quy ước đọc sơ đồ",
                "Khối màu xanh biểu diễn thành phần ứng dụng hoặc miền nghiệp vụ; khối vàng là điểm kiểm soát/giao dịch; "
                "khối đỏ là dependency/rủi ro cần chú ý; mũi tên biểu diễn lời gọi, phụ thuộc hoặc luồng dữ liệu. "
                "Quy ước này dùng chung cho toàn bộ chương và không lặp lại ở từng hình.",
            )
            common_explanation_done = True
        explanations = {
            "Sơ đồ ngữ cảnh hệ thống": "Xác định tác nhân Admin, Manager, Employee, ranh giới nền tảng và tích hợp thông báo ngoài hệ thống.",
            "Kiến trúc container": "Thể hiện Nginx, Frontend, Backend, SQL Server, Redis, lưu trữ tệp và giám sát.",
            "Kiến trúc phân lớp Backend": "Router là biên validation/authorization; Service điều phối nghiệp vụ; CRUD/ORM quản lý truy cập dữ liệu.",
            "Mô hình miền": "Project là phạm vi gốc; Backlog/Sprint/Task tổ chức kế hoạch; Employee và Project Member xác định người thực hiện.",
            "Sơ đồ triển khai": "Nginx là điểm vào production; Backend phụ thuộc SQL Server, Redis và volume; Alembic quản lý schema.",
            "Kiến trúc phân lớp Frontend": "Trang tính năng sử dụng component chung và TanStack Query; Axios là lớp giao tiếp API.",
            "Luồng dữ liệu tổng quát": "Mọi thay đổi đi từ UI qua Backend tới SQL Server rồi phản hồi về UI; cache không trở thành nguồn chuẩn.",
            "Luồng xác thực và phân quyền": "Mật khẩu được xác minh, token được cấp và mọi endpoint nhạy cảm tiếp tục kiểm tra RBAC.",
            "Luồng sự kiện thời gian thực": "Sau commit, sự kiện được phát qua WebSocket để client khác làm mới query cache.",
            "Phân rã chức năng": "Nhóm năng lực sản phẩm thành các miền ổn định, dùng làm ranh giới module và tổ chức yêu cầu.",
            "Use Case tổng quát": "Liên kết ba nhóm tác nhân với các năng lực chính mà không phụ thuộc vào kế hoạch Sprint.",
            "Activity - Đăng nhập": "Mô tả các bước validate, xác thực, kiểm tra trạng thái và tạo phiên.",
            "Activity - Tạo dự án": "Mô tả transaction tạo Project, gán thành viên, tạo Backlog ban đầu và audit.",
            "Activity - Vòng đời Task": "Mô tả trạng thái và cổng chất lượng trước khi Task đạt Done.",
            "Activity - Duyệt nghỉ phép": "Mô tả kiểm tra số dư, quyền duyệt, cập nhật số dư và thông báo.",
            "Sequence - Tạo Task": "Làm rõ trách nhiệm của React Form, Router, Service, ORM, SQL Server và WebSocket.",
            "Sequence - Thông báo": "Làm rõ việc lưu thông báo, gửi realtime/email và theo dõi kết quả.",
            "ERD rút gọn": "Tập trung quan hệ lõi giữa tổ chức, nhân viên, dự án, công việc, cộng tác và audit.",
            "Phụ thuộc module": "Identity/Organization là nền; Project/Task là lõi; Collaboration/Notification/Reporting là module hỗ trợ.",
            "Sao lưu và khôi phục": "Backup chỉ có giá trị khi checksum và restore thử nghiệm được thực hiện định kỳ.",
            "Giám sát vận hành": "Log và metrics được tập trung để dashboard và cảnh báo hỗ trợ phát hiện sự cố.",
        }
        base.add_body(doc, f"Ý nghĩa: {explanations[title]}", bold_lead="Ý nghĩa:")
        base.add_body(
            doc,
            "Ứng dụng trong thiết kế: sơ đồ là cơ sở xác định trách nhiệm thành phần, điểm kiểm soát và phạm vi kiểm thử tích hợp.",
            bold_lead="Ứng dụng trong thiết kế:",
        )

    add_part_title(doc, "3", "Thiết kế dữ liệu")
    doc.add_heading("1. Nguyên tắc thiết kế dữ liệu", level=1)
    base.add_bullets(
        doc,
        [
            "SQL Server là nguồn dữ liệu chuẩn; Redis chỉ lưu cache/trạng thái phụ trợ.",
            "Khóa ngoại bảo vệ quan hệ; các bảng quan hệ nhiều-nhiều dùng bảng nối.",
            "Xóa mềm dùng cho đối tượng nghiệp vụ cần truy vết; dữ liệu audit không bị sửa tùy ý.",
            "Ngày giờ chuẩn hóa UTC tại Backend; giao diện chuyển đổi theo múi giờ người dùng.",
            "Các trường tìm kiếm/lọc thường xuyên cần chỉ mục phù hợp và truy vấn có phân trang.",
        ],
    )
    if (MEDIA_DIR / "image1.png").exists():
        base.add_diagram(doc, MEDIA_DIR / "image1.png", "Hình. ERD tổng hợp từ tài liệu Cac_bang_table.docx", width=6.1)
    base.add_note(
        doc,
        "Hiệu chỉnh nguồn",
        "Bảng task_attachments trong tài liệu tham khảo bị chép nhầm cấu trúc của bảng tasks. "
        "Bản dưới đây đã đối chiếu app/models/task_attachment.py và migration để sửa lại các cột.",
        "FFF2CC",
    )
    for index, (name, rows) in enumerate(schema_tables(), 1):
        doc.add_page_break()
        doc.add_heading(f"{index}. Bảng {name}", level=1)
        purpose = {
            "roles": "Lưu vai trò dùng cho RBAC.",
            "departments": "Lưu cơ cấu phòng ban.",
            "teams": "Lưu nhóm trực thuộc phòng ban.",
            "employees": "Lưu tài khoản và hồ sơ nhân viên.",
            "projects": "Lưu thông tin dự án và tiến độ tổng hợp.",
            "project_members": "Liên kết nhân viên với dự án.",
            "tasks": "Lưu công việc và trạng thái thực hiện.",
            "task_assignments": "Lưu phân công nhiều người cho Task.",
            "task_comments": "Lưu trao đổi trên Task.",
            "task_checklists": "Lưu danh sách kiểm tra của Task.",
            "task_attachments": "Lưu metadata tệp và liên kết tới Task/thảo luận/feedback.",
            "notifications": "Lưu thông báo nghiệp vụ.",
            "notification_preferences": "Lưu lựa chọn kênh nhận thông báo.",
            "notification_logs": "Lưu kết quả gửi và retry.",
            "user_sessions": "Lưu phiên đăng nhập.",
            "refresh_tokens": "Lưu refresh token có thể thu hồi.",
            "token_blacklist": "Lưu token bị vô hiệu hóa.",
            "vacations": "Lưu yêu cầu nghỉ phép và phê duyệt.",
            "audit_logs": "Lưu dấu vết thao tác và thay đổi dữ liệu.",
        }[name]
        base.add_body(doc, f"Mục đích: {purpose}", bold_lead="Mục đích:")
        headers = rows[0]
        body = rows[1:]
        base.add_table(doc, headers, body, [1400, 3300, 1500, 900, 2260], font_size=7.6)
        media_index = index + 1
        media = MEDIA_DIR / f"image{media_index}.png"
        if media.exists() and name != "task_attachments":
            base.add_diagram(
                doc,
                media,
                f"Hình. Cấu trúc bảng {name}",
                width=image_width_for_page(media),
            )
        else:
            add_placeholder(
                doc,
                f"database_{name}.png",
                f"Ảnh chụp cấu trúc bảng {name} từ SQL Server Management Studio sau khi migration mới nhất được áp dụng.",
            )
        base.add_body(
            doc,
            "Lưu ý triển khai: kiểm tra khóa ngoại, index, default constraint và hành vi xóa/cập nhật bằng migration; "
            "không chỉ dựa vào ảnh thiết kế.",
        )

    doc.add_page_break()
    doc.add_heading("20. Các bảng bổ sung trong mã nguồn hiện tại", level=1)
    extra_models = [
        ("backlog_items", "Lưu Product Backlog Item liên kết Project/Topic/Sprint."),
        ("sprints", "Lưu Sprint thuộc Project."),
        ("sprint_members", "Lưu thành viên và capacity của Sprint."),
        ("sprint_snapshots", "Lưu ảnh chụp chỉ số Sprint phục vụ báo cáo."),
        ("discussion_topics", "Lưu chủ đề thảo luận."),
        ("discussion_replies", "Lưu phản hồi thảo luận."),
        ("user_feedback", "Lưu phản hồi người dùng."),
    ]
    base.add_table(doc, ["Bảng", "Vai trò"], extra_models, [2800, 6560], font_size=9.2)

    add_part_title(doc, "4", "Thiết kế giao diện và trải nghiệm")
    ui_captions = [
        ("image20.png", "Dashboard quản trị", "Tổng hợp KPI, tiến độ, phân bổ trạng thái và cảnh báo."),
        ("image21.png", "Trang quản lý nhân viên", "Danh sách, lọc, tìm kiếm và thao tác hồ sơ nhân viên."),
        ("image22.png", "Trang quản lý dự án", "Theo dõi các dự án và trạng thái tổng hợp."),
        ("image23.png", "Trang quản lý Task", "Hiển thị Task theo cột trạng thái và người phụ trách."),
        ("image24.png", "Danh sách nhân viên", "Bảng dữ liệu nhân viên dành cho quản trị."),
        ("image25.png", "Trang quản lý phòng ban", "Quản lý cơ cấu và nhân sự theo phòng ban."),
        ("image26.png", "Lịch công việc", "Tổng hợp deadline và sự kiện theo thời gian."),
        ("image27.png", "Trung tâm thông báo", "Hiển thị thông báo, trạng thái đọc và ưu tiên."),
        ("image28.png", "Cài đặt hệ thống", "Cấu hình tài khoản, bảo mật và tùy chọn hệ thống."),
        ("image29.png", "Nhật ký hệ thống", "Tra cứu sự kiện và thay đổi phục vụ audit."),
        ("image30.png", "Trang hồ sơ cá nhân", "Thông tin tài khoản và các mục liên quan cá nhân."),
        ("image31.png", "Dashboard nhân viên", "Công việc, deadline và tiến độ của người dùng."),
        ("image32.png", "Dự án của nhân viên", "Các dự án mà nhân viên tham gia."),
        ("image33.png", "Task của nhân viên", "Kanban cá nhân và công việc được giao."),
        ("image34.png", "Trang đăng nhập", "Điểm vào hệ thống và xử lý xác thực."),
    ]
    for index, (file, title, description) in enumerate(ui_captions, 1):
        if index > 1:
            doc.add_page_break()
        doc.add_heading(f"{index}. {title}", level=1)
        path = MEDIA_DIR / file
        if path.exists():
            base.add_diagram(doc, path, f"Hình. {title}", width=6.15)
        else:
            add_placeholder(doc, file, description)
        base.add_body(doc, f"Mục tiêu: {description}", bold_lead="Mục tiêu:")
        base.add_bullets(
            doc,
            [
                "Tất cả dữ liệu phải được giới hạn theo quyền và phạm vi tổ chức/dự án.",
                "Trạng thái tải, rỗng, lỗi và thành công phải rõ ràng, không làm mất dữ liệu người dùng.",
                "Thao tác nguy hiểm phải xác nhận và có Audit Log khi cần.",
            ],
        )
    add_placeholder(doc, "custom_report_builder.png", "Màn hình tạo báo cáo dự án tùy chỉnh theo bộ lọc và trường dữ liệu.")
    add_placeholder(doc, "monitoring_dashboard.png", "Dashboard Prometheus/Grafana thể hiện API latency, error rate, Redis và SQL Server.")
    add_placeholder(doc, "mobile_responsive_views.png", "Ảnh giao diện responsive trên màn hình nhỏ cho Dashboard, Project và Task.")

    add_part_title(doc, "5", "Bảo mật, kiểm thử và vận hành")
    doc.add_heading("1. Kiến trúc bảo mật", level=1)
    base.add_bullets(
        doc,
        [
            "Mật khẩu lưu dưới dạng hash mạnh; không ghi token/mật khẩu vào log.",
            "Access token ngắn hạn, refresh token có thể thu hồi và blacklist khi đăng xuất/sự cố.",
            "RBAC được kiểm tra tại Router/Service; truy vấn dữ liệu tiếp tục áp dụng scope.",
            "Upload kiểm tra MIME, dung lượng, tên tệp, quyền tải và vị trí lưu trữ.",
            "Audit ghi actor, action, entity, thời gian, IP và trước/sau khi thay đổi khi phù hợp.",
        ],
    )
    doc.add_heading("2. Chiến lược kiểm thử", level=1)
    testing = [
        ("Unit test", "Service, validation, policy, mapping và helper."),
        ("Integration test", "Router + database; transaction; quyền; migration."),
        ("Frontend test", "Component, hook, form validation và state."),
        ("E2E", "Đăng nhập, tạo dự án, tạo/gán Task, upload, nghỉ phép và thông báo."),
        ("Security test", "Broken access control, token, upload, input, audit."),
        ("Performance test", "API latency, truy vấn danh sách, WebSocket và báo cáo."),
    ]
    base.add_table(doc, ["Mức kiểm thử", "Phạm vi"], testing, [2400, 6960], font_size=9.3)
    doc.add_heading("3. Triển khai và cấu hình", level=1)
    base.add_body(
        doc,
        "Môi trường production dùng Nginx làm điểm vào, Frontend và Backend chạy trong container, SQL Server và Redis có volume. "
        "Secret phải cấp qua biến môi trường/secret manager; không commit vào mã nguồn. Alembic chịu trách nhiệm migration schema."
    )
    doc.add_heading("4. Health, readiness và observability", level=1)
    base.add_body(
        doc,
        "Liveness xác nhận tiến trình còn sống; readiness xác nhận dependency cần thiết có thể phục vụ. "
        "Structured log, correlation ID, metrics latency/error rate và dashboard cảnh báo giúp rút ngắn thời gian phát hiện sự cố."
    )
    doc.add_heading("5. Sao lưu và khôi phục", level=1)
    base.add_body(
        doc,
        "Backup SQL Server và tệp đính kèm phải cùng chính sách lưu giữ. Nhóm cần định nghĩa RPO/RTO, mã hóa bản sao, kiểm tra checksum "
        "và thực hiện restore thử nghiệm. Một bản backup chưa restore thử không được xem là bằng chứng phục hồi."
    )
    doc.add_heading("6. Rủi ro và nợ kỹ thuật", level=1)
    risks = [
        ("Redis/readiness chưa ổn định tại local", "Khởi động dependency, kiểm tra timeout và chiến lược degraded mode."),
        ("Docker engine chưa sẵn sàng", "Kiểm tra pipeline/build trên môi trường có Docker."),
        ("Dependency advisory", "Chốt phiên bản, chạy scanner và cập nhật có kiểm thử."),
        ("Summary SP không khớp", "Dùng tổng từng dòng và thêm kiểm tra công thức."),
        ("Ảnh/evidence chưa đầy đủ", "Bổ sung theo danh sách placeholder trước khi nộp."),
        ("Tên/MSSV ở nguồn không thống nhất", "Dùng Co_su_Ly_Thuyet làm danh sách chính và xác nhận lại với nhóm."),
    ]
    base.add_table(doc, ["Rủi ro", "Xử lý"], risks, [3500, 5860], font_size=9.0)
    doc.add_heading("7. Lộ trình mở rộng", level=1)
    base.add_bullets(
        doc,
        [
            "Tách Notification thành worker/service khi lưu lượng event tăng.",
            "Dùng outbox pattern để bảo đảm event không mất giữa commit và publish.",
            "Chuyển upload sang object storage khi triển khai nhiều replica.",
            "Tách Reporting/AI khỏi request API để tránh ảnh hưởng latency.",
            "Áp dụng SSO/OIDC và policy engine khi yêu cầu doanh nghiệp tăng.",
        ],
    )
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    base.add_bullets(
        doc,
        [
            "docs/architecture/SYSTEM_ARCHITECTURE.md và DOMAIN_MODEL.md.",
            "Cac_bang_table.docx - bảng dữ liệu và ảnh giao diện.",
            "Backend SQLAlchemy models và Alembic migrations.",
            "Docker Compose, Nginx và tài liệu vận hành trong repository.",
        ],
    )
    base.save_doc(doc, ARCH_DIR / "Bao_cao_Kien_truc_phan_mem_TaskSyncEnterprise.docx")


def business_days(start: date, end: date):
    current = start
    days = []
    while current <= end:
        if current.weekday() < 5:
            days.append(current)
        current += timedelta(days=1)
    return days


def decomposed_tasks(item, sprint_number):
    owners = [TEAM[4][1], TEAM[1][1], TEAM[2][1], TEAM[3][1]]
    return [
        (f"{item['id']}-01", "Phân tích yêu cầu và cập nhật tiêu chí chấp nhận", owners[0]),
        (f"{item['id']}-02", "Thiết kế API, dữ liệu và quy tắc phân quyền", owners[1]),
        (f"{item['id']}-03", "Xây dựng giao diện và tích hợp API", owners[2]),
        (f"{item['id']}-04", "Kiểm thử, ghi bằng chứng và cập nhật tài liệu", owners[3]),
    ]


def build_daily_scrum(sprint_number: int, items: list[dict], start: date, end: date):
    doc = academic_document(f"DAILY SCRUM {sprint_number} | TASKSYNCENTERPRISE")
    add_academic_cover(
        doc,
        "Mô hình phát triển phần mềm Agile",
        f"Daily Scrum {sprint_number}",
        f"Biên bản riêng cho Sprint {sprint_number} - {start.strftime('%d/%m/%Y')} đến {end.strftime('%d/%m/%Y')}",
    )
    base.add_note(
        doc,
        "Trạng thái hồ sơ",
        "Nội dung được phân bổ từ Product Backlog và tài liệu nguồn. Các cuộc họp sau ngày 29/07/2026 hoặc chưa có ảnh/video "
        "được ghi là kế hoạch dự kiến; nhóm cần xác nhận người tham dự, blocker và kết quả thực tế trước khi ký.",
        "FFF2CC",
    )
    doc.add_heading("Mục tiêu Sprint", level=1)
    base.add_body(
        doc,
        "Sprint 1 tập trung nền tảng xác thực, tổ chức, dự án, Backlog, Kanban và cộng tác."
        if sprint_number == 1
        else "Sprint 2 mở rộng Task, tài liệu, nghỉ phép/OT, thông báo, báo cáo và kiểm toán."
    )
    base.add_table(doc, ["US", "Nội dung", "SP"], [(x["id"], x["title"], x["sp"]) for x in items], [900, 7060, 1400], font_size=8.5)
    days = business_days(start, end)
    flat_tasks = [(item, task) for item in items for task in decomposed_tasks(item, sprint_number)]
    for day_index, meeting_day in enumerate(days, 1):
        doc.add_page_break()
        is_future = meeting_day > date(2026, 7, 29)
        doc.add_heading(f"Daily Scrum {day_index:02d} - {meeting_day.strftime('%d/%m/%Y')}", level=1)
        base.add_table(
            doc,
            ["Thông tin", "Nội dung"],
            [
                ("Thời lượng", "15 phút"),
                ("Điều phối", "Scrum Master / người được nhóm phân công"),
                ("Thành phần", ", ".join(name for _, name, _, _ in TEAM)),
                ("Trạng thái", "Kế hoạch dự kiến - cập nhật sau họp" if is_future else "Dự thảo tổng hợp - cần nhóm xác nhận"),
                ("Sprint Goal", "Kiểm tra tiến độ trong 24 giờ và xử lý blocker sớm."),
            ],
            [2200, 7160],
            font_size=9.2,
        )
        rows = []
        for member_index, (_, name, _, role) in enumerate(TEAM):
            task_index = (day_index * len(TEAM) + member_index) % len(flat_tasks)
            item, task = flat_tasks[task_index]
            previous_item, previous_task = flat_tasks[(task_index - len(TEAM)) % len(flat_tasks)]
            rows.append(
                (
                    name,
                    f"{previous_item['id']}: {previous_task[1]}",
                    f"{item['id']}: {task[1]}",
                    "Chưa có blocker được xác nhận; cập nhật theo nội dung họp thực tế.",
                )
            )
        base.add_table(
            doc,
            ["Thành viên", "Đã làm", "Sẽ làm", "Trở ngại"],
            rows,
            [1900, 2500, 2500, 2460],
            font_size=8.0,
        )
        doc.add_heading("Quyết định và hành động sau họp", level=2)
        base.add_bullets(
            doc,
            [
                "Cập nhật trạng thái Task và ước lượng còn lại ngay sau cuộc họp.",
                "Blocker có người chịu trách nhiệm và thời hạn xử lý; không để ghi chú chung chung.",
                "Nếu phạm vi thay đổi, Product Owner cập nhật Backlog và thông báo cho cả nhóm.",
                "Liên kết evidence: commit/PR, test, ảnh giao diện hoặc biên bản quyết định.",
            ],
        )
        add_placeholder(
            doc,
            f"daily_scrum_{sprint_number}_{meeting_day.isoformat()}.png",
            "Ảnh chụp hoặc đường dẫn video của cuộc họp; cần thấy ngày họp và thành viên tham dự.",
        )
        doc.add_heading("Xác nhận", level=2)
        base.add_table(
            doc,
            ["Vai trò", "Họ tên / Chữ ký", "Ngày xác nhận"],
            [("Scrum Master", "", ""), ("Đại diện Development Team", "", "")],
            [2400, 4560, 2400],
            font_size=9.0,
        )
    base.save_doc(doc, AGILE_DIR / f"Daily_Scrum_{sprint_number}_TaskSyncEnterprise.docx")


def build_diagram_guide():
    doc = academic_document("THUYẾT MINH SƠ ĐỒ | TASKSYNCENTERPRISE")
    add_academic_cover(
        doc,
        "Kiến trúc phần mềm và Agile",
        "Thuyết minh bộ sơ đồ",
        "Sơ đồ kiến trúc, Use Case, Activity, Sequence và Scrum",
    )
    doc.add_heading("Quy ước giải thích chung", level=1)
    base.add_body(
        doc,
        "Các sơ đồ dùng cùng ngôn ngữ màu và hướng mũi tên. Phần giải thích chung này áp dụng cho toàn bộ bộ sơ đồ; "
        "mỗi mục sau chỉ nêu mục tiêu và điểm cần đọc riêng, tránh lặp lại nội dung."
    )
    base.add_bullets(
        doc,
        [
            "Khối xanh: thành phần ứng dụng hoặc miền nghiệp vụ.",
            "Khối vàng: bước kiểm tra, giao dịch hoặc quyết định.",
            "Khối đỏ: dependency/rủi ro hoặc điểm cần kiểm soát.",
            "Mũi tên: lời gọi, phụ thuộc, luồng dữ liệu hoặc chuyển trạng thái.",
            "Mỗi sơ đồ có PNG để chèn Word và MMD để chỉnh sửa bằng Mermaid.",
        ],
    )
    descriptions = [(i, "Kiến trúc phần mềm", title, filename) for i, (filename, title) in enumerate(ARCH_DIAGRAMS, 1)]
    agile_start = len(descriptions) + 1
    agile_items = [
        ("01_Phan_ra_chuc_nang.png", "Phân rã chức năng phục vụ Product Backlog"),
        ("02_Use_Case_Tong_Quat.png", "Use Case tổng quát"),
        ("03_Activity_Task_Lifecycle.png", "Activity vòng đời Task"),
        ("04_Activity_Sprint.png", "Activity chu trình Sprint"),
        ("05_Scrum_Workflow.png", "Scrum Workflow"),
    ]
    descriptions.extend(
        (agile_start + i, "Agile", title, filename) for i, (filename, title) in enumerate(agile_items)
    )
    base.add_table(
        doc,
        ["STT", "Nhóm", "Tên sơ đồ", "File"],
        descriptions,
        [700, 1900, 4300, 2460],
        font_size=8.5,
    )
    for index, group, title, filename in descriptions:
        doc.add_page_break()
        doc.add_heading(f"{index}. {title}", level=1)
        folder = ARCH_DIAGRAM_DIR if group == "Kiến trúc phần mềm" else AGILE_DIAGRAM_DIR
        base.add_diagram(doc, folder / filename, f"Hình {index}. {title}")
        base.add_body(
            doc,
            "Cách sử dụng: đặt ngay sau đoạn mô tả phạm vi liên quan, tham chiếu bằng số hình và giải thích quyết định/ràng buộc mà sơ đồ làm rõ.",
        )
    base.save_doc(doc, DIAGRAM_DIR / "Thuyet_minh_So_do_TaskSyncEnterprise.docx")


def build_readme():
    content = """# Bộ hồ sơ Report - TaskSyncEnterprise

## Tài liệu chính

- `Môn kiến trúc phần mềm/Bao_cao_Kien_truc_phan_mem_TaskSyncEnterprise.docx`: báo cáo kiến trúc tổng hợp, không tổ chức theo Sprint.
- `Agile/Bao_cao_Agile_Scrum_TaskSyncEnterprise.docx`: báo cáo Agile chi tiết theo cấu trúc PDF mẫu 106 trang.
- `Agile/Daily_Scrum_1_TaskSyncEnterprise.docx`: biên bản riêng Sprint 1.
- `Agile/Daily_Scrum_2_TaskSyncEnterprise.docx`: biên bản riêng Sprint 2.
- `Agile/Agile_Project_Management_TaskSyncEnterprise.xlsx`: Product Backlog và quản lý phạm vi.
- `Agile/Agile_Tracking_TaskSyncEnterprise.xlsx`: phân rã công việc và theo dõi thực thi.
- `Sơ đồ Diagram/`: PNG, Mermaid và tài liệu thuyết minh sơ đồ.

## Quy ước Sprint

- Sprint cũ 1 + 2 -> Sprint 1: 02/07-30/07/2026, US01-US18, 107 SP.
- Sprint cũ 3 + 4 -> Sprint 2: 30/07-27/08/2026, US19-US40, 135 SP.
- Tổng chi tiết theo 40 User Stories: 242 SP. Summary nguồn ghi 228 SP nhưng không khớp tổng từng dòng.

## Việc nhóm cần cập nhật

1. Tìm các khung màu vàng `ẢNH CẦN BỔ SUNG` trong Word và thay bằng ảnh/video/evidence thật.
2. Xác nhận họ tên, MSSV, vai trò và chữ ký.
3. Xác nhận trạng thái Done bằng commit/PR, test, ảnh demo hoặc biên bản.
4. Các cuộc họp sau 29/07/2026 đang được ghi là kế hoạch dự kiến, không phải sự kiện đã xảy ra.
5. Mở Word, chọn toàn bộ và nhấn F9 nếu cần cập nhật mục lục/số trang sau khi thay ảnh.
"""
    (REPORT / "README_Bo_ho_so.md").write_text(content, encoding="utf-8")


def main():
    build_extended_diagrams()
    build_agile_report()
    build_architecture_report()
    build_daily_scrum(1, SPRINT_1, date(2026, 7, 2), date(2026, 7, 30))
    build_daily_scrum(2, SPRINT_2, date(2026, 7, 30), date(2026, 8, 27))
    build_diagram_guide()
    build_readme()
    print("Built report package version 2.")


if __name__ == "__main__":
    main()
