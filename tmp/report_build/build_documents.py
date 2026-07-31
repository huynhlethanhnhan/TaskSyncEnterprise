from __future__ import annotations

import math
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT = ROOT / "Report"
ARCH_DIR = REPORT / "Môn kiến trúc phần mềm"
AGILE_DIR = REPORT / "Agile"
DIAGRAM_DIR = REPORT / "Sơ đồ Diagram"
ARCH_DIAGRAM_DIR = DIAGRAM_DIR / "Kiến trúc phần mềm"
AGILE_DIAGRAM_DIR = DIAGRAM_DIR / "Agile"

for folder in (ARCH_DIR, AGILE_DIR, ARCH_DIAGRAM_DIR, AGILE_DIAGRAM_DIR):
    folder.mkdir(parents=True, exist_ok=True)


NAVY = "16324F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "0F766E"
GOLD = "B78318"
GRAY = "5B6573"
LIGHT = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
WHITE = "FFFFFF"
BLACK = "111827"
RED = "9B1C1C"
GREEN = "1F6B45"

TEAM = [
    ("A", "Huỳnh Lê Thành Nhân", "Product Owner / Full Stack Lead"),
    ("B", "Nguyễn Đức Mạnh", "Scrum Master / Backend Developer"),
    ("C", "Nguyễn Lê Huy Hoàng", "Frontend Developer"),
    ("D", "Phạm Anh Tuấn", "QA / Test Engineer"),
    ("E", "Nguyễn Anh Tuấn", "Research & Documentation"),
]


def font(size=28, bold=False):
    name = "arialbd.ttf" if bold else "arial.ttf"
    path = Path(r"C:\Windows\Fonts") / name
    return ImageFont.truetype(str(path), size=size)


def wrap_px(draw, text, fnt, max_width):
    words = text.split()
    lines, current = [], ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def box(draw, xy, text, fill="#FFFFFF", outline="#2E74B5", radius=18, fnt=None, text_color="#111827"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    fnt = fnt or font(26, True)
    lines = wrap_px(draw, text, fnt, x2 - x1 - 30)
    line_h = fnt.getbbox("Ag")[3] + 7
    start_y = y1 + (y2 - y1 - line_h * len(lines)) / 2
    for i, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=fnt)
        tx = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((tx, start_y + i * line_h), line, fill=text_color, font=fnt)


def arrow(draw, start, end, color="#52677C", width=5, label=None):
    draw.line([start, end], fill=color, width=width)
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    for delta in (2.55, -2.55):
        p = (end[0] + 18 * math.cos(ang + delta), end[1] + 18 * math.sin(ang + delta))
        draw.line([end, p], fill=color, width=width)
    if label:
        mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        draw.rounded_rectangle((mx - 90, my - 18, mx + 90, my + 18), radius=8, fill="#FFFFFF")
        bbox = draw.textbbox((0, 0), label, font=font(18))
        draw.text((mx - (bbox[2] - bbox[0]) / 2, my - 12), label, fill=color, font=font(18))


def ellipse(draw, xy, text, fill="#F8FBFF", outline="#2E74B5"):
    draw.ellipse(xy, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    fnt = font(23, True)
    lines = wrap_px(draw, text, fnt, x2 - x1 - 24)
    h = fnt.getbbox("Ag")[3] + 5
    sy = y1 + (y2 - y1 - h * len(lines)) / 2
    for i, line in enumerate(lines):
        bb = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + (x2 - x1 - bb[2] + bb[0]) / 2, sy + i * h), line, fill="#111827", font=fnt)


def actor(draw, x, y, label):
    draw.ellipse((x - 22, y - 90, x + 22, y - 46), outline="#16324F", width=4)
    draw.line((x, y - 46, x, y + 30), fill="#16324F", width=4)
    draw.line((x - 42, y - 10, x + 42, y - 10), fill="#16324F", width=4)
    draw.line((x, y + 30, x - 34, y + 78), fill="#16324F", width=4)
    draw.line((x, y + 30, x + 34, y + 78), fill="#16324F", width=4)
    bb = draw.textbbox((0, 0), label, font=font(22, True))
    draw.text((x - (bb[2] - bb[0]) / 2, y + 88), label, fill="#16324F", font=font(22, True))


def canvas(title, subtitle):
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 105), fill="#" + NAVY)
    draw.text((55, 25), title, fill="white", font=font(35, True))
    draw.text((55, 112), subtitle, fill="#" + GRAY, font=font(21))
    return img, draw


def save_diagram(path, title, subtitle, painter):
    img, draw = canvas(title, subtitle)
    painter(draw)
    img.save(path, quality=95)


def build_diagrams():
    save_diagram(
        ARCH_DIAGRAM_DIR / "01_System_Context.png",
        "Sơ đồ ngữ cảnh hệ thống",
        "Tác nhân và ranh giới của nền tảng TaskSyncEnterprise",
        lambda d: (
            actor(d, 120, 380, "Admin"),
            actor(d, 120, 600, "Manager"),
            actor(d, 120, 750, "Employee"),
            box(d, (330, 300, 610, 620), "React Web Application", "#E8F1FA", "#2E74B5"),
            box(d, (740, 260, 1170, 660), "TaskSyncEnterprise\nREST API + WebSocket", "#EEF8F5", "#0F766E", fnt=font(31, True)),
            box(d, (1280, 330, 1530, 530), "Email / Push Provider", "#FFF7E8", "#B78318"),
            arrow(d, (165, 380), (330, 380)),
            arrow(d, (165, 600), (330, 520)),
            arrow(d, (165, 750), (330, 600)),
            arrow(d, (610, 460), (740, 460), label="HTTPS / WS"),
            arrow(d, (1170, 430), (1280, 430), label="tùy chọn"),
        ),
    )

    def paint_container(d):
        nodes = {
            "Browser": (50, 360, 250, 530),
            "Nginx": (315, 360, 515, 530),
            "Frontend": (590, 210, 860, 380),
            "Backend": (590, 520, 860, 700),
            "SQL": (1010, 170, 1270, 330),
            "Redis": (1010, 380, 1270, 540),
            "Storage": (1010, 590, 1270, 750),
            "Metrics": (1340, 380, 1550, 540),
        }
        labels = {
            "Browser": "Trình duyệt",
            "Nginx": "Nginx\nReverse Proxy",
            "Frontend": "React 19 + Vite\nSPA",
            "Backend": "FastAPI\nModular Monolith",
            "SQL": "SQL Server 2022\nSource of Truth",
            "Redis": "Redis 7\nCache-aside",
            "Storage": "Upload Volume\nAttachments",
            "Metrics": "Prometheus\nGrafana",
        }
        colors = {
            "Browser": ("#F2F4F7", "#5B6573"),
            "Nginx": ("#F2F4F7", "#5B6573"),
            "Frontend": ("#E8F1FA", "#2E74B5"),
            "Backend": ("#EEF8F5", "#0F766E"),
            "SQL": ("#FFF7E8", "#B78318"),
            "Redis": ("#FFF1F1", "#9B1C1C"),
            "Storage": ("#F6F0FF", "#6B46A1"),
            "Metrics": ("#F2F4F7", "#5B6573"),
        }
        for key, xy in nodes.items():
            box(d, xy, labels[key], colors[key][0], colors[key][1], fnt=font(24, True))
        arrow(d, (250, 445), (315, 445))
        arrow(d, (515, 430), (590, 300))
        arrow(d, (515, 460), (590, 610))
        arrow(d, (860, 610), (1010, 250), label="ORM")
        arrow(d, (860, 610), (1010, 460), label="cache")
        arrow(d, (860, 620), (1010, 670), label="file")
        arrow(d, (1270, 460), (1340, 460), label="metrics")

    save_diagram(
        ARCH_DIAGRAM_DIR / "02_Container_Architecture.png",
        "Kiến trúc container",
        "Các container triển khai và phụ thuộc vận hành",
        paint_container,
    )

    def paint_layers(d):
        layers = [
            ("Router / API v1", "#E8F1FA", "#2E74B5"),
            ("Pydantic Schema + RBAC Boundary", "#EAF6FF", "#3B82F6"),
            ("Service - Business Workflow", "#EEF8F5", "#0F766E"),
            ("CRUD / Repository", "#F6F0FF", "#6B46A1"),
            ("SQLAlchemy ORM + Session", "#FFF7E8", "#B78318"),
            ("SQL Server", "#FDECEC", "#9B1C1C"),
        ]
        y = 170
        for i, (label, fill, outline) in enumerate(layers):
            box(d, (340, y, 980, y + 88), label, fill, outline, fnt=font(25, True))
            if i < len(layers) - 1:
                arrow(d, (660, y + 88), (660, y + 120))
            y += 115
        box(d, (1100, 250, 1500, 390), "CacheInvalidator + Redis", "#FFF1F1", "#9B1C1C", fnt=font(24, True))
        box(d, (1100, 500, 1500, 640), "WebSocket Manager", "#EEF8F5", "#0F766E", fnt=font(24, True))
        arrow(d, (980, 440), (1100, 320), label="invalidate")
        arrow(d, (980, 440), (1100, 570), label="publish")
        d.text((980, 720), "Domain event chỉ phát sau khi transaction commit.", fill="#" + RED, font=font(18, True))

    save_diagram(
        ARCH_DIAGRAM_DIR / "03_Backend_Layers.png",
        "Kiến trúc phân lớp Backend",
        "Router → Service → CRUD → ORM → Database",
        paint_layers,
    )

    def paint_domain(d):
        nodes = {
            "Project": (650, 180, 950, 290),
            "Epic": (250, 360, 520, 470),
            "Sprint": (665, 360, 935, 470),
            "Backlog": (1080, 360, 1390, 470),
            "Task": (650, 570, 950, 680),
            "Member": (160, 660, 500, 770),
            "Employee": (1080, 660, 1390, 770),
        }
        labels = {
            "Project": "PROJECT",
            "Epic": "TOPIC / EPIC",
            "Sprint": "SPRINT",
            "Backlog": "BACKLOG ITEM",
            "Task": "TASK",
            "Member": "PROJECT MEMBER",
            "Employee": "EMPLOYEE",
        }
        for k, xy in nodes.items():
            fill = "#E8F1FA" if k in ("Project", "Sprint", "Task") else "#F2F4F7"
            box(d, xy, labels[k], fill, "#2E74B5", fnt=font(24, True))
        arrow(d, (750, 290), (420, 360), label="1..n")
        arrow(d, (800, 290), (800, 360), label="1..n")
        arrow(d, (850, 290), (1235, 360), label="1..n")
        arrow(d, (385, 470), (690, 590), label="nhóm")
        arrow(d, (800, 470), (800, 570), label="thực thi")
        arrow(d, (1235, 470), (910, 590), label="chuyển thành")
        arrow(d, (650, 630), (500, 710), label="giao")
        arrow(d, (950, 630), (1080, 710), label="nhận")

    save_diagram(
        ARCH_DIAGRAM_DIR / "04_Domain_Model.png",
        "Mô hình miền Project - Sprint - Task",
        "Quan hệ cốt lõi giữa phạm vi dự án, backlog và người thực hiện",
        paint_domain,
    )

    def paint_deploy(d):
        box(d, (60, 360, 260, 520), "Client", "#F2F4F7", "#5B6573")
        box(d, (330, 310, 590, 570), "Nginx\n:80 / :443", "#E8F1FA", "#2E74B5")
        box(d, (690, 170, 970, 330), "Frontend\n:8080", "#E8F1FA", "#2E74B5")
        box(d, (690, 520, 970, 680), "Backend\n:8000", "#EEF8F5", "#0F766E")
        box(d, (1080, 150, 1360, 310), "SQL Server\n:1433", "#FFF7E8", "#B78318")
        box(d, (1080, 380, 1360, 540), "Redis\n:6379", "#FFF1F1", "#9B1C1C")
        box(d, (1080, 610, 1360, 770), "Persistent\nVolumes", "#F6F0FF", "#6B46A1")
        box(d, (1390, 150, 1550, 310), "sqlserver-\ninit", "#F2F4F7", "#5B6573", fnt=font(20, True))
        arrow(d, (260, 440), (330, 440))
        arrow(d, (590, 400), (690, 250))
        arrow(d, (590, 480), (690, 600))
        arrow(d, (970, 580), (1080, 230))
        arrow(d, (970, 600), (1080, 460))
        arrow(d, (970, 620), (1080, 690))
        arrow(d, (1390, 230), (1360, 230))

    save_diagram(
        ARCH_DIAGRAM_DIR / "05_Deployment.png",
        "Sơ đồ triển khai",
        "Topology Docker Compose cho môi trường production",
        paint_deploy,
    )

    def paint_function(d):
        box(d, (575, 160, 1025, 280), "TASKSYNCENTERPRISE", "#16324F", "#16324F", fnt=font(31, True), text_color="#FFFFFF")
        modules = [
            ("Xác thực & RBAC", (70, 410, 340, 530)),
            ("Tổ chức & Nhân sự", (365, 410, 635, 530)),
            ("Dự án & Backlog", (660, 410, 930, 530)),
            ("Task & Cộng tác", (955, 410, 1225, 530)),
            ("Nghỉ phép & OT", (1250, 410, 1530, 530)),
        ]
        for label, xy in modules:
            box(d, xy, label, "#E8F1FA", "#2E74B5", fnt=font(21, True))
            arrow(d, (800, 280), ((xy[0] + xy[2]) / 2, xy[1]))
        subs = [
            "Đăng nhập\nPhiên đăng nhập",
            "Phòng ban\nTeam / Employee",
            "Epic\nSprint / Backlog",
            "Kanban\nComment / File",
            "Đăng ký\nDuyệt / Chính sách",
        ]
        for (_, xy), label in zip(modules, subs):
            cx = (xy[0] + xy[2]) / 2
            box(d, (xy[0], 650, xy[2], 790), label, "#F8FAFC", "#9AA7B4", fnt=font(20))
            arrow(d, (cx, 530), (cx, 650))

    save_diagram(
        AGILE_DIAGRAM_DIR / "01_Phan_ra_chuc_nang.png",
        "Sơ đồ phân rã chức năng",
        "Phân nhóm phạm vi sản phẩm phục vụ Product Backlog",
        paint_function,
    )

    def paint_usecase(d):
        actor(d, 110, 360, "Admin")
        actor(d, 110, 650, "Manager")
        actor(d, 1490, 500, "Employee")
        d.rounded_rectangle((300, 160, 1300, 820), radius=28, outline="#" + BLUE, width=4)
        cases = [
            ((380, 220, 720, 330), "Quản lý tài khoản & RBAC"),
            ((850, 220, 1190, 330), "Quản lý tổ chức & nhân sự"),
            ((380, 420, 720, 530), "Lập kế hoạch Project / Sprint"),
            ((850, 420, 1190, 530), "Quản lý Task / Kanban"),
            ((380, 620, 720, 730), "Duyệt nghỉ phép / OT"),
            ((850, 620, 1190, 730), "Theo dõi thông báo / báo cáo"),
        ]
        for xy, label in cases:
            ellipse(d, xy, label)
        for y in (275, 675):
            arrow(d, (155, 360 if y == 275 else 650), (380, y))
        arrow(d, (155, 650), (380, 475))
        arrow(d, (1445, 500), (1190, 475))
        arrow(d, (1445, 500), (1190, 675))
        arrow(d, (155, 360), (850, 275))

    save_diagram(
        AGILE_DIAGRAM_DIR / "02_Use_Case_Tong_Quat.png",
        "Use Case tổng quát",
        "Ba nhóm tác nhân chính và các năng lực nghiệp vụ",
        paint_usecase,
    )

    def paint_task_activity(d):
        steps = [
            ("Bắt đầu", "#16324F", "#16324F", "#FFFFFF"),
            ("Tạo Task / Backlog Item", "#E8F1FA", "#2E74B5", "#111827"),
            ("Kiểm tra quyền & dữ liệu", "#FFF7E8", "#B78318", "#111827"),
            ("Gán người phụ trách", "#E8F1FA", "#2E74B5", "#111827"),
            ("In Progress", "#EEF8F5", "#0F766E", "#111827"),
            ("Review Acceptance Criteria", "#FFF7E8", "#B78318", "#111827"),
            ("Done + Audit Event", "#E7F6ED", "#1F6B45", "#111827"),
        ]
        x = 50
        for i, (label, fill, outline, tc) in enumerate(steps):
            box(d, (x, 330, x + 180, 520), label, fill, outline, fnt=font(19, True), text_color=tc)
            if i < len(steps) - 1:
                arrow(d, (x + 180, 425), (x + 215, 425))
            x += 215
        d.text((1030, 610), "Không đạt", fill="#" + RED, font=font(22, True))
        d.line((1215, 520, 1215, 670, 1000, 670, 1000, 520), fill="#" + RED, width=4)
        arrow(d, (1000, 670), (1000, 525), color="#" + RED)

    save_diagram(
        AGILE_DIAGRAM_DIR / "03_Activity_Task_Lifecycle.png",
        "Activity Diagram - Vòng đời Task",
        "Từ backlog đến Done, có kiểm tra quyền, review và audit",
        paint_task_activity,
    )

    def paint_sprint_activity(d):
        steps = [
            ("Product Backlog", "#E8F1FA", "#2E74B5"),
            ("Sprint Planning", "#FFF7E8", "#B78318"),
            ("Sprint Backlog", "#E8F1FA", "#2E74B5"),
            ("Daily Scrum + Development", "#EEF8F5", "#0F766E"),
            ("Sprint Review", "#F6F0FF", "#6B46A1"),
            ("Retrospective", "#FFF7E8", "#B78318"),
            ("Increment", "#E7F6ED", "#1F6B45"),
        ]
        coords = [(80, 280), (350, 280), (620, 280), (890, 280), (1160, 280), (1160, 590), (760, 590)]
        for (label, fill, outline), (x, y) in zip(steps, coords):
            box(d, (x, y, x + 220, y + 130), label, fill, outline, fnt=font(20, True))
        for a, b in zip(coords[:4], coords[1:5]):
            arrow(d, (a[0] + 220, a[1] + 65), (b[0], b[1] + 65))
        arrow(d, (1270, 410), (1270, 590))
        arrow(d, (1160, 655), (980, 655))
        arrow(d, (760, 655), (300, 655), label="cập nhật backlog")
        arrow(d, (300, 655), (190, 410))

    save_diagram(
        AGILE_DIAGRAM_DIR / "04_Activity_Sprint.png",
        "Activity Diagram - Chu trình Sprint",
        "Các sự kiện Scrum và vòng phản hồi sang Product Backlog",
        paint_sprint_activity,
    )

    def paint_scrum(d):
        box(d, (80, 330, 330, 500), "Product Owner\nProduct Backlog", "#E8F1FA", "#2E74B5", fnt=font(22, True))
        box(d, (420, 330, 670, 500), "Sprint Planning", "#FFF7E8", "#B78318", fnt=font(22, True))
        box(d, (760, 330, 1010, 500), "Sprint Backlog", "#E8F1FA", "#2E74B5", fnt=font(22, True))
        box(d, (1100, 220, 1440, 390), "Development + Daily Scrum", "#EEF8F5", "#0F766E", fnt=font(22, True))
        box(d, (1100, 560, 1440, 730), "Increment + Review", "#E7F6ED", "#1F6B45", fnt=font(22, True))
        box(d, (680, 620, 960, 770), "Retrospective", "#F6F0FF", "#6B46A1", fnt=font(22, True))
        arrow(d, (330, 415), (420, 415))
        arrow(d, (670, 415), (760, 415))
        arrow(d, (1010, 415), (1100, 305))
        arrow(d, (1270, 390), (1270, 560))
        arrow(d, (1100, 645), (960, 695))
        arrow(d, (680, 695), (205, 500), label="adapt")

    save_diagram(
        AGILE_DIAGRAM_DIR / "05_Scrum_Workflow.png",
        "Scrum Workflow",
        "Luồng kế hoạch - thực thi - kiểm tra - thích ứng",
        paint_scrum,
    )

    mmd = {
        ARCH_DIAGRAM_DIR / "01_System_Context.mmd": """flowchart LR
Admin --> Web[React Web Application]
Manager --> Web
Employee --> Web
Web --> Platform[TaskSyncEnterprise REST API + WebSocket]
Platform -. notification .-> Provider[Email / Push Provider]""",
        ARCH_DIAGRAM_DIR / "02_Container_Architecture.mmd": """flowchart TB
Browser --> Nginx
Nginx --> Frontend[React + Vite SPA]
Nginx --> Backend[FastAPI Modular Monolith]
Backend --> SQL[SQL Server - Source of Truth]
Backend --> Redis[Redis - Cache Aside]
Backend --> Storage[Upload Volume]
Backend --> Metrics[Prometheus / Grafana]""",
        ARCH_DIAGRAM_DIR / "03_Backend_Layers.mmd": """flowchart TB
Request --> Router[Router / API v1]
Router --> Schema[Pydantic + RBAC]
Schema --> Service[Business Service]
Service --> CRUD
CRUD --> ORM[SQLAlchemy ORM]
ORM --> DB[SQL Server]
Service --> Cache[CacheInvalidator + Redis]
Service --> WS[WebSocket Manager]""",
        ARCH_DIAGRAM_DIR / "04_Domain_Model.mmd": """erDiagram
PROJECT ||--o{ TOPIC : defines
PROJECT ||--o{ SPRINT : plans
PROJECT ||--o{ BACKLOG_ITEM : owns
PROJECT ||--o{ TASK : delivers
TOPIC ||--o{ BACKLOG_ITEM : groups
SPRINT ||--o{ TASK : executes
EMPLOYEE ||--o{ TASK_ASSIGNMENT : receives
TASK ||--o{ TASK_ASSIGNMENT : assigns""",
        ARCH_DIAGRAM_DIR / "05_Deployment.mmd": """flowchart LR
Client --> Nginx
Nginx --> Frontend
Nginx --> Backend
Backend --> SQLServer
Backend --> Redis
Backend --> Volumes
SQLInit[sqlserver-init] --> SQLServer""",
        AGILE_DIAGRAM_DIR / "01_Phan_ra_chuc_nang.mmd": """flowchart TB
TSE[TaskSyncEnterprise] --> Auth[Xác thực & RBAC]
TSE --> Org[Tổ chức & Nhân sự]
TSE --> Project[Dự án & Backlog]
TSE --> Task[Task & Cộng tác]
TSE --> Leave[Nghỉ phép & OT]""",
        AGILE_DIAGRAM_DIR / "02_Use_Case_Tong_Quat.mmd": """flowchart LR
Admin --> RBAC[Quản lý tài khoản & RBAC]
Admin --> Org[Quản lý tổ chức & nhân sự]
Manager --> Plan[Lập kế hoạch Project / Sprint]
Manager --> Kanban[Quản lý Task / Kanban]
Manager --> Approve[Duyệt nghỉ phép / OT]
Employee --> Kanban
Employee --> Notify[Theo dõi thông báo / báo cáo]""",
        AGILE_DIAGRAM_DIR / "03_Activity_Task_Lifecycle.mmd": """flowchart LR
Start --> Create[Tạo Task]
Create --> Validate[Kiểm tra quyền & dữ liệu]
Validate --> Assign[Gán người phụ trách]
Assign --> Doing[In Progress]
Doing --> Review[Review Acceptance Criteria]
Review -->|Đạt| Done[Done + Audit Event]
Review -->|Không đạt| Doing""",
        AGILE_DIAGRAM_DIR / "04_Activity_Sprint.mmd": """flowchart LR
PB[Product Backlog] --> Planning[Sprint Planning]
Planning --> SB[Sprint Backlog]
SB --> Daily[Daily Scrum + Development]
Daily --> Review[Sprint Review]
Review --> Retro[Retrospective]
Retro --> Increment
Retro --> PB""",
        AGILE_DIAGRAM_DIR / "05_Scrum_Workflow.mmd": """flowchart LR
PO[Product Owner / Product Backlog] --> Planning
Planning --> SB[Sprint Backlog]
SB --> Dev[Development + Daily Scrum]
Dev --> Increment
Increment --> Review
Review --> Retro
Retro --> PO""",
    }
    for path, content in mmd.items():
        path.write_text(content + "\n", encoding="utf-8")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    rfonts = run._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.text = running_title
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.runs[0], size=9, color=GRAY, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("TaskSyncEnterprise | Trang ")
    set_run_font(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)
    return doc


def add_cover(doc, course, title, subtitle, date_text="TP. Hồ Chí Minh, tháng 7 năm 2026"):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("TRƯỜNG ĐẠI HỌC NGUYỄN TẤT THÀNH")
    set_run_font(r, size=13, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VIỆN ĐÀO TẠO QUỐC TẾ NIIE")
    set_run_font(r, size=12, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(30)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(course.upper())
    set_run_font(r, size=12, color=GOLD, bold=True)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size=26, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, size=14, color=DARK_BLUE, italic=True)
    p.paragraph_format.space_after = Pt(36)

    rows = [
        ("Giảng viên hướng dẫn", "ThS. Trần Thanh Nhã"),
        ("Đề tài", "Nền tảng quản lý công việc và nhân sự doanh nghiệp TaskSyncEnterprise"),
        ("Nhóm thực hiện", "Scrum Team TaskSyncEnterprise - 5 thành viên"),
        ("Phiên bản tài liệu", "1.0 - Chuẩn hóa 2 Sprint"),
    ]
    add_table(doc, ["Thông tin", "Nội dung"], rows, [2700, 6660], font_size=10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run(date_text)
    set_run_font(r, size=11, color=GRAY, italic=True)
    doc.add_page_break()


def add_body(doc, text, bold_lead=None, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def add_note(doc, label, text, color=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, font_size=9, header_fill=LIGHT):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run("" if value is None else str(value))
            set_run_font(r, size=font_size, color=BLACK)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_diagram(doc, image_path, caption, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    p.paragraph_format.space_after = Pt(4)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(0)
    c.paragraph_format.space_after = Pt(10)
    r = c.add_run(caption)
    set_run_font(r, size=9, color=GRAY, italic=True)


def add_sources(doc, sources):
    doc.add_heading("Tài liệu nguồn và đối chiếu", level=1)
    for src in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(src)


def save_doc(doc, path):
    doc.core_properties.title = path.stem
    doc.core_properties.subject = "TaskSyncEnterprise report package"
    doc.core_properties.author = "Scrum Team TaskSyncEnterprise"
    doc.core_properties.keywords = "TaskSyncEnterprise, Agile, Software Architecture, Scrum"
    doc.save(path)


def build_architecture_report():
    doc = style_document(Document(), "BÁO CÁO KIẾN TRÚC PHẦN MỀM | TASKSYNCENTERPRISE")
    add_cover(
        doc,
        "Môn Kiến trúc phần mềm",
        "BÁO CÁO KIẾN TRÚC PHẦN MỀM",
        "Modular Monolith phân lớp, realtime và triển khai container",
    )
    doc.add_heading("Tóm tắt điều hành", level=1)
    add_body(
        doc,
        "TaskSyncEnterprise là nền tảng web quản lý công việc, dự án và nhân sự doanh nghiệp. "
        "Giải pháp hiện tại lựa chọn Modular Monolith để giữ chi phí triển khai thấp, nhưng tổ chức mã nguồn theo ranh giới miền rõ ràng để có thể tách dịch vụ khi quy mô tăng."
    )
    add_note(
        doc,
        "Kết luận kiến trúc",
        "SQL Server là nguồn dữ liệu chuẩn; Redis chỉ là cache-aside. RBAC được thực thi tại Backend. "
        "Domain event chỉ phát sau khi transaction commit thành công.",
    )
    doc.add_heading("1. Bối cảnh và phạm vi", level=1)
    add_body(
        doc,
        "Hệ thống phục vụ ba nhóm người dùng chính: Administrator, Manager/Team Leader và Employee. "
        "Các miền nghiệp vụ gồm Organization, Project, Epic/Topic, Product Backlog, Sprint, Task, Collaboration, File, Vacation, Overtime, Notification và Audit."
    )
    add_bullets(
        doc,
        [
            "Frontend: React 19, Vite, React Router, TanStack Query và lớp dịch vụ Axios.",
            "Backend: Python/FastAPI, Pydantic, SQLAlchemy và Alembic.",
            "Dữ liệu: Microsoft SQL Server 2022; Redis 7 hỗ trợ cache và trạng thái vận hành.",
            "Realtime: WebSocket cập nhật các thay đổi nghiệp vụ mà không cần tải lại trang.",
            "Triển khai: Nginx, Docker Compose, persistent volume và overlay monitoring tùy chọn.",
        ],
    )
    add_diagram(doc, ARCH_DIAGRAM_DIR / "01_System_Context.png", "Hình 1. Sơ đồ ngữ cảnh hệ thống TaskSyncEnterprise")

    doc.add_heading("2. Các động lực kiến trúc", level=1)
    add_table(
        doc,
        ["Động lực", "Yêu cầu kiến trúc", "Cách đáp ứng"],
        [
            ("Dễ bảo trì", "Ranh giới module rõ; phụ thuộc một chiều", "Router → Service → CRUD → ORM"),
            ("Bảo mật", "Xác thực, RBAC và audit ở server", "JWT/Session, policy backend, audit log"),
            ("Realtime", "Thay đổi hiển thị trên nhiều client", "Domain event + WebSocket + query invalidation"),
            ("Khả dụng", "Cache không là điểm lỗi duy nhất", "Cache-aside; SQL Server giữ dữ liệu chuẩn"),
            ("Triển khai", "Đơn giản cho nhóm 5 người", "Modular Monolith trong Docker Compose"),
            ("Mở rộng", "Có thể tách tác vụ nặng", "Ranh giới Notification, Reporting và AI"),
        ],
        [1800, 3300, 4260],
        font_size=9,
    )

    doc.add_heading("3. Phong cách kiến trúc", level=1)
    doc.add_heading("3.1. Modular Monolith", level=2)
    add_body(
        doc,
        "Kiến trúc được đóng gói như một ứng dụng Backend duy nhất nhưng chia module theo miền. "
        "Lựa chọn này phù hợp với quy mô nhóm, giảm độ phức tạp quan sát và đồng bộ dữ liệu so với microservices, đồng thời tránh gom toàn bộ logic vào một khối không cấu trúc."
    )
    doc.add_heading("3.2. Client - Server và kiến trúc phân lớp", level=2)
    add_body(
        doc,
        "React SPA giao tiếp REST API và WebSocket. Backend phân lớp để ràng buộc validation, authorization, business workflow và persistence. "
        "Frontend không phải là lớp bảo mật; việc ẩn nút chỉ cải thiện trải nghiệm người dùng."
    )
    add_diagram(doc, ARCH_DIAGRAM_DIR / "02_Container_Architecture.png", "Hình 2. Kiến trúc container")
    add_diagram(doc, ARCH_DIAGRAM_DIR / "03_Backend_Layers.png", "Hình 3. Kiến trúc phân lớp Backend")

    doc.add_heading("4. Mô hình miền và dữ liệu", level=1)
    add_body(
        doc,
        "Project là phạm vi liên kết Epic, Backlog Item, Sprint và Task. Backlog Item biểu diễn User Story có giá trị nghiệp vụ, ưu tiên và Story Point; khi được lựa chọn vào Sprint, item được phân rã thành Task thực thi."
    )
    add_diagram(doc, ARCH_DIAGRAM_DIR / "04_Domain_Model.png", "Hình 4. Mô hình miền Project - Sprint - Task")
    add_table(
        doc,
        ["Invariant", "Ý nghĩa"],
        [
            ("Đồng nhất Project", "Task, Sprint, Backlog và Epic phải thuộc cùng một Project."),
            ("Quyền truy cập", "Người được giao Task phải là thành viên có quyền truy cập Project."),
            ("Sprint đóng", "Không thêm công việc vào Sprint đã hoàn tất nếu chưa reopen."),
            ("Capacity", "Tổng Story Point cam kết không vượt capacity đã thống nhất."),
            ("Audit", "Soft delete và trạng thái nghiệp vụ phải giữ được dấu vết thay đổi."),
        ],
        [2300, 7060],
        font_size=9.5,
    )

    doc.add_heading("5. Luồng ghi dữ liệu và realtime", level=1)
    add_bullets(
        doc,
        [
            "Client gửi POST/PATCH/DELETE tới API.",
            "Backend kiểm tra schema, xác thực và quyền trước khi chạy workflow.",
            "Service gọi CRUD/ORM trong transaction; SQL Server commit dữ liệu.",
            "Sau commit, hệ thống phát domain event như task.changed hoặc vacation.changed.",
            "Client khác nhận sự kiện, invalidate query key và refetch dữ liệu mới.",
        ],
        numbered=True,
    )
    add_note(
        doc,
        "Nguyên tắc nhất quán",
        "Không phát event trước commit. Nếu Redis ngừng hoạt động, nghiệp vụ chính vẫn đọc/ghi từ SQL Server; cache phải được xem là tối ưu hóa, không phải nguồn dữ liệu.",
    )

    doc.add_heading("6. Bảo mật và kiểm soát truy cập", level=1)
    add_table(
        doc,
        ["Lớp kiểm soát", "Biện pháp"],
        [
            ("Transport", "HTTPS tại Nginx; WebSocket qua kết nối được xác thực."),
            ("Identity", "Login, access token/session, refresh và blacklist khi cần."),
            ("Authorization", "RBAC tại Router/Service; kiểm tra phạm vi Project/Department."),
            ("Input", "Pydantic validation, kiểm tra file, giới hạn kích thước và loại dữ liệu."),
            ("Persistence", "Transaction, unique/FK constraints, migration Alembic."),
            ("Audit", "Structured logging và audit trail cho thay đổi nhạy cảm."),
        ],
        [2300, 7060],
        font_size=9.5,
    )

    doc.add_heading("7. Triển khai và vận hành", level=1)
    add_diagram(doc, ARCH_DIAGRAM_DIR / "05_Deployment.png", "Hình 5. Sơ đồ triển khai Docker Compose")
    add_body(
        doc,
        "Ở production, Nginx là điểm vào duy nhất. Frontend và Backend chạy trong container riêng; SQL Server và Redis có volume. "
        "Container sqlserver-init bảo đảm database tồn tại, trong khi Alembic chịu trách nhiệm cập nhật schema."
    )
    add_note(
        doc,
        "Trạng thái kiểm tra ngày 29/07/2026",
        "331 kiểm thử Backend và 22 kiểm thử Frontend được tài liệu audit ghi nhận đạt. Tuy nhiên Redis local đang DOWN, readiness trả 503, Docker Desktop engine chưa chạy và còn dependency advisory. "
        "Vì vậy tài liệu này không kết luận production-ready trước khi hoàn thành các điều kiện khắc phục.",
        color="FFF7E8",
    )

    doc.add_heading("8. Rủi ro, trade-off và hướng mở rộng", level=1)
    add_table(
        doc,
        ["Rủi ro / trade-off", "Xử lý đề xuất"],
        [
            ("Monolith tăng coupling khi module lớn", "Giữ interface nội bộ; cấm truy cập chéo model tùy tiện."),
            ("Mất event giữa commit và publish", "Áp dụng transactional outbox khi cần bảo đảm delivery."),
            ("Upload volume khó scale nhiều replica", "Chuyển sang object storage khi triển khai phân tán."),
            ("Báo cáo/AI làm chậm API", "Tách worker hoặc service bất đồng bộ."),
            ("Residual dependency risk", "Theo dõi advisory và thay dependency khi có bản vá khả dụng."),
        ],
        [3100, 6260],
        font_size=9.2,
    )
    doc.add_heading("9. Kết luận", level=1)
    add_body(
        doc,
        "Kiến trúc Modular Monolith đáp ứng tốt phạm vi đồ án và năng lực nhóm 5 người. Giá trị quan trọng nhất là ranh giới module, kiểm soát quyền tại Backend, SQL Server làm nguồn dữ liệu chuẩn và luồng realtime sau commit. "
        "Các hướng tách service chỉ nên thực hiện khi có số liệu tải và yêu cầu vận hành đủ rõ."
    )
    add_sources(
        doc,
        [
            r"docs/architecture/SYSTEM_ARCHITECTURE.md",
            r"docs/architecture/DOMAIN_MODEL.md",
            r"docs/reports/CONSOLIDATED_AUDIT.md",
            r"Report/Tham khảo/Cac_bang_table.docx",
            r"Report/Tham khảo/bao-cao-do-an-tasksyncenterprise (1).docx",
        ],
    )
    save_doc(doc, ARCH_DIR / "Bao_cao_Kien_truc_phan_mem_TaskSyncEnterprise.docx")


def build_agile_report():
    doc = style_document(Document(), "BÁO CÁO AGILE / SCRUM | TASKSYNCENTERPRISE")
    add_cover(
        doc,
        "Mô hình phát triển phần mềm Agile",
        "BÁO CÁO AGILE / SCRUM",
        "Product Backlog 40 User Stories và kế hoạch 2 Sprint đã gộp",
    )
    doc.add_heading("Tóm tắt", level=1)
    add_body(
        doc,
        "Báo cáo trình bày cách nhóm áp dụng Scrum cho TaskSyncEnterprise, chuẩn hóa Product Backlog 40 User Stories thuộc 12 Epics và gộp bốn Sprint ban đầu thành hai Sprint theo yêu cầu giảng viên."
    )
    add_note(
        doc,
        "Quy tắc bắt buộc",
        "Sprint cũ 1 + 2 được gộp thành Sprint 1 mới; Sprint cũ 3 + 4 được gộp thành Sprint 2 mới. "
        "Mọi bảng, Daily Scrum và Release Plan trong bộ hồ sơ sử dụng quy ước này.",
    )

    doc.add_heading("1. Project Charter", level=1)
    add_table(
        doc,
        ["Thuộc tính", "Nội dung"],
        [
            ("Tên dự án", "TaskSyncEnterprise"),
            ("Mục tiêu", "Quản lý công việc, dự án, tổ chức, nhân sự, nghỉ phép và thông báo realtime."),
            ("Khung làm việc", "Scrum với Product Backlog, Sprint Backlog, Daily Scrum, Review và Retrospective."),
            ("Quy mô", "5 thành viên; 40 User Stories; 12 Epics; 242 Story Points theo tổng từng dòng backlog."),
            ("Công nghệ", "FastAPI, React/Vite, SQL Server, Redis, WebSocket và Docker Compose."),
        ],
        [2500, 6860],
        font_size=9.5,
    )

    doc.add_heading("2. Scrum Team và trách nhiệm", level=1)
    add_table(
        doc,
        ["Mã", "Thành viên", "Vai trò", "Trách nhiệm chính"],
        [
            (code, name, role, {
                "A": "Tầm nhìn sản phẩm, kiến trúc, ưu tiên backlog, code review.",
                "B": "Facilitate Scrum, Backend API, database và workflow nghiệp vụ.",
                "C": "React UI, Kanban, form và trải nghiệm người dùng.",
                "D": "Test plan, API/E2E, regression và quality gate.",
                "E": "Nghiên cứu, biên bản Scrum, báo cáo và chuẩn hóa tài liệu.",
            }[code])
            for code, name, role in TEAM
        ],
        [600, 2100, 2500, 4160],
        font_size=8.6,
    )
    add_note(
        doc,
        "Điểm cần xác nhận",
        "Một số tài liệu nguồn ghi khác nhau về họ tên người phụ trách tài liệu và họ của Product Owner. "
        "Bộ hồ sơ này dùng danh sách trong workbook/báo cáo tổng hợp mới nhất; nhóm cần đối chiếu MSSV trước khi nộp.",
        color="FFF7E8",
    )

    doc.add_heading("3. Cơ sở Scrum áp dụng", level=1)
    add_bullets(
        doc,
        [
            "Ba trụ cột: Minh bạch, Kiểm tra và Thích ứng.",
            "Năm giá trị: Cam kết, Tập trung, Cởi mở, Tôn trọng và Dũng cảm.",
            "Artifacts: Product Backlog, Sprint Backlog và Increment.",
            "Events: Sprint Planning, Daily Scrum, Sprint Review và Sprint Retrospective.",
        ],
    )
    add_diagram(doc, AGILE_DIAGRAM_DIR / "05_Scrum_Workflow.png", "Hình 1. Scrum Workflow áp dụng cho TaskSyncEnterprise")

    doc.add_heading("4. Cấu trúc Product Backlog", level=1)
    add_body(
        doc,
        "Product Backlog được nhóm theo 12 Epics: Access & Security, Employee, Department, Project, Task/Kanban, Collaboration, File, Vacation, Overtime, Notification, Analytics và Audit. "
        "Mỗi item có User Story, Priority, Story Point, Business Value, Assignee, Target Sprint và Acceptance Criteria."
    )
    add_diagram(doc, AGILE_DIAGRAM_DIR / "01_Phan_ra_chuc_nang.png", "Hình 2. Phân rã chức năng làm cơ sở tổ chức backlog")
    add_table(
        doc,
        ["Nhóm chức năng", "User Stories tiêu biểu", "Giá trị"],
        [
            ("Xác thực & RBAC", "US01-US04", "Truy cập an toàn và phân quyền đúng."),
            ("Project / Task", "US05-US07, US15-US20, US30", "Lập kế hoạch và thực thi công việc."),
            ("Tổ chức / Nhân sự", "US09-US14", "Chuẩn hóa cơ cấu và hồ sơ nhân viên."),
            ("File / Notification", "US08, US21, US29, US31-US34", "Cộng tác và cập nhật kịp thời."),
            ("Vacation / Overtime", "US22-US28", "Số hóa quy trình hành chính."),
            ("Analytics / Audit", "US35-US40", "Đo lường và truy vết hoạt động."),
        ],
        [2300, 3100, 3960],
        font_size=9,
    )

    doc.add_heading("5. Chuẩn hóa Sprint theo yêu cầu giảng viên", level=1)
    add_table(
        doc,
        ["Sprint cũ", "Thời gian cũ", "Sprint mới", "Phạm vi", "Story Points"],
        [
            ("Sprint 1", "02/07-16/07", "Sprint 1", "US01-US08", "39"),
            ("Sprint 2", "16/07-30/07", "Sprint 1", "US09-US18", "68"),
            ("Sprint 3", "30/07-13/08", "Sprint 2", "US19-US29", "64"),
            ("Sprint 4", "13/08-27/08", "Sprint 2", "US30-US40", "71"),
        ],
        [1300, 1800, 1300, 2800, 2160],
        font_size=8.8,
    )
    add_note(
        doc,
        "Baseline sau khi gộp",
        "Sprint 1: 02/07-30/07/2026, 18 User Stories, 107 SP. "
        "Sprint 2: 30/07-27/08/2026, 22 User Stories, 135 SP. Tổng: 40 User Stories, 242 SP. "
        "Trang Summary của file nguồn ghi 228 SP nhưng không khớp phép cộng từng dòng; bộ hồ sơ dùng tổng dòng để có thể kiểm tra.",
    )

    doc.add_heading("6. Kế hoạch Sprint 1", level=1)
    add_table(
        doc,
        ["Thuộc tính", "Nội dung"],
        [
            ("Sprint Goal", "Thiết lập nền tảng an toàn, cơ cấu tổ chức, hồ sơ nhân sự và năng lực Project/Task cơ bản."),
            ("Phạm vi", "US01-US18; Auth/RBAC, Project Wizard, Kanban, Employee, Department, Backlog và Collaboration."),
            ("Capacity", "107 Story Points sau khi gộp Sprint cũ 1 và 2."),
            ("Definition of Done", "Code review; test đạt; không lỗi nghiêm trọng; tài liệu/API cập nhật; demo được."),
            ("Các vấn đề ghi nhận", "Dead-code frontend, loading giả, HTTP 409 thiếu thông điệp và import Excel."),
        ],
        [2500, 6860],
        font_size=9.2,
    )
    doc.add_heading("6.1. Retrospective Sprint 1", level=2)
    add_table(
        doc,
        ["Nhóm", "Nội dung"],
        [
            ("Went well", "Kiến trúc, Auth/RBAC, schema tổ chức và Kanban được thống nhất."),
            ("Needs improvement", "Dọn code cũ chậm; validation và thông điệp lỗi chưa đồng đều."),
            ("Action", "Static scan trước release; kiểm tra state machine ở Backend; validate batch import trước commit."),
        ],
        [2200, 7160],
        font_size=9.2,
    )

    doc.add_page_break()
    doc.add_heading("7. Kế hoạch Sprint 2", level=1)
    add_table(
        doc,
        ["Thuộc tính", "Nội dung"],
        [
            ("Sprint Goal", "Hoàn thiện Task, Vacation/OT, Notification realtime, Analytics và Audit."),
            ("Phạm vi", "US19-US40; Task editor, mention, cleanup, leave/OT, WebSocket, báo cáo và audit."),
            ("Capacity", "135 Story Points sau khi gộp Sprint cũ 3 và 4."),
            ("Rủi ro", "CI timeout, Redis availability, orphaned files, transaction quỹ phép và dependency advisory."),
            ("Kết quả kỳ vọng", "Increment demo được; test regression đạt; tài liệu kiến trúc và Agile đồng bộ."),
        ],
        [2500, 6860],
        font_size=9.2,
    )
    add_note(
        doc,
        "Mốc tương lai",
        "Theo ngày lập hồ sơ 29/07/2026, các mốc từ 30/07 đến 27/08 là kế hoạch/baseline. "
        "Nhóm cần cập nhật trạng thái thực tế và bằng chứng họp trước khi dùng làm biên bản đã hoàn thành.",
        color="FFF7E8",
    )

    doc.add_heading("8. Các Scrum Events", level=1)
    add_table(
        doc,
        ["Sự kiện", "Timebox", "Đầu vào", "Đầu ra"],
        [
            ("Sprint Planning", "Tối đa 4 giờ/Sprint", "Product Backlog, capacity", "Sprint Goal và Sprint Backlog"),
            ("Daily Scrum", "15 phút/ngày", "Tiến độ 24 giờ, blocker", "Kế hoạch ngày và action xử lý"),
            ("Sprint Review", "Tối đa 2 giờ", "Increment và Acceptance Criteria", "Feedback và backlog cập nhật"),
            ("Retrospective", "Tối đa 1.5 giờ", "Dữ liệu Sprint", "1-3 hành động cải tiến"),
        ],
        [1700, 1900, 2800, 2960],
        font_size=8.8,
    )
    add_diagram(doc, AGILE_DIAGRAM_DIR / "04_Activity_Sprint.png", "Hình 3. Activity Diagram chu trình Sprint")

    doc.add_heading("9. Definition of Ready và Definition of Done", level=1)
    add_table(
        doc,
        ["Definition of Ready", "Definition of Done"],
        [
            ("User Story đúng mẫu As a / I want / So that.", "Acceptance Criteria được kiểm tra và Product Owner chấp nhận."),
            ("Priority, Business Value và Story Point đã thống nhất.", "Code review đạt; test liên quan chạy xanh."),
            ("Phụ thuộc, dữ liệu và thiết kế đủ rõ.", "Không có lỗi Blocker/Critical còn mở."),
            ("Assignee và Sprint được xác định.", "Tài liệu, migration và cấu hình được cập nhật."),
        ],
        [4680, 4680],
        font_size=9.2,
    )

    doc.add_heading("10. Quản trị chất lượng và đo lường", level=1)
    add_bullets(
        doc,
        [
            "Velocity = tổng Story Point hoàn thành trong Sprint.",
            "Completion Rate = số User Stories Done / tổng User Stories cam kết.",
            "Defect trend theo Sprint và severity; mọi bug có Root Cause, Fix và Lesson Learned.",
            "Không kết luận production-ready chỉ dựa vào số test; phải kiểm tra dependency, infrastructure, backup/restore và health.",
        ],
    )
    add_body(
        doc,
        "Workbook Product_Backlog_TaskSyncEnterprise.xlsx đi kèm là nguồn dữ liệu cập nhật cho 40 User Stories, hai Sprint Backlog, Dashboard và Release Plan. "
        "Nhật ký Daily Scrum nằm trong file Daily_Scrum_2_Sprint_TaskSyncEnterprise.docx."
    )

    doc.add_heading("11. Kết luận", level=1)
    add_body(
        doc,
        "Việc gộp Sprint giúp báo cáo phù hợp yêu cầu học phần mà không làm mất truy vết: mỗi User Story vẫn giữ Feature, Epic, Assignee và Story Point. "
        "Nhóm cần duy trì tính minh bạch bằng cách tách rõ dữ liệu đã xác nhận với kế hoạch tương lai, cập nhật Daily Scrum và trạng thái backlog từ bằng chứng thực tế."
    )
    add_sources(
        doc,
        [
            r"Report/Tham khảo/Co_su_Ly_Thuyet.docx",
            r"Report/Tham khảo/LTHN_Agile_Final.pdf (chỉ tham khảo bố cục, không dùng nghiệp vụ Hilo Cinema)",
            r"Report/Tham khảo/Agile_Project_Management.xlsx",
            r"Report/Tham khảo/Agile_Tracking.xlsx",
            r"Report/Tham khảo/bao-cao-thong-ke-bug-timeline (1).xlsx",
            "Daily Sprint 1: https://docs.google.com/document/d/1fa-B87BGnb2J-axm0YS6oYN5BeOzqFjK_acfUXfYw_E/edit",
        ],
    )
    save_doc(doc, AGILE_DIR / "Bao_cao_Agile_Scrum_TaskSyncEnterprise.docx")


DAILY_MEETINGS = [
    ("Sprint 1", "03/07/2026", "Khởi tạo Auth và database", [
        ("A", "Chốt kiến trúc và backlog.", "Dựng JWT/RBAC middleware.", "Cần thống nhất permission matrix."),
        ("B", "Chuẩn bị schema SQL Server.", "Tạo roles, departments, teams, employees.", "Không."),
        ("C", "Dựng khung React/Vite.", "Hoàn thiện Login và Auth Provider.", "Chờ API login."),
        ("D", "Soạn test plan.", "Viết test login hợp lệ/không hợp lệ.", "Thiếu seed account."),
        ("E", "Tổng hợp yêu cầu.", "Cập nhật Domain Model.", "Không."),
    ]),
    ("Sprint 1", "08/07/2026", "Project Wizard và phân quyền", [
        ("A", "Hoàn thành Auth cơ bản.", "Review Project Wizard và API Project.", "RBAC edge cases."),
        ("B", "CRUD organization.", "Thêm Project Member validation.", "Dữ liệu seed chưa đủ."),
        ("C", "Login và protected routes.", "Làm Project Wizard.", "Chờ schema API."),
        ("D", "Test Auth/RBAC.", "Kiểm thử quyền Admin/Manager/Employee.", "Cần permission matrix."),
        ("E", "Cập nhật biên bản.", "Viết mô tả Product Backlog.", "Không."),
    ]),
    ("Sprint 1", "14/07/2026", "Kanban và quản lý công việc", [
        ("A", "Thiết kế Task workflow.", "Hoàn thiện Kanban mutation.", "Đồng bộ status canonical."),
        ("B", "API Project/Task.", "Bổ sung transaction và conflict 409.", "Một số case thiếu message."),
        ("C", "Kanban UI.", "Kết nối drag-and-drop với API.", "Cache chưa invalidate."),
        ("D", "Regression Auth.", "Test Task status và quyền.", "Thiếu dữ liệu dependency."),
        ("E", "Mô tả use case.", "Vẽ activity Task lifecycle.", "Không."),
    ]),
    ("Sprint 1", "18/07/2026", "Employee và Department", [
        ("A", "Review Kanban.", "Xử lý query invalidation.", "Không."),
        ("B", "Schema Employee.", "API Department transfer.", "Workflow xung đột."),
        ("C", "Employee list.", "Form Department và Manager assignment.", "Chờ API transfer."),
        ("D", "Test Kanban.", "Test HTTP 409 và permission.", "Thông điệp lỗi chưa rõ."),
        ("E", "Cập nhật báo cáo.", "Chuẩn hóa bảng nhân sự.", "Không."),
    ]),
    ("Sprint 1", "24/07/2026", "Bulk import và cleanup", [
        ("A", "Audit frontend.", "Xóa dead-code và dependency thừa.", "Có file JSX legacy."),
        ("B", "Bulk import.", "Mở rộng parser date và unique email.", "Nhiều định dạng ngày."),
        ("C", "Settings UI.", "Thay loading giả bằng Provider thật.", "Không."),
        ("D", "Test import.", "Regression file Excel.", "Cần fixture dữ liệu lỗi."),
        ("E", "Ghi nhận bug.", "Cập nhật Root Cause/Fix/Lesson.", "Không."),
    ]),
    ("Sprint 1", "29/07/2026", "Tổng duyệt Sprint 1", [
        ("A", "Code review và kiến trúc.", "Chốt scope US01-US18.", "Chờ xác nhận acceptance."),
        ("B", "Sửa lỗi import/transfer.", "Rà soát migration và seed.", "Không."),
        ("C", "Hoàn thiện UI.", "Fix các lỗi responsive.", "Một số nhãn chưa thống nhất."),
        ("D", "Chạy regression.", "Lập Sprint Review checklist.", "Docker engine local chưa chạy."),
        ("E", "Gộp Sprint 1+2 cũ.", "Đồng bộ report, backlog và diagram.", "Cần xác nhận tên/MSSV."),
    ]),
    ("Sprint 2", "31/07/2026", "Kế hoạch Task nâng cao", [
        ("A", "Chuẩn bị Sprint Goal.", "Thiết kế Task editor và mention.", "Mốc cần xác nhận thực tế."),
        ("B", "Rà soát API Task.", "Tạo cleanup job.", "Mốc kế hoạch."),
        ("C", "Chuẩn bị UI.", "Task editor và comment feed.", "Mốc kế hoạch."),
        ("D", "Chuẩn bị test.", "Test file security và mention.", "Mốc kế hoạch."),
        ("E", "Cập nhật Sprint Backlog.", "Ghi biên bản.", "Mốc kế hoạch."),
    ]),
    ("Sprint 2", "05/08/2026", "Vacation và Overtime", [
        ("A", "Review workflow.", "Kiểm tra transaction và audit.", "Mốc kế hoạch."),
        ("B", "Leave balance.", "API vacation/OT.", "Policy cần PO xác nhận."),
        ("C", "Form vacation.", "Approvals panel.", "Mốc kế hoạch."),
        ("D", "Test workflow.", "Case reject/cancel/restore.", "Mốc kế hoạch."),
        ("E", "Mô tả use case.", "Cập nhật Acceptance Criteria.", "Mốc kế hoạch."),
    ]),
    ("Sprint 2", "11/08/2026", "Realtime notification", [
        ("A", "Thiết kế event.", "WebSocket publish sau commit.", "Mốc kế hoạch."),
        ("B", "Notification API.", "Retry và preference.", "Redis availability."),
        ("C", "Notification hook.", "Invalidate query cache.", "Mốc kế hoạch."),
        ("D", "Test hai browser.", "Kiểm tra không cần F5.", "Mốc kế hoạch."),
        ("E", "Cập nhật kiến trúc.", "Vẽ sequence/realtime flow.", "Mốc kế hoạch."),
    ]),
    ("Sprint 2", "17/08/2026", "Analytics và audit", [
        ("A", "Review dashboard.", "Burnup predictor.", "Mốc kế hoạch."),
        ("B", "Audit log.", "Query filter và immutable trail.", "Mốc kế hoạch."),
        ("C", "Dashboard UI.", "Resource allocation map.", "Mốc kế hoạch."),
        ("D", "Test report.", "Security/audit regression.", "Mốc kế hoạch."),
        ("E", "Soạn report.", "Tổng hợp metrics.", "Mốc kế hoạch."),
    ]),
    ("Sprint 2", "23/08/2026", "Hardening và release", [
        ("A", "Review PR.", "Sửa cache-aside và CI.", "Mốc kế hoạch."),
        ("B", "Cleanup job.", "Rà soát Docker/.env.", "Mốc kế hoạch."),
        ("C", "Realtime UI.", "Fix contract/responsive.", "Mốc kế hoạch."),
        ("D", "Regression.", "Chạy quality gate.", "Dependency advisory."),
        ("E", "Cập nhật lesson.", "Chuẩn bị Sprint Review.", "Mốc kế hoạch."),
    ]),
    ("Sprint 2", "27/08/2026", "Sprint Review và Retrospective", [
        ("A", "Tổng duyệt.", "Demo Increment và chốt backlog.", "Mốc kế hoạch."),
        ("B", "Rà soát backend.", "Xác nhận migration/health.", "Mốc kế hoạch."),
        ("C", "Rà soát frontend.", "Demo luồng chính.", "Mốc kế hoạch."),
        ("D", "Chạy test.", "Báo cáo defect/residual risk.", "Mốc kế hoạch."),
        ("E", "Hoàn thiện hồ sơ.", "Lưu biên bản và action item.", "Mốc kế hoạch."),
    ]),
]


def build_daily_scrum():
    doc = style_document(Document(), "DAILY SCRUM | TASKSYNCENTERPRISE")
    add_cover(
        doc,
        "Mô hình phát triển phần mềm Agile",
        "NHẬT KÝ DAILY SCRUM",
        "Sprint 1 và Sprint 2 sau khi gộp",
    )
    doc.add_heading("Hướng dẫn sử dụng", level=1)
    add_body(
        doc,
        "Daily Scrum được timebox 15 phút. Mỗi thành viên nêu công việc đã làm, kế hoạch 24 giờ tiếp theo và trở ngại. "
        "Scrum Master ghi action xử lý blocker; Product Owner chỉ tham gia khi cần làm rõ ưu tiên."
    )
    add_note(
        doc,
        "Tính xác thực",
        "Các mốc đến 29/07/2026 được tổng hợp từ tài liệu dự án. Các mốc 30/07-27/08/2026 là baseline/kịch bản dự kiến và phải được nhóm thay bằng biên bản thực tế trước khi nộp như chứng cứ đã họp.",
        color="FFF7E8",
    )
    add_table(
        doc,
        ["Mã", "Thành viên", "Vai trò"],
        [(code, name, role) for code, name, role in TEAM],
        [900, 3300, 5160],
        font_size=9.2,
    )
    for idx, (sprint, date, focus, rows) in enumerate(DAILY_MEETINGS, start=1):
        doc.add_page_break()
        doc.add_heading(f"Daily Scrum {idx:02d} - {sprint}", level=1)
        add_table(
            doc,
            ["Ngày", "Thời lượng", "Trọng tâm", "Tham dự"],
            [(date, "07:00-07:15", focus, "A, B, C, D, E")],
            [1600, 1700, 4360, 1700],
            font_size=9,
            header_fill=LIGHT_BLUE,
        )
        add_table(
            doc,
            ["TV", "Đã làm", "Sẽ làm", "Trở ngại"],
            rows,
            [700, 2860, 2860, 2940],
            font_size=8.4,
        )
        doc.add_heading("Action sau Daily", level=2)
        if "kế hoạch" in rows[0][3].lower():
            add_bullets(
                doc,
                [
                    "Xác nhận tiến độ thật trên Product Backlog trước cuộc họp tiếp theo.",
                    "Gắn evidence/PR/test run vào User Story liên quan.",
                    "Cập nhật blocker và người chịu trách nhiệm xử lý.",
                ],
            )
        else:
            add_bullets(
                doc,
                [
                    "Scrum Master theo dõi blocker và cập nhật trạng thái trong ngày.",
                    "Assignee cập nhật Task/Backlog sau khi hoàn thành Acceptance Criteria.",
                    "QA bổ sung test case cho lỗi hoặc thay đổi workflow phát sinh.",
                ],
            )
        add_body(doc, "Người chủ trì: B - Nguyễn Đức Mạnh | Thư ký: E - Nguyễn Anh Tuấn", italic=True)
    save_doc(doc, AGILE_DIR / "Daily_Scrum_2_Sprint_TaskSyncEnterprise.docx")


DIAGRAM_DESCRIPTIONS = [
    ("Kiến trúc phần mềm", "01_System_Context.png", "Sơ đồ ngữ cảnh hệ thống", "Xác định tác nhân, ranh giới hệ thống và kênh giao tiếp chính.", [
        "Admin, Manager và Employee thao tác qua React Web Application.",
        "Nền tảng cung cấp REST API và WebSocket.",
        "Email/Push là tích hợp tùy chọn, không phải nguồn dữ liệu chuẩn.",
    ]),
    ("Kiến trúc phần mềm", "02_Container_Architecture.png", "Kiến trúc container", "Mô tả các container và phụ thuộc vận hành.", [
        "Nginx là điểm vào production.",
        "Frontend và Backend tách container nhưng cùng một sản phẩm.",
        "SQL Server giữ dữ liệu chuẩn; Redis chỉ hỗ trợ cache-aside.",
    ]),
    ("Kiến trúc phần mềm", "03_Backend_Layers.png", "Kiến trúc phân lớp Backend", "Giải thích ranh giới trách nhiệm trong mã nguồn.", [
        "Router chịu trách nhiệm validation và authorization boundary.",
        "Service chứa workflow nghiệp vụ; CRUD tập trung persistence.",
        "Event realtime chỉ phát sau khi commit.",
    ]),
    ("Kiến trúc phần mềm", "04_Domain_Model.png", "Mô hình miền", "Thống nhất quan hệ Project, Epic, Sprint, Backlog, Task và Employee.", [
        "Project là phạm vi gốc.",
        "Backlog Item/User Story được chọn vào Sprint và phân rã thành Task.",
        "Task Assignment nối Task với Employee có quyền truy cập Project.",
    ]),
    ("Kiến trúc phần mềm", "05_Deployment.png", "Sơ đồ triển khai", "Trình bày topology Docker Compose.", [
        "Nginx định tuyến tới Frontend và Backend.",
        "SQL Server, Redis và upload sử dụng volume.",
        "sqlserver-init bảo đảm database tồn tại; Alembic cập nhật schema.",
    ]),
    ("Agile", "01_Phan_ra_chuc_nang.png", "Phân rã chức năng", "Dùng để nhóm 40 User Stories thành các miền dễ quản lý.", [
        "Mỗi nhánh là một nhóm năng lực nghiệp vụ.",
        "Phân rã giúp định nghĩa Epic và tránh backlog trùng lặp.",
        "Có thể dùng trực tiếp ở phần khảo sát/phân tích yêu cầu.",
    ]),
    ("Agile", "02_Use_Case_Tong_Quat.png", "Use Case tổng quát", "Mô tả năng lực nhìn từ ba nhóm tác nhân.", [
        "Admin quản trị tài khoản, tổ chức và chính sách.",
        "Manager lập kế hoạch, quản lý Task và duyệt workflow.",
        "Employee thực thi Task, cộng tác và theo dõi thông báo.",
    ]),
    ("Agile", "03_Activity_Task_Lifecycle.png", "Activity - Vòng đời Task", "Dùng để giải thích luồng từ Backlog đến Done.", [
        "Validation và RBAC diễn ra trước mutation.",
        "Task chỉ Done khi Acceptance Criteria đạt.",
        "Nếu review không đạt, Task quay lại In Progress.",
    ]),
    ("Agile", "04_Activity_Sprint.png", "Activity - Chu trình Sprint", "Thể hiện vòng kiểm tra và thích ứng của Scrum.", [
        "Planning tạo Sprint Goal và Sprint Backlog.",
        "Daily Scrum điều chỉnh kế hoạch ngày.",
        "Review và Retrospective tạo feedback cho Sprint tiếp theo.",
    ]),
    ("Agile", "05_Scrum_Workflow.png", "Scrum Workflow", "Tóm tắt luồng trách nhiệm Product Owner và Development Team.", [
        "Product Owner ưu tiên Product Backlog.",
        "Development Team sở hữu cách thực hiện Sprint Backlog.",
        "Review kiểm tra Increment; Retrospective cải tiến quy trình.",
    ]),
]


def build_diagram_guide():
    doc = style_document(Document(), "THUYẾT MINH SƠ ĐỒ | TASKSYNCENTERPRISE")
    add_cover(
        doc,
        "Kiến trúc phần mềm và Agile",
        "THUYẾT MINH BỘ SƠ ĐỒ",
        "Sơ đồ kiến trúc, phân rã chức năng, Use Case và Activity",
    )
    doc.add_heading("Danh mục sơ đồ", level=1)
    add_table(
        doc,
        ["STT", "Nhóm", "Tên sơ đồ", "Mục đích"],
        [(i, group, title, purpose) for i, (group, _, title, purpose, _) in enumerate(DIAGRAM_DESCRIPTIONS, 1)],
        [700, 1800, 2800, 4060],
        font_size=8.7,
    )
    add_note(
        doc,
        "File nguồn",
        "Mỗi sơ đồ có bản PNG để chèn báo cáo và bản .mmd để chỉnh sửa bằng Mermaid. "
        "Các sơ đồ đã được tách vào hai thư mục Kiến trúc phần mềm và Agile.",
    )
    for i, (group, filename, title, purpose, notes) in enumerate(DIAGRAM_DESCRIPTIONS, 1):
        doc.add_page_break()
        doc.add_heading(f"{i}. {title}", level=1)
        folder = ARCH_DIAGRAM_DIR if group == "Kiến trúc phần mềm" else AGILE_DIAGRAM_DIR
        add_diagram(doc, folder / filename, f"Hình {i}. {title}")
        add_body(doc, f"Mục đích: {purpose}", bold_lead="Mục đích:")
        doc.add_heading("Cách giải thích trong báo cáo", level=2)
        add_bullets(doc, notes)
        doc.add_heading("Vị trí đề xuất", level=2)
        add_body(
            doc,
            "Chèn sơ đồ ngay sau đoạn giới thiệu phạm vi liên quan; dùng caption thống nhất và tham chiếu bằng số hình trong phần phân tích."
        )
    save_doc(doc, DIAGRAM_DIR / "Thuyet_minh_So_do_TaskSyncEnterprise.docx")


def build_readme():
    content = """# Bộ hồ sơ Report - TaskSyncEnterprise

## Cấu trúc

- `Môn kiến trúc phần mềm/`: báo cáo kiến trúc phần mềm.
- `Agile/`: báo cáo Agile/Scrum, Daily Scrum và workbook Product Backlog.
- `Sơ đồ Diagram/`: sơ đồ PNG, nguồn Mermaid và tài liệu thuyết minh.
- `Tham khảo/`: tài liệu nguồn do nhóm cung cấp, được giữ nguyên.

## Quy ước Sprint đã chuẩn hóa

- Sprint cũ 1 + Sprint cũ 2 → Sprint 1 mới: 02/07-30/07/2026, US01-US18, 107 SP.
- Sprint cũ 3 + Sprint cũ 4 → Sprint 2 mới: 30/07-27/08/2026, US19-US40, 135 SP.

Tổng chi tiết: 242 SP. File Summary nguồn ghi 228 SP nhưng không khớp tổng từng dòng Product Backlog.

## Lưu ý trước khi nộp

1. Xác nhận lại họ tên, MSSV và vai trò do tài liệu nguồn có một số khác biệt.
2. Các mốc sau ngày 29/07/2026 là baseline/kế hoạch; thay bằng evidence thực tế nếu dùng làm biên bản đã họp.
3. Cập nhật số trang/mục lục trong Word nếu chỉnh sửa nội dung.
4. Không mô tả hệ thống là production-ready khi Redis/Docker/dependency risk chưa được xác nhận lại.
"""
    (REPORT / "README_Bo_ho_so.md").write_text(content, encoding="utf-8")


def main():
    build_diagrams()
    build_architecture_report()
    build_agile_report()
    build_daily_scrum()
    build_diagram_guide()
    build_readme()
    print("Built report documents and diagrams.")


if __name__ == "__main__":
    main()
