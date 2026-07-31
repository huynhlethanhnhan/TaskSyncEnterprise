from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"E:\TaskSyncEnterprise")
SOURCE = ROOT / "Report" / "Tham khảo"
OUT = ROOT / "tmp" / "report_build" / "reference_extract"
OUT.mkdir(parents=True, exist_ok=True)


def extract_pdf(path: Path) -> dict:
    pages = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            pages.append({"page": index, "text": text})
    (OUT / f"{path.stem}_pages.json").write_text(
        json.dumps(pages, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (OUT / f"{path.stem}.txt").write_text(
        "\n\n".join(f"===== TRANG {p['page']} =====\n{p['text']}" for p in pages),
        encoding="utf-8",
    )
    return {
        "file": path.name,
        "page_count": len(pages),
        "nonempty_pages": sum(bool(p["text"].strip()) for p in pages),
        "characters": sum(len(p["text"]) for p in pages),
    }


def extract_docx(path: Path) -> dict:
    doc = Document(path)
    lines = []
    for idx, p in enumerate(doc.paragraphs, start=1):
        text = p.text.strip()
        if text:
            lines.append(f"[P{idx}][{p.style.name}] {text}")
    table_dump = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        table_dump.append({"table": table_index, "rows": rows})
    (OUT / f"{path.stem}_paragraphs.txt").write_text("\n".join(lines), encoding="utf-8")
    (OUT / f"{path.stem}_tables.json").write_text(
        json.dumps(table_dump, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    rel_images = [
        rel.target_ref
        for rel in doc.part.rels.values()
        if "image" in rel.reltype
    ]
    return {
        "file": path.name,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "image_rel_count": len(rel_images),
        "nonempty_paragraphs": len(lines),
    }


def extract_docx_media(path: Path) -> int:
    target = OUT / f"{path.stem}_media"
    target.mkdir(parents=True, exist_ok=True)
    count = 0
    with zipfile.ZipFile(path) as archive:
        for member in archive.namelist():
            if not member.startswith("word/media/") or member.endswith("/"):
                continue
            output = target / Path(member).name
            output.write_bytes(archive.read(member))
            count += 1
    return count


def build_media_contact(path: Path) -> None:
    folder = OUT / f"{path.stem}_media"
    files = sorted(folder.iterdir())
    cards = []
    font = ImageFont.load_default()
    for file in files:
        try:
            image = Image.open(file).convert("RGB")
        except Exception:
            continue
        scale = min(360 / image.width, 240 / image.height, 1)
        image = image.resize(
            (max(1, int(image.width * scale)), max(1, int(image.height * scale))),
            Image.Resampling.LANCZOS,
        )
        card = Image.new("RGB", (390, 285), "white")
        draw = ImageDraw.Draw(card)
        draw.text((10, 8), file.name, fill="black", font=font)
        card.paste(image, ((390 - image.width) // 2, 35))
        cards.append(card)
    if not cards:
        return
    cols = 4
    rows = (len(cards) + cols - 1) // cols
    contact = Image.new("RGB", (cols * 390, rows * 285), (225, 230, 235))
    for index, card in enumerate(cards):
        contact.paste(card, ((index % cols) * 390, (index // cols) * 285))
    contact.save(OUT / f"{path.stem}_contact.png")


summary = []
summary.append(extract_pdf(SOURCE / "LTHN_Agile_Final.pdf"))
for name in (
    "Co_su_Ly_Thuyet.docx",
    "Cac_bang_table.docx",
    "bao-cao-do-an-tasksyncenterprise (1).docx",
):
    path = SOURCE / name
    item = extract_docx(path)
    item["extracted_media_count"] = extract_docx_media(path)
    build_media_contact(path)
    summary.append(item)

(OUT / "summary.json").write_text(
    json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
)
print(json.dumps(summary, ensure_ascii=False, indent=2))
