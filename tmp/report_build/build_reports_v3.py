from __future__ import annotations

import json
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_documents as base
import build_reports_v2 as legacy


ROOT = Path(r"E:\TaskSyncEnterprise")
REPORT = ROOT / "Report"
AGILE_DIR = REPORT / "Agile"
ARCH_DIR = REPORT / "Môn kiến trúc phần mềm"
DIAGRAM_DIR = REPORT / "Sơ đồ Diagram"
AGILE_DIAGRAM_DIR = DIAGRAM_DIR / "Agile"
DATA = json.loads((ROOT / "tmp" / "report_build" / "report_data_v3.json").read_text(encoding="utf-8"))

TEAM = [
    ("A", "Huỳnh Lê Thành Nhân", "2000004897", "Product Owner / Full-stack Developer"),
    ("B", "Nguyễn Đức Mạnh", "2200010420", "Scrum Master tạm thời / Hỗ trợ kiểm thử"),
    ("C", "Nguyễn Lê Huy Hoàng", "2311554285", "Hỗ trợ kiểm thử và demo"),
    ("D", "Phạm Tuấn Anh", "", "Sơ đồ và tài liệu được phân công"),
    ("E", "Nguyễn Anh Tuấn", "", "Nghiên cứu tài liệu được phân công"),
]
ARCH_AUTHOR = [TEAM[0]]


def normalized_backlog() -> list[dict]:
    result = []
    for item in DATA["backlog"]:
        result.append(
            {
                "id": item["id"],
                "jira": item["jira"],
                "epic": item["epic"],
                "feature": item["epic_name"],
                "title_en": item["title_en"],
                "title": item["title"],
                "description_en": "",
                "priority": item["priority"],
                "sp": int(item["sp"]),
                "value": "Cao" if item["priority"] == "Cao" else "Trung bình",
                "source_assignee": item["jira_assignee"],
                "sprint": item["sprint"],
                "source_status": item["jira_status"],
                "verified_status": item["verified_status"],
                "verified_owner": item["verified_owner"],
                "acceptance": item["acceptance"],
                "source": item["source"],
            }
        )
    return result


BACKLOG = normalized_backlog()
SPRINT_1 = [item for item in BACKLOG if item["sprint"] == "Sprint 1"]
SPRINT_2 = [item for item in BACKLOG if item["sprint"] == "Sprint 2"]


def cover(doc: Document, course: str, title: str, subtitle: str) -> None:
    if "agile" in course.lower():
        course = "Mô hình Agile"
    elif "kiến trúc" in course.lower():
        course = "Môn Kiến Trúc Phần Mềm"
    for text, size in (
        ("TRƯỜNG ĐẠI HỌC NGUYỄN TẤT THÀNH", 14),
        ("VIỆN ĐÀO TẠO QUỐC TẾ NIIE", 14),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        base.set_run_font(run, "Times New Roman", size, "000000", True)
    logo = legacy.THEORY_MEDIA_DIR / "image1.jpg"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(1.1))
    doc.add_paragraph()
    for text, size in (
        ("BÁO CÁO MÔN HỌC", 15),
        (course.upper(), 15),
        (title.upper(), 19),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text)
        base.set_run_font(run, "Times New Roman", size, "000000", True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run(subtitle)
    base.set_run_font(run, "Times New Roman", 13, "444444", False, True)

    rows = [
        ("Giảng viên hướng dẫn", "ThS. Trần Thanh Nhã"),
        ("Nhóm", "JD"),
        ("Lớp học phần", "........................................................"),
        ("Học kỳ / Năm học", "Học kỳ 3 - Năm học 2025-2026"),
    ]
    rows.extend(
        (f"Thành viên {code}", f"{student_id or '................'} - {name} ({role})")
        for code, name, student_id, role in legacy.TEAM
    )
    base.add_table(doc, ["Thông tin", "Nội dung"], rows, [2500, 6860], font_size=9.0, header_fill="D9EAF7")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("TP. Hồ Chí Minh, tháng 7 năm 2026")
    base.set_run_font(run, "Times New Roman", 13, "000000", False, True)
    doc.add_page_break()


def front_matter(doc: Document, document_kind: str, figures: list[str], tables: list[str]) -> None:
    doc.add_heading("LỜI CẢM ƠN", level=1)
    base.add_body(
        doc,
        "Nhóm xin chân thành cảm ơn ThS. Trần Thanh Nhã đã hướng dẫn cách tổ chức hồ sơ theo hai Sprint, "
        "góp ý về phạm vi giữa kỳ và yêu cầu trình bày bằng chứng rõ ràng. Báo cáo ưu tiên tính trung thực: "
        "trạng thái trên Jira được tách khỏi trạng thái đã xác minh bằng mã nguồn, Git, kiểm thử và demo.",
    )
    doc.add_page_break()
    doc.add_heading("CƠ SỞ XÂY DỰNG TÀI LIỆU", level=1)
    sources = [
        ("1", "Báo cáo mẫu", "LTHN_Agile_Final.pdf - dùng cấu trúc và cách tổ chức, không sao chép nội dung đề tài khác"),
        ("2", "Cơ sở lý thuyết", "Co_su_Ly_Thuyet.docx"),
        ("3", "Cơ sở dữ liệu và giao diện", "Cac_bang_table.docx"),
        ("4", "Jira", "Dự án JD_Scrum_TaskSnycEnterprise - 50 issue"),
        ("5", "Git", "79 commit trong giai đoạn 04/07-23/07/2026"),
        ("6", "Mã nguồn", r"E:\TaskSyncEnterprise"),
        ("7", "Thông tin nhóm", "Nội dung Product Owner cung cấp trong quá trình rà soát"),
    ]
    base.add_table(doc, ["STT", "Nguồn", "Cách sử dụng"], sources, [700, 2100, 6560], font_size=9.0)
    doc.add_page_break()
    doc.add_heading("DANH MỤC HÌNH ẢNH", level=1)
    base.add_table(doc, ["STT", "Tên hình"], [(i, name) for i, name in enumerate(figures, 1)], [800, 8560], font_size=8.4)
    doc.add_page_break()
    doc.add_heading("DANH MỤC BẢNG BIỂU", level=1)
    base.add_table(doc, ["STT", "Tên bảng"], [(i, name) for i, name in enumerate(tables, 1)], [800, 8560], font_size=8.4)
    doc.add_page_break()
    legacy.add_toc_field(doc)


def figure_caption(doc: Document, filename: str, description: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"Hình X. {description}")
    base.set_run_font(run, "Times New Roman", 10, "666666", False, True)


def configure_legacy(team: list[tuple]) -> None:
    legacy.TEAM = team
    legacy.BACKLOG = BACKLOG
    legacy.SPRINT_1 = SPRINT_1
    legacy.SPRINT_2 = SPRINT_2
    legacy.EPIC_VI.update(DATA["epics"])
    legacy.add_academic_cover = cover
    legacy.add_front_matter = front_matter
    legacy.add_placeholder = figure_caption


def add_status_legend(doc: Document) -> None:
    doc.add_heading("Quy ước trạng thái", level=2)
    rows = [
        ("Hoàn thành và có thể demo", "Có mã nguồn và luồng chính có thể trình bày."),
        ("Hoàn thành phần chính", "Đã có chức năng cốt lõi nhưng vẫn cần kiểm thử thêm."),
        ("Sơ khai / Đang hoàn thiện", "Đã có một phần giao diện hoặc xử lý."),
        ("Chưa xác minh", "Chưa đủ bằng chứng để kết luận hoàn thành."),
        ("Dự kiến Sprint 2", "Không tính vào kết quả giữa kỳ."),
    ]
    base.add_table(doc, ["Trạng thái", "Ý nghĩa"], rows, [3100, 6260], font_size=9.0, header_fill="D9EAF7")


def build_midterm_report() -> None:
    configure_legacy(TEAM)
    doc = legacy.academic_document("BÁO CÁO GIỮA KỲ SPRINT 1 | TASKSYNCENTERPRISE")
    cover(
        doc,
        "Mô hình Agile",
        "TaskSyncEnterprise",
        "Báo cáo giữa kỳ - Sprint 1 (02/07/2026 đến 23/07/2026)",
    )
    figures = [
        "Phân rã chức năng TaskSyncEnterprise",
        "Use Case tổng quát",
        "Activity Diagram vòng đời công việc",
        "Quy trình Scrum",
        "Giao diện đăng nhập",
        "Giao diện dự án",
        "Giao diện công việc và Kanban",
        "Trung tâm thông báo và lịch công việc",
    ]
    tables = [
        "Thành viên và vai trò",
        "Quy ước gộp Sprint",
        "Phạm vi Product Backlog Sprint 1",
        "Timeline thực hiện",
        "Kết quả chức năng",
        "Daily Scrum và cập nhật tiến độ",
        "Lỗi và biện pháp xử lý",
        "Kế hoạch Sprint 2",
    ]
    front_matter(doc, "Giữa kỳ", figures, tables)

    legacy.add_part_title(doc, "1", "Tổng quan đề tài", first=True)
    doc.add_heading("1. Bối cảnh", level=1)
    base.add_body(
        doc,
        "TaskSyncEnterprise là nền tảng quản lý công việc doanh nghiệp, tập trung vào xác thực, tổ chức nhân sự, "
        "dự án, công việc, bảng Kanban, lịch, thông báo và khả năng theo dõi vận hành. Báo cáo giữa kỳ chỉ trình bày "
        "kết quả của Sprint 1 sau khi giảng viên yêu cầu gộp Sprint cũ 1 và 2.",
    )
    doc.add_heading("2. Mục tiêu Sprint 1", level=1)
    base.add_body(
        doc,
        "Tạo Increment có thể trình bày với Login/Auth, Project, Task/Kanban, Employee ở mức sơ khai, "
        "Notification và Calendar; đồng thời ổn định nền tảng Backend/Frontend, database, logging, cache và CI.",
    )
    doc.add_heading("3. Quy ước gộp Sprint", level=1)
    base.add_table(
        doc,
        ["Tên sử dụng", "Nguồn được gộp", "Thời gian", "Mục đích"],
        [
            ("Sprint 1", "Sprint cũ 1 + 2", "02/07-23/07/2026", "Báo cáo giữa kỳ"),
            ("Sprint 2", "Sprint cũ 3 + 4", "24/07-20/08/2026", "Kế hoạch và tiếp tục xác minh"),
        ],
        [1700, 2400, 2100, 3160],
        font_size=9.2,
        header_fill="D9EAF7",
    )
    doc.add_heading("4. Nhóm Scrum", level=1)
    base.add_table(doc, ["Mã", "Họ tên", "MSSV", "Vai trò"], TEAM, [650, 2200, 1700, 4810], font_size=8.9)
    base.add_note(
        doc,
        "Cách đọc đóng góp",
        "Vai trò được giao không đồng nghĩa kết quả đã được xác minh. Báo cáo đóng góp thành viên được tách thành tài liệu riêng và dùng cách diễn đạt trung tính.",
        "DDEBF7",
    )

    legacy.add_part_title(doc, "2", "Product Backlog và phạm vi Sprint 1")
    doc.add_heading("1. Cơ sở lập Product Backlog", level=1)
    base.add_body(
        doc,
        "Product Backlog gồm 50 User Story: 38 Story liên kết Jira JD-14 đến JD-51 và 12 Story được bổ sung từ mã nguồn/Git "
        "để phản ánh các chức năng thật như Calendar, Notification Center, Dashboard, health check, logging, Redis và Docker/CI.",
    )
    add_status_legend(doc)
    doc.add_heading("2. Danh sách Sprint 1", level=1)
    rows = [
        (
            item["id"],
            item["jira"] or "-",
            item["title"],
            item["priority"],
            item["sp"],
            item["verified_status"],
        )
        for item in SPRINT_1
    ]
    base.add_table(doc, ["Mã", "Jira", "User Story", "Ưu tiên", "SP", "Xác minh"], rows, [650, 750, 3820, 1000, 600, 2540], font_size=7.5)
    base.add_diagram(doc, AGILE_DIAGRAM_DIR / "01_Phan_ra_chuc_nang.png", "Hình 1. Phân rã chức năng TaskSyncEnterprise")
    base.add_diagram(doc, AGILE_DIAGRAM_DIR / "02_Use_Case_Tong_Quat.png", "Hình 2. Use Case tổng quát")

    legacy.add_part_title(doc, "3", "Quá trình thực hiện Sprint 1")
    doc.add_heading("1. Giai đoạn chuẩn bị", level=1)
    timeline = [
        ("18/06-24/06", "Thảo luận và tranh luận đề tài; phản hồi nhóm còn hạn chế."),
        ("25/06", "Product Owner đề xuất Sprint Planning; chưa nhận đủ phản hồi vào buổi tối."),
        ("26/06-01/07", "Tạo Epic/User Story ban đầu trên Jira và chuẩn bị môi trường."),
        ("02/07-08/07", "Khởi động Sprint, ổn định cấu trúc, database và Login/Auth."),
        ("09/07-15/07", "Củng cố Backend, Project, Task/Kanban, Notification nền tảng và quan sát hệ thống."),
        ("16/07-23/07", "Hoàn thiện hạ tầng, giao diện, kiểm thử chấp nhận, sửa lỗi và chuẩn bị demo."),
    ]
    base.add_table(doc, ["Thời gian", "Nội dung"], timeline, [1900, 7460], font_size=9.1)
    doc.add_heading("2. Các phase kỹ thuật phát sinh", level=1)
    phases = [
        ("Phase 1", "Dọn dẹp và ổn định dự án", "Được dùng làm định hướng xử lý lỗi."),
        ("Phase 2", "Xác minh hạ tầng", "Tiếp tục cập nhật trong Sprint 2."),
        ("Phase 3", "Nền tảng cơ sở dữ liệu", "Database đã có; tiếp tục rà schema và migration."),
        ("Phase 4", "Xác thực", "Logic được làm lại và bổ sung các phần còn thiếu."),
        ("Phase 5", "Phân quyền", "Có nền tảng RBAC; tiếp tục kiểm thử trong Sprint 2."),
    ]
    base.add_table(doc, ["Phase", "Tên tiếng Việt", "Cách sử dụng"], phases, [1200, 3000, 5160], font_size=9.0)
    base.add_note(
        doc,
        "Lưu ý",
        "Năm phase trên là kế hoạch kỹ thuật khi hệ thống phát sinh lỗi, không được trình bày như năm Sprint và không thay thế Product Backlog.",
        "FFF2CC",
    )
    doc.add_heading("3. Daily Scrum và cập nhật tiến độ", level=1)
    daily_summary = [
        ("02/07", "Bắt đầu Sprint", "Chốt phạm vi, rà dự án và ưu tiên Login/Auth."),
        ("06/07", "Cập nhật tại trường", "Rà Backend/Frontend, lỗi 404 và trạng thái đăng nhập."),
        ("08/07", "Cập nhật tại trường", "Ổn định database, cấu hình và nền tảng Backend."),
        ("13/07", "Cập nhật tại trường", "Health, logging, cache, Notification và kiểm thử tích hợp."),
        ("15/07", "Cập nhật tại trường", "Chạy thử Login, Task, Project; giới hạn phạm vi demo."),
        ("20/07", "Cập nhật tại trường", "Docker, reverse proxy, backup và chuẩn bị demo."),
        ("22/07", "Cập nhật tại trường", "Rà Project, Task, Notification, Calendar và blocker."),
        ("23/07", "Sprint Review", "Chốt kết quả đã xác minh và chuyển phần chưa ổn định sang Sprint 2."),
    ]
    base.add_table(doc, ["Ngày", "Loại", "Nội dung chính"], daily_summary, [1300, 2300, 5760], font_size=8.8)
    base.add_body(
        doc,
        "Các buổi trao đổi chính diễn ra trực tiếp tại trường vào thứ Hai và thứ Tư, thường từ 12:00 đến 13:00, "
        "thời lượng 15-30 phút; buổi tranh luận có thể kéo dài khoảng một giờ. Các ngày còn lại chỉ được gọi là cập nhật tiến độ, "
        "không khẳng định toàn bộ nhóm đã họp.",
    )

    legacy.add_part_title(doc, "4", "Kết quả kỹ thuật và minh chứng")
    doc.add_heading("1. Kết quả chức năng tại mốc 23/07", level=1)
    result_rows = [
        ("Login/Auth", "Hoàn thành phần chính; có thể demo.", "Git, mã nguồn, Jira JD-14 đến JD-17"),
        ("Task/Kanban", "Hoàn thành phần chính; có thể demo.", "Git, mã nguồn, Jira JD-27 đến JD-31"),
        ("Project", "Có thể demo luồng chính.", "Mã nguồn ProjectPage/API"),
        ("Employee", "Sơ khai.", "Mã nguồn EmployeePage/API"),
        ("Dashboard", "Sơ khai.", "DashboardPage và dashboard API"),
        ("Notification", "Đã có nền tảng và giao diện; cần kiểm thử thêm.", "Notification framework, NotificationsPage"),
        ("Calendar", "Đã có và được hỗ trợ kiểm thử.", "CalendarPage"),
        ("Department/Team", "Chưa chạy ổn định tại mốc giữa kỳ.", "Chuyển tiếp Sprint 2"),
        ("Nghỉ phép, Audit, báo cáo nâng cao", "Có cấu trúc hoặc mã nguồn nhưng không tính là kết quả giữa kỳ.", "Ngoài phạm vi demo Sprint 1"),
    ]
    base.add_table(doc, ["Chức năng", "Kết quả", "Bằng chứng"], result_rows, [2300, 3600, 3460], font_size=8.8)
    for description in (
        "Giao diện đăng nhập sau khi kiểm thử",
        "Giao diện quản lý dự án",
        "Giao diện quản lý công việc và bảng Kanban",
        "Giao diện quản lý nhân viên ở mức sơ khai",
        "Trung tâm thông báo và lịch công việc",
    ):
        figure_caption(doc, "", description)

    doc.add_heading("2. Bằng chứng Git", level=1)
    base.add_body(
        doc,
        "Trong giai đoạn 04/07-23/07/2026, Git ghi nhận 79 commit dưới ba biến thể tên tác giả "
        "Huynh Le Thanh Nhan, huynh và huynhlethanhnhan. Không có commit mang tên bốn thành viên còn lại trong giai đoạn này. "
        "Vì vậy Git được dùng để xác minh đóng góp kỹ thuật của Huỳnh Lê Thành Nhân, không dùng để phủ nhận các hoạt động kiểm thử trực tiếp chưa được ghi nhận.",
    )
    commits = [
        ("04/07", "e4b318a", "Khởi tạo TaskSync Enterprise V2 Full-stack."),
        ("07/07", "8db87cd", "Chuẩn hóa model SQLAlchemy."),
        ("09/07", "71d3564", "Hoàn thiện nền tảng hạ tầng Backend."),
        ("10/07", "17fb551", "Khung hạ tầng và Notification."),
        ("11/07", "83d2a69", "Redis caching."),
        ("12/07", "e9918fe-3a542b4", "Health, metrics, logging và tracing."),
        ("20/07", "052f835", "Chuẩn bị bản phát hành v1.0.0-rc1."),
        ("23/07", "4475139", "Thiết kế lại UI/UX, tích hợp workflow và kiểm toán production."),
    ]
    base.add_table(doc, ["Ngày", "Commit", "Nội dung"], commits, [1300, 2200, 5860], font_size=8.8)
    doc.add_heading("3. Lỗi và cách xử lý", level=1)
    bugs = [
        ("Database", "Khác biệt SQL Server/SQLite, migration và default", "Chuẩn hóa model, migration và test isolation."),
        ("Login/404", "Điều hướng hoặc route không tồn tại", "Rà ProtectedRoute, AppRouter và trạng thái xác thực."),
        ("Cấu hình", "Thông tin môi trường không đồng bộ", "Dùng cấu hình mẫu và loại bỏ đường dẫn máy cá nhân."),
        ("CI/Test", "Timeout và dependency", "Tách test, sửa dependency và workflow."),
        ("Notification", "Thiếu cơ chế đa kênh/retry", "Bổ sung framework, repository và dịch vụ."),
    ]
    base.add_table(doc, ["Nhóm lỗi", "Hiện tượng/nguyên nhân", "Xử lý"], bugs, [1900, 3500, 3960], font_size=8.9)

    legacy.add_part_title(doc, "5", "Đặc tả các luồng demo chính")
    for number, item_id in enumerate(("US01", "US02", "US11", "US14", "US15", "US43", "US45"), 1):
        item = next(item for item in BACKLOG if item["id"] == item_id)
        legacy.add_use_case_spec(doc, item, f"5.{number}")
    base.add_diagram(doc, AGILE_DIAGRAM_DIR / "03_Activity_Task_Lifecycle.png", "Hình 3. Activity Diagram vòng đời công việc")
    base.add_diagram(doc, AGILE_DIAGRAM_DIR / "05_Scrum_Workflow.png", "Hình 4. Quy trình Scrum")

    legacy.add_part_title(doc, "6", "Sprint Review và kế hoạch tiếp theo")
    doc.add_heading("1. Sprint Review", level=1)
    base.add_bullets(
        doc,
        [
            "Increment có thể trình bày các luồng Login/Auth, Project, Task/Kanban, Employee sơ khai, Notification và Calendar.",
            "Một số module đã tồn tại trong mã nguồn nhưng chưa đủ kiểm thử để tính là hoàn thành giữa kỳ.",
            "Jira có 14 issue Done, 3 In Progress, 21 To Do và 12 Story bổ sung chưa có issue riêng.",
            "Trạng thái người được giao trên Jira không được dùng thay cho bằng chứng đóng góp.",
        ],
    )
    doc.add_heading("2. Sprint Retrospective", level=1)
    retro = [
        ("Điểm tốt", "Product Owner duy trì tiến độ, sửa lỗi liên tục và tạo được phiên bản demo."),
        ("Khó khăn", "Phản hồi nhóm chậm, thiếu bằng chứng họp và một số công việc được giao chưa có sản phẩm bàn giao."),
        ("Cải tiến", "Giới hạn WIP, yêu cầu bằng chứng cho mỗi hạng mục, cập nhật Jira ngay sau khi xác minh."),
        ("Hành động", "Sprint 2 tập trung kiểm thử, ổn định Department/Team và hoàn thiện module còn lại."),
    ]
    base.add_table(doc, ["Nhóm", "Nội dung"], retro, [2000, 7360], font_size=9.0)
    doc.add_heading("3. Kế hoạch Sprint 2", level=1)
    plan_rows = [(item["id"], item["title"], item["sp"], "Dự kiến/tiếp tục xác minh") for item in SPRINT_2]
    base.add_table(doc, ["Mã", "User Story", "SP", "Trạng thái"], plan_rows, [700, 5160, 700, 2800], font_size=8.0)
    doc.add_heading("4. Thuật ngữ", level=1)
    glossary = [
        ("Agile", "Phương pháp phát triển linh hoạt", "Cách tổ chức dự án"),
        ("Sprint", "Chu kỳ phát triển có mục tiêu", "Sprint 1 và Sprint 2"),
        ("Product Backlog", "Danh sách yêu cầu sản phẩm", "Excel và báo cáo"),
        ("User Story", "Câu chuyện người dùng", "Đơn vị yêu cầu"),
        ("Daily Scrum", "Cập nhật tiến độ hằng ngày", "Biên bản Sprint"),
        ("Authentication", "Xác thực", "Đăng nhập và phiên"),
        ("Authorization", "Phân quyền", "Vai trò và quyền truy cập"),
        ("Project Cleanup & Stabilization", "Dọn dẹp và ổn định dự án", "Phase kỹ thuật"),
        ("Infrastructure Validation", "Xác minh hạ tầng", "Phase kỹ thuật"),
        ("Database Foundation", "Nền tảng cơ sở dữ liệu", "Phase kỹ thuật"),
    ]
    base.add_table(doc, ["Thuật ngữ", "Nghĩa tiếng Việt", "Vị trí sử dụng"], glossary, [3000, 3600, 2760], font_size=8.8)
    base.save_doc(doc, AGILE_DIR / "Bao_cao_Giua_ky_Sprint_1_TaskSyncEnterprise.docx")


DAILY_EVENTS = [
    (date(2026, 7, 2), "Bắt đầu Sprint", "Chốt phạm vi, kiểm tra dự án và ưu tiên Login/Auth."),
    (date(2026, 7, 6), "Cập nhật tại trường", "Rà Backend/Frontend, trạng thái đăng nhập và lỗi 404."),
    (date(2026, 7, 8), "Cập nhật tại trường", "Ổn định database, cấu hình và nền tảng Backend."),
    (date(2026, 7, 13), "Cập nhật tại trường", "Rà health, logging, cache, Notification và kiểm thử."),
    (date(2026, 7, 15), "Cập nhật tại trường", "Chạy thử Login, Project, Task và phạm vi demo."),
    (date(2026, 7, 20), "Cập nhật tại trường", "Docker, reverse proxy, backup và chuẩn bị demo."),
    (date(2026, 7, 22), "Cập nhật tại trường", "Rà Project, Task, Notification, Calendar và blocker."),
    (date(2026, 7, 23), "Sprint Review", "Chốt kết quả đã xác minh và chuyển phần chưa ổn định sang Sprint 2."),
]


def build_daily_1() -> None:
    configure_legacy(TEAM)
    doc = legacy.academic_document("DAILY SCRUM SPRINT 1 | TASKSYNCENTERPRISE")
    cover(doc, "Mô hình Agile", "Daily Scrum Sprint 1", "Biên bản giữa kỳ - 02/07/2026 đến 23/07/2026")
    base.add_note(
        doc,
        "Phạm vi biên bản",
        "Các buổi chính được tái dựng từ thông tin Product Owner, Jira và lịch sử Git. Nhóm họp trực tiếp nên không có ảnh/video. "
        "Những ngày không có cuộc họp chính thức được gọi là cập nhật tiến độ, không khẳng định toàn bộ nhóm đã tham dự.",
        "DDEBF7",
    )
    doc.add_heading("Mục tiêu Sprint", level=1)
    base.add_body(
        doc,
        "Tạo Increment có thể demo Login/Auth, Project, Task/Kanban, Employee sơ khai, Notification và Calendar; "
        "đồng thời ổn định nền tảng kỹ thuật.",
    )
    base.add_table(doc, ["Mã", "User Story", "SP"], [(x["id"], x["title"], x["sp"]) for x in SPRINT_1], [700, 7260, 1400], font_size=8.0)
    for index, (meeting_date, meeting_type, focus) in enumerate(DAILY_EVENTS, 1):
        doc.add_page_break()
        doc.add_heading(f"{index}. {meeting_type} - {meeting_date.strftime('%d/%m/%Y')}", level=1)
        duration = "30-60 phút" if meeting_type == "Sprint Review" else "15-30 phút"
        base.add_table(
            doc,
            ["Thông tin", "Nội dung"],
            [
                ("Hình thức", "Trực tiếp tại trường hoặc cập nhật ngắn theo tiến độ"),
                ("Khung giờ thường dùng", "12:00-13:00"),
                ("Thời lượng", duration),
                ("Trọng tâm", focus),
                ("Nguồn tái dựng", "Thông tin Product Owner, Jira và lịch sử Git"),
            ],
            [2400, 6960],
            font_size=9.0,
        )
        nhan_done = {
            1: "Khởi tạo phạm vi Sprint, rà cấu trúc Full-stack và danh sách lỗi.",
            2: "Rà Backend/Frontend và luồng đăng nhập.",
            3: "Chuẩn hóa database/model và nền tảng Backend.",
            4: "Hoàn thiện health, logging, cache và Notification nền tảng.",
            5: "Rà giao diện và các route chính phục vụ demo.",
            6: "Hoàn thiện hạ tầng Docker, Nginx và sao lưu.",
            7: "Tích hợp Project, Task, Notification và Calendar.",
            8: "Sửa lỗi chấp nhận cuối, hoàn thiện CI và tài liệu kỹ thuật.",
        }[index]
        nhan_next = {
            1: "Ưu tiên Login/Auth, database và cấu hình.",
            2: "Sửa lỗi 404, phiên đăng nhập và cấu hình.",
            3: "Củng cố Project, Task, cache và Notification.",
            4: "Kiểm thử tích hợp, giao diện và monitoring.",
            5: "Ổn định các chức năng demo và sửa lỗi.",
            6: "Chốt lỗi nghiêm trọng trước Sprint Review.",
            7: "Chuẩn bị demo và phân loại phần chuyển Sprint 2.",
            8: "Cập nhật Product Backlog và bắt đầu kế hoạch Sprint 2.",
        }[index]
        rows = [
            ("Huỳnh Lê Thành Nhân", nhan_done, nhan_next, "Khối lượng lớn; nhiều lỗi tích hợp; phản hồi nhóm hạn chế."),
        ]
        if index in {2, 4, 5, 6, 7, 8}:
            rows.append(
                (
                    "Nguyễn Đức Mạnh",
                    "Hỗ trợ kiểm thử, Calendar/Notification hoặc trao đổi lỗi 404 theo từng thời điểm.",
                    "Tiếp tục kiểm thử luồng được giao và phản hồi lỗi.",
                    "Hạn chế thời gian do môn học và công việc.",
                )
            )
        if index == 5:
            rows.append(
                (
                    "Nguyễn Lê Huy Hoàng",
                    "Chạy thử local và hỗ trợ demo cơ bản.",
                    "Phản hồi lỗi nhìn thấy trên giao diện.",
                    "Chưa có biên bản kiểm thử chi tiết.",
                )
            )
        rows.extend(
            [
                ("Phạm Tuấn Anh", "Không có nội dung cập nhật được cung cấp.", "Tiếp tục hạng mục tự chọn trên Jira.", "Chưa có sản phẩm bàn giao để xác minh."),
                ("Nguyễn Anh Tuấn", "Không có nội dung cập nhật được cung cấp.", "Nghiên cứu tài liệu theo phân công.", "Chưa có sản phẩm bàn giao để xác minh."),
            ]
        )
        base.add_table(doc, ["Thành viên", "Đã làm", "Sẽ làm", "Trở ngại"], rows, [1900, 2600, 2500, 2360], font_size=7.8)
        doc.add_heading("Quyết định sau cập nhật", level=2)
        base.add_bullets(
            doc,
            [
                "Chỉ chuyển trạng thái hoàn thành khi có mã nguồn, kiểm thử hoặc demo.",
                "Product Owner tiếp tục xử lý blocker kỹ thuật và thu hẹp phạm vi demo khi cần.",
                "Công việc được giao nhưng thiếu sản phẩm bàn giao giữ trạng thái chưa xác minh.",
            ],
        )
        figure_caption(doc, "", f"Minh chứng cập nhật tiến độ ngày {meeting_date.strftime('%d/%m/%Y')}")
    base.save_doc(doc, AGILE_DIR / "Daily_Scrum_1_TaskSyncEnterprise.docx")


def business_days(start: date, end: date) -> list[date]:
    days = []
    current = start
    while current <= end:
        if current.weekday() < 5:
            days.append(current)
        current += timedelta(days=1)
    return days


def build_daily_2() -> None:
    configure_legacy(TEAM)
    doc = legacy.academic_document("DAILY SCRUM SPRINT 2 | TASKSYNCENTERPRISE")
    cover(doc, "Mô hình Agile", "Daily Scrum Sprint 2", "Kế hoạch dự kiến - 24/07/2026 đến 20/08/2026")
    base.add_note(
        doc,
        "Trạng thái",
        "Toàn bộ nội dung trong tài liệu này là kế hoạch Sprint 2. Chỉ chuyển sang kết quả thực tế sau khi nhóm cập nhật Jira, mã nguồn, kiểm thử hoặc demo.",
        "E7E6E6",
    )
    doc.add_heading("Mục tiêu Sprint 2", level=1)
    base.add_body(
        doc,
        "Ổn định Department/Team, tiếp tục kiểm thử các chức năng sơ khai, hoàn thiện nghỉ phép, làm thêm giờ, báo cáo, audit "
        "và triển khai năm phase kỹ thuật dọn dẹp - hạ tầng - database - xác thực - phân quyền.",
    )
    base.add_table(doc, ["Mã", "User Story", "SP"], [(x["id"], x["title"], x["sp"]) for x in SPRINT_2], [700, 7260, 1400], font_size=8.0)
    days = business_days(date(2026, 7, 24), date(2026, 8, 20))
    for index, meeting_date in enumerate(days, 1):
        item = SPRINT_2[(index - 1) % len(SPRINT_2)]
        doc.add_page_break()
        doc.add_heading(f"Kế hoạch Daily Scrum {index:02d} - {meeting_date.strftime('%d/%m/%Y')}", level=1)
        base.add_table(
            doc,
            ["Thông tin", "Nội dung"],
            [
                ("Trạng thái", "Dự kiến - cập nhật sau khi thực hiện"),
                ("Thời lượng", "15 phút"),
                ("Trọng tâm", f"{item['id']} - {item['title']}"),
                ("Mục tiêu", "Kiểm tra tiến độ, blocker và bằng chứng trong 24 giờ tiếp theo."),
            ],
            [2400, 6960],
            font_size=9.0,
        )
        rows = [
            ("Huỳnh Lê Thành Nhân", "Rà yêu cầu, triển khai và tích hợp chức năng chính.", "Dự kiến"),
            ("Nguyễn Đức Mạnh", "Hỗ trợ kiểm thử, Calendar/Notification và nội dung Login/Auth/404.", "Dự kiến"),
            ("Nguyễn Lê Huy Hoàng", "Chạy test demo trên local và ghi lỗi cơ bản.", "Dự kiến"),
            ("Phạm Tuấn Anh", "Cập nhật sơ đồ theo hạng mục được giao.", "Dự kiến - cần sản phẩm bàn giao"),
            ("Nguyễn Anh Tuấn", "Nghiên cứu và tổng hợp tài liệu.", "Dự kiến - cần sản phẩm bàn giao"),
        ]
        base.add_table(doc, ["Thành viên", "Kế hoạch", "Trạng thái"], rows, [2400, 4860, 2100], font_size=8.5)
        doc.add_heading("Ba câu hỏi Daily Scrum", level=2)
        base.add_bullets(
            doc,
            [
                "Từ lần cập nhật trước đã hoàn thành phần nào?",
                "Trong 24 giờ tới sẽ hoàn thành phần nào?",
                "Trở ngại nào đang ảnh hưởng Sprint Goal và ai chịu trách nhiệm xử lý?",
            ],
        )
        figure_caption(doc, "", f"Minh chứng Daily Scrum Sprint 2 ngày {meeting_date.strftime('%d/%m/%Y')}")
    base.save_doc(doc, AGILE_DIR / "Daily_Scrum_2_TaskSyncEnterprise.docx")


def build_contribution_report() -> None:
    configure_legacy(TEAM)
    doc = legacy.academic_document("ĐÓNG GÓP THÀNH VIÊN SPRINT 1 | TASKSYNCENTERPRISE")
    cover(doc, "Mô hình Agile", "Đóng góp thành viên Sprint 1", "Báo cáo xác minh trung tính - không sử dụng tỷ lệ phần trăm")
    doc.add_heading("1. Mục đích", level=1)
    base.add_body(
        doc,
        "Tài liệu phân biệt vai trò được giao, trạng thái Jira và kết quả có thể xác minh. "
        "Không sử dụng từ ngữ phê bình cá nhân và không kết luận chỉ dựa vào tên người được giao trên Jira.",
    )
    doc.add_heading("2. Nguồn xác minh", level=1)
    base.add_bullets(
        doc,
        [
            "Jira: 50 issue, gồm 14 Done, 5 In Progress và 31 To Do tại thời điểm rà soát.",
            "Git: 79 commit trong giai đoạn 04/07-23/07, đều thuộc các biến thể tên Huỳnh Lê Thành Nhân.",
            "Mã nguồn: Frontend, Backend, database, kiểm thử và tài liệu kỹ thuật.",
            "Thông tin Product Owner: hoạt động trực tiếp, hỗ trợ kiểm thử và tình trạng bàn giao.",
        ],
    )
    summary = [
        ("Huỳnh Lê Thành Nhân", "Product Owner / Full-stack", "Đóng góp chính", "79 commit; triển khai Full-stack; tích hợp; debug; CI và tài liệu kỹ thuật."),
        ("Nguyễn Đức Mạnh", "Scrum Master tạm thời / Hỗ trợ", "Có đóng góp và hỗ trợ", "Kiểm thử Login/Auth; hỗ trợ lỗi 404; Calendar; hỗ trợ Notification; viết phần Login/Auth và 404."),
        ("Nguyễn Lê Huy Hoàng", "Hỗ trợ kiểm thử/demo", "Hỗ trợ kiểm thử", "Chạy local và test demo cơ bản; cần xác minh từng hạng mục Jira."),
        ("Phạm Tuấn Anh", "Sơ đồ/tài liệu được phân công", "Chưa có sản phẩm bàn giao để xác minh", "Có hạng mục tự chọn trên Jira; chưa có minh chứng bàn giao được Product Owner xác nhận."),
        ("Nguyễn Anh Tuấn", "Nghiên cứu tài liệu được phân công", "Chưa có sản phẩm bàn giao để xác minh", "Chưa có tài liệu tổng hợp được bàn giao để xác minh."),
    ]
    base.add_table(doc, ["Thành viên", "Vai trò", "Mức đóng góp", "Cơ sở nhận định"], summary, [1900, 2400, 2100, 2960], font_size=8.0)

    profiles = [
        (
            "Huỳnh Lê Thành Nhân",
            [
                ("Vai trò", "Product Owner / Full-stack Developer"),
                ("Công việc", "Phân tích, kiến trúc, database, Backend, Frontend, tích hợp, debug, CI và tài liệu kỹ thuật."),
                ("Kết quả", "Sản phẩm chạy gần đầy đủ; phạm vi giữa kỳ tập trung chức năng có thể demo và phần sơ khai được ghi rõ."),
                ("Bằng chứng", "79 commit Sprint 1, Jira, mã nguồn và giao diện."),
                ("Trở ngại", "Khối lượng lớn, lỗi tích hợp và phản hồi nhóm hạn chế."),
            ],
        ),
        (
            "Nguyễn Đức Mạnh",
            [
                ("Vai trò", "Scrum Master tạm thời / Hỗ trợ kiểm thử"),
                ("Công việc", "Kiểm thử Login/Auth; hỗ trợ debug lỗi 404; làm Calendar; hỗ trợ kiểm thử Notification."),
                ("Báo cáo", "Phụ trách nội dung Login/Auth và lỗi 404."),
                ("Kết quả", "Có đóng góp và hỗ trợ; không ghi là người viết chính toàn bộ Login/Auth."),
                ("Trở ngại", "Hạn chế thời gian do môn học và công việc."),
            ],
        ),
        (
            "Nguyễn Lê Huy Hoàng",
            [
                ("Vai trò", "Hỗ trợ kiểm thử và demo"),
                ("Công việc", "Chạy thử local, test demo và phản hồi lỗi cơ bản."),
                ("Jira", "Tài khoản MSSV 2311554285 có thể thuộc thành viên này; chưa dùng làm kết luận chắc chắn."),
                ("Kết quả", "Ghi nhận mức hỗ trợ kiểm thử; cần bổ sung biên bản hoặc kết quả test cho từng Story."),
            ],
        ),
        (
            "Phạm Tuấn Anh",
            [
                ("Vai trò", "Sơ đồ và tài liệu được phân công"),
                ("Jira", "Có hạng mục tự chọn, gồm nhập nhân viên, trưởng phòng và tạo dự án."),
                ("Kết quả", "Chưa có sản phẩm bàn giao để Product Owner xác minh tại thời điểm rà soát."),
                ("Cách cập nhật", "Khi có sơ đồ/tài liệu, bổ sung đường dẫn và chuyển trạng thái sau khi review."),
            ],
        ),
        (
            "Nguyễn Anh Tuấn",
            [
                ("Vai trò", "Nghiên cứu tài liệu được phân công"),
                ("Công việc", "Đọc tài liệu và nghiên cứu nội dung báo cáo theo phân công."),
                ("Kết quả", "Chưa có sản phẩm bàn giao để xác minh tại thời điểm rà soát."),
                ("Cách cập nhật", "Bổ sung tài liệu tổng hợp, nguồn tham khảo và phần nội dung được chấp nhận."),
            ],
        ),
    ]
    for index, (name, rows) in enumerate(profiles, 1):
        doc.add_page_break()
        doc.add_heading(f"{index}. {name}", level=1)
        base.add_table(doc, ["Hạng mục", "Nội dung"], rows, [2200, 7160], font_size=9.0)
        doc.add_heading("Kết luận xác minh", level=2)
        base.add_body(
            doc,
            next(row[3] for row in summary if row[0] == name),
        )
        figure_caption(doc, "", f"Minh chứng đóng góp của {name}")
    base.save_doc(doc, AGILE_DIR / "Bao_cao_dong_gop_thanh_vien_Sprint_1_TaskSyncEnterprise.docx")


def build_full_reports() -> None:
    configure_legacy(TEAM)
    legacy.build_agile_report()
    configure_legacy(ARCH_AUTHOR)
    legacy.build_architecture_report()
    legacy.build_diagram_guide()


def build_readme() -> None:
    text = """# Bộ hồ sơ Report - TaskSyncEnterprise

## Quy ước chính

- Sprint 1 = Sprint cũ 1 + 2, từ 02/07/2026 đến 23/07/2026, là phạm vi giữa kỳ.
- Sprint 2 = Sprint cũ 3 + 4, từ 24/07/2026 đến 20/08/2026, là kế hoạch/tiếp tục xác minh.
- Product Backlog gồm 50 User Story: 38 từ Jira và 12 bổ sung từ mã nguồn/Git.
- Trạng thái Jira và trạng thái đã xác minh được tách riêng.

## Hồ sơ Agile

- `Bao_cao_Giua_ky_Sprint_1_TaskSyncEnterprise.docx`: bản nộp giữa kỳ rút gọn.
- `Bao_cao_Agile_Scrum_TaskSyncEnterprise.docx`: báo cáo Agile đầy đủ.
- `Daily_Scrum_1_TaskSyncEnterprise.docx`: biên bản Sprint 1.
- `Daily_Scrum_2_TaskSyncEnterprise.docx`: kế hoạch dự kiến Sprint 2.
- `Bao_cao_dong_gop_thanh_vien_Sprint_1_TaskSyncEnterprise.docx`: xác minh đóng góp trung tính.
- `Agile_Project_Management_TaskSyncEnterprise.xlsx`: Product Backlog, Sprint và Epic.
- `Agile_Tracking_TaskSyncEnterprise.xlsx`: công việc, Daily, lỗi, rủi ro và bằng chứng.

## Hồ sơ Kiến trúc

- `Môn kiến trúc phần mềm/Bao_cao_Kien_truc_phan_mem_TaskSyncEnterprise.docx`: báo cáo kiến trúc do Huỳnh Lê Thành Nhân thực hiện.
- `Sơ đồ Diagram/`: sơ đồ PNG, Mermaid và tài liệu thuyết minh.

## Nội dung cần nhóm tự cập nhật

1. Chèn ảnh vào các vị trí có chú thích `Hình X`.
2. Bổ sung MSSV Phạm Tuấn Anh và Nguyễn Anh Tuấn.
3. Điền lớp học phần.
4. Cập nhật mục lục Word bằng `Ctrl+A`, sau đó nhấn `F9`.
5. Chỉ đổi trạng thái khi có mã nguồn, kiểm thử, demo hoặc sản phẩm bàn giao.
"""
    (REPORT / "README_Bo_ho_so.md").write_text(text, encoding="utf-8")


def main() -> None:
    legacy.build_extended_diagrams()
    build_midterm_report()
    build_daily_1()
    build_daily_2()
    build_contribution_report()
    build_full_reports()
    build_readme()
    print("Built report package version 3.")


if __name__ == "__main__":
    main()
