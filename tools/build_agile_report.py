from __future__ import annotations

import os
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"E:\TaskSyncEnterprise")
SOURCE = Path(r"C:\Users\huynh\Downloads\Docs_Midel (2).docx")
OUT_DIR = ROOT / "outputs" / "agile_report_final"
OUT_DOCX = OUT_DIR / "TaskSyncEnterprise_BaoCao_Agile_Final.docx"
ASSET_DIR = OUT_DIR / "assets"
SCREEN_DIR = ROOT / "docs" / "testing" / "screenshots" / "codex"

NAVY = "1F4E78"
TEAL = "0F6B78"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FA"
PALE_GREEN = "E2F0D9"
PALE_AMBER = "FFF2CC"
PALE_GRAY = "F2F2F2"
DARK = RGBColor(31, 78, 120)
GRAY = RGBColor(89, 89, 89)


SPRINT1 = [
    ("JD-14", "US01", "Đăng nhập người dùng an toàn", 3, "Huỳnh Lê Thành Nhân", "Jira Sprint 1"),
    ("JD-15", "US02", "Tự động gia hạn phiên đăng nhập", 5, "Nguyễn Đức Mạnh", "Jira Sprint 1"),
    ("JD-16", "US03", "Thiết lập danh mục vai trò", 5, "Nguyễn Lê Huy Hoàng", "Jira Sprint 1"),
    ("JD-17", "US04", "Kiểm soát truy cập theo vai trò", 8, "Nguyễn Lê Huy Hoàng", "Jira Sprint 1"),
    ("JD-24", "US11", "Tạo dự án", 3, "Phạm Tuấn Anh", "Jira Sprint 1"),
    ("JD-25", "US12", "Phân công thành viên vào dự án", 5, "Huỳnh Lê Thành Nhân", "Jira Sprint 1"),
    ("JD-28", "US15", "Theo dõi công việc trên bảng Kanban", 5, "Huỳnh Lê Thành Nhân", "Jira Sprint 1"),
    ("JD-35", "US22", "Tải nhiều tài liệu cùng lúc", 5, "Huỳnh Lê Thành Nhân", "Jira Sprint 1"),
    ("JD-18", "US05", "Tạo hồ sơ nhân viên", 5, "Nguyễn Đức Mạnh", "Jira Sprint 2"),
    ("JD-19", "US06", "Chỉnh sửa thông tin tài khoản", 8, "Phạm Tuấn Anh", "Jira Sprint 2"),
    ("JD-20", "US07", "Gán người quản lý trực tiếp", 3, "Nguyễn Đức Mạnh", "Jira Sprint 2"),
    ("JD-21", "US08", "Thiết lập cơ cấu phòng ban", 5, "Nguyễn Đức Mạnh", "Jira Sprint 2"),
    ("JD-22", "US09", "Bổ nhiệm trưởng phòng", 5, "Phạm Tuấn Anh", "Jira Sprint 2"),
    ("JD-23", "US10", "Quản lý chuyển phòng ban", 13, "Huỳnh Lê Thành Nhân", "Jira Sprint 2"),
    ("JD-29", "US16", "Theo dõi quan hệ phụ thuộc công việc", 5, "Huỳnh Lê Thành Nhân", "Jira Sprint 2"),
    ("JD-31", "US18", "Phân rã yêu cầu thành công việc nhỏ", 8, "Phạm Anh Tuấn", "Jira Sprint 2"),
    ("JD-32", "US19", "Trao đổi bằng luồng bình luận", 8, "Nguyễn Lê Huy Hoàng", "Jira Sprint 2"),
]

SPRINT2_JIRA3 = [
    ("JD-60", "US41", "Quản lý phiên đăng nhập", 5, "Unassigned"),
    ("JD-59", "US40", "Xử lý phiên hết hạn", 5, "Unassigned"),
    ("JD-58", "US39", "Đăng xuất an toàn", 5, "Unassigned"),
    ("JD-62", "US43", "Xem lịch công việc", 5, "Unassigned"),
    ("JD-27", "US14", "Tạo và chỉnh sửa công việc", 5, "Unassigned"),
    ("JD-54", "US17", "Lọc danh sách yêu cầu nâng cao", 5, "Nguyễn Lê Huy Hoàng"),
    ("JD-37", "US24", "Gửi yêu cầu nghỉ phép", 3, "Unassigned"),
    ("JD-38", "US25", "Phê duyệt yêu cầu nghỉ phép", 5, "Unassigned"),
    ("JD-63", "US45", "Xem trung tâm thông báo", 5, "Nguyễn Lê Huy Hoàng"),
    ("JD-44", "US31", "Nhận cảnh báo thời gian thực", 8, "Unassigned"),
    ("JD-64", "US46", "Đánh dấu thông báo đã đọc", 5, "Unassigned"),
    ("JD-61", "US42", "Xem bảng điều khiển tổng quan", 3, "Unassigned"),
    ("JD-55", "US20", "Thông báo khi người dùng được nhắc tên", None, "Unassigned"),
    ("JD-50", "US37", "Tạo báo cáo dự án tùy chỉnh", None, "Unassigned"),
    ("JD-51", "US38", "Dự báo tiến độ hoàn thành dự án", None, "Unassigned"),
    ("JD-52", "US47", "Kiểm tra tình trạng hệ thống", None, "Unassigned"),
]

SPRINT2_JIRA4 = [
    ("JD-42", "US29", "Phê duyệt yêu cầu làm thêm giờ"),
    ("JD-43", "US30", "Kiểm tra quy định làm thêm giờ"),
    ("JD-41", "US28", "Gửi yêu cầu làm thêm giờ"),
    ("JD-40", "US27", "Hủy yêu cầu nghỉ khẩn cấp"),
    ("JD-39", "US26", "Tính số dư phép năm"),
    ("JD-26", "US13", "Theo dõi cột mốc dự án"),
    ("JD-34", "US21", "Bảo vệ tệp đính kèm công việc"),
    ("JD-45", "US32", "Cài đặt kênh nhận thông báo"),
    ("JD-46", "US33", "Nhận email tổng hợp hằng ngày"),
    ("JD-47", "US34", "Gửi lại thông báo thất bại"),
    ("JD-48", "US35", "Theo dõi bảng điều khiển hiệu suất"),
    ("JD-49", "US36", "Theo dõi phân bổ nguồn lực phòng ban"),
    ("JD-53", "US48", "Ghi nhật ký có cấu trúc"),
    ("JD-57", "US50", "Triển khai bằng Docker và kiểm tra CI"),
    ("JD-56", "US49", "Sử dụng bộ nhớ đệm Redis"),
]


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row):
    repeat_header(row)
    prevent_row_split(row)


def set_run_font(run, size=12, bold=False, color=None, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, 10)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Mục lục sẽ được cập nhật khi mở bằng Microsoft Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, placeholder, end])


def add_para(doc, text="", *, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=6, first_line=True, color=None, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    if first_line and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Cm(1)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.25
        set_run_font(p.add_run(item), 12)


def add_heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 14, 3: 12}[level], bold=True, color=DARK)
    return p


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 110, 140, 110, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(title + ": "), 11, bold=True, color=DARK)
    set_run_font(p.add_run(text), 11)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths=None, font_size=9, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B7C9D6")
        borders.append(tag)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j, header in enumerate(headers):
        cell = hdr.cells[j]
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(str(header)), font_size, bold=True, color=RGBColor(255, 255, 255))
        set_cell_margins(cell)
        if widths:
            cell.width = Cm(widths[j])
    for i, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if i % 2:
            for cell in row.cells:
                shade(cell, PALE_GRAY)
        for j, value in enumerate(row_data):
            cell = row.cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = alignments[j] if alignments else (WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            set_run_font(p.add_run("" if value is None else str(value)), font_size)
            if widths:
                cell.width = Cm(widths[j])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run(text), 10, italic=True, color=GRAY)
    return p


def add_figure(doc, image_path: Path, caption: str, width=6.2):
    if not image_path.exists():
        add_callout(doc, "Thiếu hình", f"Không tìm thấy {image_path.name}", PALE_AMBER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = DARK
        if style_name == "Heading 1":
            style.paragraph_format.page_break_before = False
    for style_name in ("List Bullet", "List Number"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)


def set_update_fields(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def extract_logo():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    target = ASSET_DIR / "source-logo.jpg"
    try:
        with zipfile.ZipFile(SOURCE) as zf:
            data = zf.read("word/media/image1.jpg")
        target.write_bytes(data)
        return target
    except Exception:
        return None


def add_cover(doc, logo):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HỒ CHÍ MINH"), 14, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("KHOA CÔNG NGHỆ THÔNG TIN"), 13, bold=True, color=DARK)
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(1.25))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("BÁO CÁO ĐỒ ÁN MÔN HỌC"), 18, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("PHÁT TRIỂN PHẦN MỀM THEO AGILE"), 18, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    set_run_font(p.add_run("TASKSYNCENTERPRISE"), 22, bold=True, color=RGBColor(15, 107, 120))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Hệ thống quản lý công việc và cộng tác nội bộ doanh nghiệp"), 13, italic=True, color=GRAY)
    doc.add_paragraph()
    rows = [
        ("2000004897", "Huỳnh Lê Thành Nhân", "Product Owner / Thành viên phát triển"),
        ("2200010420", "Nguyễn Đức Mạnh", "Scrum Master / Thành viên phát triển"),
        ("2311558672", "Phạm Tuấn Anh", "Thành viên phát triển"),
        ("2311559121", "Phạm Anh Tuấn", "Thành viên phát triển"),
        ("2311554285", "Nguyễn Lê Huy Hoàng", "Thành viên phát triển"),
    ]
    add_table(doc, ["MSSV", "Họ và tên", "Vai trò trong nhóm"], rows, widths=[3.0, 5.8, 7.5], font_size=10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    set_run_font(p.add_run("TP. Hồ Chí Minh, tháng 08 năm 2026"), 12, italic=True)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(SOURCE))
    clear_body(doc)
    configure_styles(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "TASKSYNCENTERPRISE — BÁO CÁO ĐỒ ÁN AGILE"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(hp.runs[0], 9, color=GRAY)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    add_page_number(fp)
    set_update_fields(doc)

    logo = extract_logo()
    add_cover(doc, logo)

    add_heading(doc, "LỜI CAM ĐOAN VÀ PHẠM VI XÁC MINH", 1, page_break=True)
    add_para(doc, "Nhóm cam đoan nội dung báo cáo được xây dựng từ mã nguồn TaskSyncEnterprise, bảng công việc Jira, bảng tính Agile và kết quả kiểm thử tại thời điểm ngày 10/08/2026. Báo cáo phân biệt rõ ba mức: chức năng đã được Jira xác nhận Done; nền tảng kỹ thuật đã tồn tại trong mã nguồn nhưng chưa được Jira nghiệm thu; và hạng mục mở rộng chỉ nằm trong kế hoạch.")
    add_callout(doc, "Quy ước quan trọng", "Sprint 1 của báo cáo = Jira Sprint 1 + Jira Sprint 2. Sprint 2 của báo cáo = Jira Sprint 3 + Jira Sprint 4. Trạng thái Jira là nguồn quản lý phạm vi; mã nguồn, kiểm thử và ảnh chụp là bằng chứng kỹ thuật bổ sung.")
    add_para(doc, "Bản báo cáo này sửa lại các số liệu chưa nhất quán trong bản nháp. Cụ thể, tổng chính thức của Sprint 1 là 17 issue và 99 Story Point, không sử dụng tổng 135 Story Point của bảng cũ. Sprint 2 có 31 issue; tại ngày đối chiếu chỉ 59 Story Point đã được ước lượng và chưa có Story Point nào ở trạng thái Done.")

    add_heading(doc, "MỤC LỤC", 1, page_break=True)
    p = doc.add_paragraph()
    add_toc_field(p)
    add_para(doc, "Ghi chú: Microsoft Word sẽ tự cập nhật số trang của mục lục khi mở tài liệu.", italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, size=10, color=GRAY)

    add_heading(doc, "DANH MỤC BẢNG VÀ HÌNH ẢNH", 1, page_break=True)
    rows = [
        ("Bảng 3.1", "Các lớp kiến trúc của TaskSyncEnterprise"),
        ("Bảng 3.2", "Các thực thể dữ liệu cốt lõi"),
        ("Bảng 4.1", "Phạm vi Sprint 1 — 17 issue, 99 SP"),
        ("Bảng 5.1", "Phạm vi Jira Sprint 3 trong Sprint 2 báo cáo"),
        ("Bảng 5.2", "Phạm vi Jira Sprint 4 — hạng mục mở rộng"),
        ("Bảng 6.1", "Kết quả kiểm thử và build"),
        ("Hình 6.1–6.9", "Các màn hình chức năng đã chạy thử trong mã nguồn"),
    ]
    add_table(doc, ["Ký hiệu", "Nội dung"], rows, widths=[3.2, 13.0], font_size=10)

    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1, page_break=True)
    add_heading(doc, "1.1. Bối cảnh và bài toán", 2)
    add_para(doc, "Trong doanh nghiệp, dữ liệu nhân sự, phòng ban, dự án, Sprint, backlog, công việc và thông báo thường bị phân tán ở nhiều công cụ. Việc thiếu một không gian làm việc thống nhất làm tăng thời gian cập nhật, khó truy vết trách nhiệm và giảm khả năng quan sát tiến độ. TaskSyncEnterprise được xây dựng như một nền tảng quản lý công việc nội bộ theo định hướng Agile, hỗ trợ quản trị cơ cấu tổ chức và thực thi công việc trên cùng hệ thống.")
    add_heading(doc, "1.2. Mục tiêu", 2)
    add_bullets(doc, [
        "Xây dựng cơ chế đăng nhập, quản lý phiên và phân quyền theo vai trò.",
        "Quản lý nhân viên, phòng ban, đội nhóm, dự án, Sprint, Epic, Backlog Item và Task.",
        "Hỗ trợ theo dõi công việc bằng danh sách/Kanban, bình luận, tệp đính kèm và thông báo.",
        "Tổ chức quá trình phát triển theo Scrum, quản lý phạm vi trên Jira và minh chứng bằng mã nguồn, kiểm thử và màn hình chạy thử.",
    ])
    add_heading(doc, "1.3. Phạm vi báo cáo", 2)
    add_para(doc, "Báo cáo tập trung vào hai Sprint học phần. Sprint 1 báo cáo tổng hợp hai Sprint đầu trên Jira và là phần đã hoàn thành. Sprint 2 báo cáo tổng hợp Jira Sprint 3 và Sprint 4, nhưng tại thời điểm chốt báo cáo vẫn là kế hoạch. Các chức năng mở rộng như phép năm, làm thêm giờ, báo cáo nâng cao, Redis và Docker được mô tả theo đúng trạng thái, không ghi nhận là hoàn thành nếu Jira chưa Done.")
    add_heading(doc, "1.4. Nguồn dữ liệu dùng để đánh giá", 2)
    for label, url in [
        ("Mã nguồn nhánh develop", "https://github.com/huynhlethanhnhan/TaskSyncEnterprise/tree/develop"),
        ("Jira Backlog dự án JD", "https://task-snycs-enterprise.atlassian.net/jira/software/projects/JD/boards/1/backlog"),
        ("Bảng tính Agile đã đối chiếu", "https://docs.google.com/spreadsheets/d/1pBk8gqkxyRddmChIlGe78xFSgaf-65pRsKoh2rqwkx4/edit"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(label + ": "), 11, bold=True)
        add_hyperlink(p, url, url)

    add_heading(doc, "CHƯƠNG 2. PHƯƠNG PHÁP AGILE/SCRUM ÁP DỤNG", 1, page_break=True)
    add_heading(doc, "2.1. Vai trò Scrum", 2)
    rows = [
        ("Product Owner", "Huỳnh Lê Thành Nhân", "Sắp xếp Product Backlog, làm rõ Acceptance Criteria, xác nhận ưu tiên và phạm vi báo cáo."),
        ("Scrum Master", "Nguyễn Đức Mạnh", "Điều phối Sprint, hỗ trợ loại bỏ trở ngại, nhắc cập nhật Jira và bằng chứng."),
        ("Development Team", "Toàn bộ 5 thành viên", "Phân tích, phát triển, kiểm thử, sửa lỗi, tài liệu hóa và demo Increment."),
    ]
    add_table(doc, ["Vai trò", "Phụ trách", "Trách nhiệm"], rows, widths=[3.1, 4.2, 9.0], font_size=10)
    add_heading(doc, "2.2. Các sự kiện Scrum", 2)
    add_bullets(doc, [
        "Sprint Planning: thống nhất Sprint Goal, chọn backlog theo năng lực và xác định tiêu chí hoàn thành.",
        "Daily Scrum: cập nhật việc đã làm, việc sẽ làm, trở ngại và bằng chứng liên quan.",
        "Sprint Review: trình diễn Increment, so sánh với Acceptance Criteria và ghi nhận phản hồi.",
        "Sprint Retrospective: xem lại cách cộng tác, chất lượng kỹ thuật và hành động cải tiến cho Sprint sau.",
    ])
    add_heading(doc, "2.3. Definition of Ready và Definition of Done", 2)
    rows = [
        ("Definition of Ready", "User Story có mục tiêu rõ, Acceptance Criteria kiểm thử được, có phụ thuộc/owner, có ước lượng hoặc ghi rõ chưa ước lượng."),
        ("Definition of Done", "Mã nguồn đã tích hợp; kiểm thử liên quan đạt; build thành công; phân quyền và lỗi chính được kiểm tra; có bằng chứng demo; Jira chuyển Done sau khi Product Owner xác nhận."),
    ]
    add_table(doc, ["Chuẩn", "Điều kiện"], rows, widths=[4.0, 12.3], font_size=10)
    add_heading(doc, "2.4. Quy tắc gộp Sprint trong báo cáo", 2)
    rows = [
        ("Sprint 1 báo cáo", "Jira Sprint 1 + Jira Sprint 2", "02/07–30/07/2026", "17 issue", "99", "99"),
        ("Sprint 2 báo cáo", "Jira Sprint 3 + Jira Sprint 4", "30/07–28/08/2026", "31 issue", "59 đã ước lượng", "0"),
    ]
    add_table(doc, ["Sprint", "Nguồn Jira", "Khoảng thời gian", "Phạm vi", "Tổng SP", "SP Done"], rows, widths=[3.0, 4.2, 3.5, 2.3, 2.3, 2.0], font_size=9)

    add_heading(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1, page_break=True)
    add_heading(doc, "3.1. Kiến trúc tổng thể", 2)
    add_para(doc, "TaskSyncEnterprise sử dụng kiến trúc modular monolith. Giao diện React gọi REST API FastAPI; các cập nhật thời gian thực sử dụng WebSocket; dữ liệu nghiệp vụ lưu trên SQL Server; Redis được thiết kế cho cache và hỗ trợ hạ tầng. Nginx/Docker đóng vai trò reverse proxy và đóng gói triển khai. Cách tổ chức theo module giúp nhóm phát triển nhanh trong phạm vi môn học nhưng vẫn giữ ranh giới nghiệp vụ rõ ràng.")
    rows = [
        ("Presentation", "React 19, TypeScript, Vite, Tailwind, TanStack Query", "Trang đăng nhập, dashboard, danh mục, Kanban và form thao tác."),
        ("API/Application", "FastAPI, Pydantic, REST, WebSocket", "Xác thực request, điều phối use case, chuẩn hóa response và sự kiện thời gian thực."),
        ("Domain/Service", "Auth, RBAC, Employee, Department, Team, Project, Sprint, Backlog, Task, Notification", "Đóng gói luật nghiệp vụ và quyền truy cập."),
        ("Persistence", "SQLAlchemy, Alembic, SQL Server 2022", "Lưu dữ liệu, quan hệ và migration."),
        ("Infrastructure", "Redis, Nginx, Docker", "Cache, reverse proxy, đóng gói và vận hành."),
    ]
    add_table(doc, ["Lớp", "Công nghệ/thành phần", "Trách nhiệm"], rows, widths=[3.0, 6.0, 7.3], font_size=9)
    add_caption(doc, "Bảng 3.1. Các lớp kiến trúc của TaskSyncEnterprise")
    add_heading(doc, "3.2. Phân rã miền nghiệp vụ", 2)
    add_bullets(doc, [
        "Identity & Access: người dùng, phiên đăng nhập, vai trò, quyền và kiểm soát truy cập.",
        "Organization Directory: nhân viên, phòng ban, đội nhóm và quan hệ quản lý.",
        "Project Delivery: dự án, thành viên dự án, Sprint, Epic, Backlog Item, Task và phụ thuộc.",
        "Collaboration: bình luận, tệp đính kèm, thông báo và cập nhật thời gian thực.",
        "Operational Support: audit log, kiểm tra health, cache và cấu hình triển khai.",
    ])
    add_heading(doc, "3.3. Thiết kế cơ sở dữ liệu", 2)
    rows = [
        ("employees", "Hồ sơ nhân viên và thông tin tài khoản", "roles, departments, teams, tasks"),
        ("roles", "Danh mục vai trò và nền tảng RBAC", "employees, permissions"),
        ("departments", "Cơ cấu phòng ban, trưởng phòng", "employees, teams"),
        ("teams", "Đội nhóm trực thuộc phòng ban", "employees, projects"),
        ("projects", "Thông tin dự án và thành viên", "sprints, epics, backlog_items, tasks"),
        ("sprints", "Timebox của dự án", "backlog_items, tasks"),
        ("backlog_items", "User Story/Backlog Item và ưu tiên", "epics, sprints, tasks"),
        ("tasks", "Đơn vị công việc, trạng thái, người phụ trách", "employees, projects, sprints"),
        ("notifications", "Thông báo theo người nhận và trạng thái đọc", "employees, tasks"),
        ("audit_logs", "Dấu vết hành động quan trọng", "employees và thực thể nghiệp vụ"),
        ("vacations", "Yêu cầu nghỉ phép và trạng thái phê duyệt", "employees, approvers"),
    ]
    add_table(doc, ["Thực thể", "Vai trò", "Quan hệ chính"], rows, widths=[3.2, 6.0, 7.1], font_size=9)
    add_caption(doc, "Bảng 3.2. Các thực thể dữ liệu cốt lõi được xác minh trong mã nguồn")
    add_heading(doc, "3.4. Luồng xử lý tiêu biểu", 2)
    rows = [
        ("Đăng nhập", "Người dùng nhập email/mật khẩu", "Auth API xác minh tài khoản và mật khẩu băm", "Phát hành token/phiên và chuyển đến Dashboard"),
        ("Tạo/cập nhật Task", "Người dùng gửi dữ liệu và project/sprint", "API kiểm tra quyền, validate trạng thái và ghi CSDL", "UI làm mới Kanban; sự kiện thông báo được phát nếu cần"),
        ("Phân quyền", "Người dùng gọi endpoint được bảo vệ", "Middleware/Dependency đọc danh tính và vai trò", "Cho phép hoặc trả lỗi 401/403, đồng thời có thể ghi audit"),
        ("Thông báo", "Sự kiện nghiệp vụ phát sinh", "Service tạo notification cho người nhận", "UI nhận danh sách/real-time và đánh dấu đã đọc"),
    ]
    add_table(doc, ["Luồng", "Bước 1", "Bước 2", "Bước 3"], rows, widths=[2.6, 4.5, 5.0, 4.2], font_size=8)
    add_heading(doc, "3.5. Use Case trọng tâm và tiêu chí chấp nhận", 2)
    rows = [
        ("US01", "Đăng nhập an toàn", "Đăng nhập đúng chuyển đến Dashboard; sai thông tin báo lỗi; API bảo vệ từ chối khi chưa đăng nhập."),
        ("US04", "Kiểm soát truy cập theo vai trò", "Người không đủ quyền không xem/thực hiện thao tác quản trị; quyền đúng được phép."),
        ("US11", "Tạo dự án", "Dự án lưu đủ trường bắt buộc; hiển thị trong danh sách; thành viên được phân công theo quyền."),
        ("US15", "Theo dõi bằng Kanban", "Task hiển thị đúng cột trạng thái; thay đổi trạng thái cập nhật server và giao diện."),
        ("US19", "Luồng bình luận", "Thành viên liên quan có thể thêm và xem bình luận theo Task; dữ liệu được sắp xếp theo thời gian."),
        ("US22", "Tải nhiều tài liệu", "Người dùng chọn nhiều tệp hợp lệ và nhận một gói tải xuống; quyền truy cập được kiểm tra."),
    ]
    add_table(doc, ["US", "Tên", "Acceptance Criteria rút gọn"], rows, widths=[2.0, 5.0, 9.3], font_size=9)
    add_heading(doc, "3.6. Ràng buộc chất lượng", 2)
    add_bullets(doc, [
        "Bảo mật: mật khẩu băm, token/phiên có thời hạn, RBAC phía API, không tin cậy kiểm tra quyền chỉ ở giao diện.",
        "Tính đúng đắn: migration có kiểm soát, validation request/response, quan hệ dữ liệu rõ ràng.",
        "Khả năng kiểm thử: tách service/router, có kiểm thử backend và frontend tự động.",
        "Khả năng triển khai: cấu hình theo môi trường, Docker/Nginx và health check; các phần chưa Jira Done chỉ ghi nhận là nền tảng kỹ thuật.",
    ])

    add_heading(doc, "CHƯƠNG 4. THỰC HIỆN SPRINT 1 (JIRA SPRINT 1 + 2)", 1, page_break=True)
    add_heading(doc, "4.1. Sprint Goal và phạm vi", 2)
    add_para(doc, "Sprint 1 của báo cáo kéo dài từ 02/07 đến 30/07/2026 và gộp hai timebox đầu trên Jira. Mục tiêu là tạo nền tảng xác thực/phân quyền, danh bạ tổ chức, dự án và các chức năng công việc/cộng tác đủ để chạy demo. Jira ghi nhận toàn bộ 17 issue hoàn thành, tổng 99/99 Story Point Done.")
    add_callout(doc, "Kết quả quản lý", "17 issue — 99 Story Point — 100% Story Point ở trạng thái Done trên Jira. Đây là số liệu chính thức dùng trong báo cáo.", PALE_GREEN)
    s1_rows = [(i + 1, x[5], x[0], x[1], x[2], x[3], x[4], "Done") for i, x in enumerate(SPRINT1)]
    add_table(doc, ["#", "Jira Sprint", "Issue", "US", "Chức năng", "SP", "Phụ trách", "Trạng thái"], s1_rows,
              widths=[0.8, 2.2, 1.8, 1.3, 5.0, 1.0, 3.0, 1.5], font_size=7.5)
    add_caption(doc, "Bảng 4.1. Phạm vi Sprint 1 đã đối chiếu trực tiếp với Jira")
    add_heading(doc, "4.2. Sprint Planning", 2)
    add_para(doc, "Nhóm ưu tiên các hạng mục có tính nền tảng trước: đăng nhập, gia hạn phiên, vai trò và RBAC; sau đó mở rộng sang hồ sơ nhân viên, cơ cấu phòng ban, dự án và Kanban. Các Story Point được sử dụng như thước đo tương đối về độ phức tạp, không quy đổi trực tiếp thành giờ công.")
    add_bullets(doc, [
        "Jira Sprint 1 (02/07–16/07): 8 issue, 39 SP — nền tảng xác thực, vai trò, dự án, Kanban và tải tài liệu.",
        "Jira Sprint 2 (16/07–30/07): 9 issue, 60 SP — nhân viên, phòng ban, chuyển phòng, phụ thuộc Task, công việc con và bình luận.",
        "Rủi ro chính: tích hợp API–UI, kiểm soát quyền nhất quán và sai lệch dữ liệu/migration.",
    ])
    add_heading(doc, "4.3. Tổ chức thực hiện và Daily Scrum", 2)
    rows = [
        ("02/07", "Chốt mục tiêu, rà backlog và phân công nhóm nền tảng Auth/RBAC.", "Phạm vi và owner rõ ràng."),
        ("06/07", "Rà API/Frontend, xử lý lỗi 404 và trạng thái đăng nhập.", "Ổn định luồng đăng nhập."),
        ("13/07", "Tích hợp database, logging, cache/notification nền tảng và kiểm thử.", "Giảm lỗi tích hợp."),
        ("16/07", "Review Jira Sprint 1, chuyển sang danh bạ tổ chức và tác vụ.", "39 SP hoàn thành."),
        ("22/07", "Rà Project, Department/Team, Task, bình luận và blocker.", "Bổ sung bằng chứng demo."),
        ("30/07", "Review Jira Sprint 2, chạy test/build và chốt phạm vi báo cáo.", "60 SP hoàn thành."),
    ]
    add_table(doc, ["Mốc", "Nội dung Daily/Review", "Kết quả"], rows, widths=[2.2, 9.0, 5.1], font_size=9)
    add_heading(doc, "4.4. Increment bàn giao", 2)
    add_bullets(doc, [
        "Nhóm xác thực và phân quyền: đăng nhập, gia hạn phiên, danh mục vai trò và RBAC.",
        "Nhóm tổ chức: hồ sơ nhân viên, quản lý trực tiếp, phòng ban, trưởng phòng và chuyển phòng.",
        "Nhóm dự án: tạo dự án và phân công thành viên.",
        "Nhóm công việc/cộng tác: Kanban, phụ thuộc, công việc con, luồng bình luận và tải nhiều tài liệu.",
    ])
    add_heading(doc, "4.5. Sprint Review", 2)
    add_para(doc, "Tại Sprint Review, nhóm đối chiếu từng issue với trạng thái Jira, mã nguồn, kiểm thử và màn hình chạy thử. Toàn bộ 17 issue đã ở trạng thái Done. Tuy nhiên, báo cáo không dùng trạng thái Jira như bằng chứng duy nhất: các luồng chính vẫn được chạy kiểm thử và build lại ở Chương 6.")
    add_heading(doc, "4.6. Sprint Retrospective", 2)
    rows = [
        ("Start", "Cập nhật Jira ngay sau khi xác minh Acceptance Criteria; lưu ảnh demo theo chức năng."),
        ("Stop", "Không cộng SP từ bảng nháp; không xem nền tảng kỹ thuật là chức năng đã nghiệm thu."),
        ("Continue", "Chạy test/build trước Review; chia module theo miền; dùng checklist RBAC/migration."),
        ("Action", "Giới hạn WIP, hoàn tất ước lượng và tách hạng mục bắt buộc với mở rộng."),
    ]
    add_table(doc, ["Nhóm", "Nội dung"], rows, widths=[3.0, 13.3], font_size=9)

    add_heading(doc, "CHƯƠNG 5. THỰC HIỆN SPRINT 2 (JIRA SPRINT 3 + 4)", 1, page_break=True)
    add_heading(doc, "5.1. Trạng thái tại ngày chốt báo cáo", 2)
    add_para(doc, "Sprint 2 của báo cáo bao gồm Jira Sprint 3 (30/07–13/08) và Jira Sprint 4 (14/08–28/08). Tại thời điểm đối chiếu 10/08/2026, cả 31 issue đều đang To Do. Vì vậy chương này trình bày kế hoạch thực hiện, mức sẵn sàng và bằng chứng nền tảng; không ghi nhận kết quả hoàn thành.")
    add_callout(doc, "Kết luận Sprint 2", "31 issue — 59 Story Point đã được ước lượng — 0 Story Point Done. Toàn bộ Jira Sprint 4 được xem là phạm vi mở rộng/tự nghiên cứu cho đến khi được ước lượng và nghiệm thu.", PALE_AMBER)
    add_heading(doc, "5.2. Jira Sprint 3 — phạm vi ưu tiên", 2)
    s3_rows = [(i + 1, x[0], x[1], x[2], "Chưa ước lượng" if x[3] is None else x[3], x[4], "To Do") for i, x in enumerate(SPRINT2_JIRA3)]
    add_table(doc, ["#", "Issue", "US", "Chức năng", "SP", "Phụ trách", "Trạng thái"], s3_rows,
              widths=[0.8, 1.8, 1.3, 6.1, 2.0, 3.0, 1.5], font_size=8)
    add_caption(doc, "Bảng 5.1. Jira Sprint 3: 16 issue, 59 SP đã ước lượng, tất cả To Do")
    add_para(doc, "Định hướng ưu tiên của Sprint 3 là hoàn thiện vòng đời phiên đăng nhập, chỉnh sửa Task, lịch công việc, thông báo, nghỉ phép và dashboard. Bốn issue chưa có Story Point cần được refinement trước khi cam kết năng lực Sprint.")
    add_heading(doc, "5.3. Jira Sprint 4 — phạm vi mở rộng", 2)
    s4_rows = [(i + 1, x[0], x[1], x[2], "Chưa ước lượng", "To Do", "Mở rộng/tự nghiên cứu") for i, x in enumerate(SPRINT2_JIRA4)]
    add_table(doc, ["#", "Issue", "US", "Chức năng", "SP", "Trạng thái", "Phân loại báo cáo"], s4_rows,
              widths=[0.8, 1.8, 1.3, 6.2, 2.0, 1.5, 2.7], font_size=8)
    add_caption(doc, "Bảng 5.2. Jira Sprint 4: 15 issue chưa ước lượng, chưa triển khai theo trạng thái Jira")
    add_heading(doc, "5.4. Phân biệt nền tảng kỹ thuật và chức năng đã nghiệm thu", 2)
    rows = [
        ("Calendar, Dashboard, Notification", "Đã có page/API hoặc nền tảng trong mã nguồn", "Jira Sprint 3 vẫn To Do", "Chỉ ghi nhận mức sẵn sàng kỹ thuật; chưa tính Done."),
        ("Health, logging, Redis, Docker", "Có cấu hình/module/hạ tầng trong repository", "Các US47–US50 chưa Done", "Dùng làm bằng chứng nền tảng và hướng mở rộng."),
        ("Nghỉ phép/làm thêm giờ", "Có schema hoặc một phần nền tảng", "Jira Sprint 3–4 chưa Done", "Cần refinement, UI/API, test và Review trước nghiệm thu."),
    ]
    add_table(doc, ["Nhóm", "Bằng chứng mã nguồn", "Trạng thái quản lý", "Cách trình bày"], rows, widths=[3.3, 5.0, 3.7, 4.3], font_size=9)
    add_heading(doc, "5.5. Kế hoạch thực hiện", 2)
    add_bullets(doc, [
        "Refinement: bổ sung Acceptance Criteria, phụ thuộc, owner và Story Point cho tất cả issue còn thiếu.",
        "Thứ tự triển khai: phiên đăng nhập/Task trước; Calendar/Notification/Dashboard sau; nghỉ phép và báo cáo nâng cao theo năng lực còn lại.",
        "Chính sách WIP: mỗi thành viên không đồng thời kéo quá hai hạng mục; issue bị chặn phải ghi blocker trong Daily Scrum.",
        "Kiểm thử: thêm test cho luồng mới, chạy hồi quy backend/frontend, build production và demo theo Acceptance Criteria.",
        "Jira Sprint 4: chỉ đưa vào cam kết sau khi Sprint 3 ổn định; nếu không đủ năng lực thì chuyển Product Backlog và ghi rõ lý do.",
    ])
    add_heading(doc, "5.6. Rủi ro và biện pháp", 2)
    rows = [
        ("31 issue nhưng chỉ 59 SP được ước lượng", "Không dự báo được năng lực", "Refinement và ước lượng trước Planning/điều chỉnh Sprint scope."),
        ("Nhiều module phụ thuộc Auth/RBAC", "Lỗi tích hợp hoặc hở quyền", "Checklist quyền, test 401/403 và review API."),
        ("Nhầm nền tảng với Done", "Báo cáo sai mức hoàn thành", "Chỉ tính Done khi Jira, Acceptance Criteria và bằng chứng khớp nhau."),
        ("Bundle frontend lớn", "Tải trang chậm", "Code splitting theo route và đo lại kích thước bundle."),
    ]
    add_table(doc, ["Rủi ro", "Tác động", "Biện pháp"], rows, widths=[5.0, 4.2, 7.1], font_size=9)

    add_heading(doc, "CHƯƠNG 6. KẾT QUẢ KỸ THUẬT VÀ BẰNG CHỨNG", 1, page_break=True)
    add_heading(doc, "6.1. Kết quả kiểm thử độc lập", 2)
    rows = [
        ("Backend", "pytest", "437/437 test passed", "262,06 giây", "Đạt"),
        ("Frontend", "Vitest", "28/28 test passed", "Lần chạy xác minh", "Đạt"),
        ("Frontend production", "Vite build", "Build thành công", "4,05 giây", "Đạt, có cảnh báo bundle lớn"),
    ]
    add_table(doc, ["Phạm vi", "Công cụ", "Kết quả", "Thời gian", "Đánh giá"], rows, widths=[3.2, 3.0, 4.0, 3.0, 3.1], font_size=9)
    add_caption(doc, "Bảng 6.1. Kết quả chạy lại trên mã nguồn local ngày 10/08/2026")
    add_callout(doc, "Điểm cần cải tiến", "File JavaScript production khoảng 1.378,94 kB (gzip 378,46 kB). Đây không làm build thất bại nhưng nên tách bundle theo route/module trong Sprint tiếp theo.", PALE_AMBER)
    add_heading(doc, "6.2. Màn hình chức năng", 2)
    figures = [
        ("01-login.png", "Hình 6.1. Màn hình đăng nhập — bằng chứng cho nhóm Auth."),
        ("02-dashboard.png", "Hình 6.2. Dashboard tổng quan của hệ thống."),
        ("03-departments.png", "Hình 6.3. Danh sách phòng ban."),
        ("05-teams.png", "Hình 6.4. Danh sách đội nhóm."),
        ("07-projects.png", "Hình 6.5. Danh sách dự án."),
        ("08-sprints.png", "Hình 6.6. Quản lý Sprint trong dự án."),
        ("10-backlog.png", "Hình 6.7. Product/Project Backlog trên ứng dụng."),
        ("11-tasks.png", "Hình 6.8. Danh sách và quản lý công việc."),
        ("12-notifications.png", "Hình 6.9. Trung tâm thông báo — bằng chứng nền tảng, chưa thay thế trạng thái Jira Sprint 2."),
    ]
    for idx, (name, cap) in enumerate(figures):
        add_figure(doc, SCREEN_DIR / name, cap, width=6.15)
        if idx in (1, 3, 5, 7):
            doc.add_page_break()
    add_heading(doc, "6.3. Truy vết giữa yêu cầu và bằng chứng", 2)
    rows = [
        ("Jira", "Trạng thái issue, Story Point, Sprint, assignee", "Nguồn quản lý phạm vi và tiến độ."),
        ("Mã nguồn", "Router/service/model/page/component", "Chứng minh tồn tại triển khai kỹ thuật."),
        ("Kiểm thử", "437 backend, 28 frontend, production build", "Chứng minh hồi quy tại thời điểm báo cáo."),
        ("Ảnh chạy thử", "Login, Dashboard, Department, Team, Project, Sprint, Backlog, Task, Notification", "Chứng minh luồng UI có thể trình bày."),
        ("Bảng tính", "Dashboard Báo cáo, Sprint 1, Sprint 2, Product Backlog", "Nguồn tổng hợp dành cho giảng viên và nhóm."),
    ]
    add_table(doc, ["Nguồn", "Bằng chứng", "Vai trò"], rows, widths=[3.0, 7.0, 6.3], font_size=9)

    add_heading(doc, "CHƯƠNG 7. ĐÁNH GIÁ VÀ KẾT LUẬN", 1, page_break=True)
    add_heading(doc, "7.1. Kết quả đạt được", 2)
    add_para(doc, "Nhóm đã hình thành một hệ thống quản lý công việc doanh nghiệp với nền tảng kiến trúc, cơ sở dữ liệu, xác thực/phân quyền, cơ cấu tổ chức, dự án và công việc. Sprint 1 có bằng chứng quản lý rõ ràng với 99/99 Story Point Done. Mã nguồn hiện tại vượt qua 437 kiểm thử backend, 28 kiểm thử frontend và production build.")
    add_heading(doc, "7.2. Hạn chế", 2)
    add_bullets(doc, [
        "Sprint 2 chưa hoàn thành theo Jira; nhiều issue chưa có Story Point và assignee.",
        "Một số module có nền tảng trong mã nguồn nhưng chưa đủ điều kiện nghiệm thu theo Acceptance Criteria.",
        "Tài liệu Daily Scrum/Review cần được cập nhật đều và gắn link bằng chứng trực tiếp hơn.",
        "Bundle frontend lớn; cần tối ưu tải trang và đo hiệu năng trước khi triển khai thực tế.",
    ])
    add_heading(doc, "7.3. Bài học Agile", 2)
    add_para(doc, "Bài học quan trọng nhất là tính minh bạch: trạng thái quản lý, số liệu ước lượng và bằng chứng kỹ thuật phải được đối chiếu thay vì suy diễn từ một nguồn duy nhất. Definition of Done giúp ngăn việc xem một module ‘đã có mã nguồn’ như một Increment đã được chấp nhận. Việc giới hạn WIP, refinement trước Planning và cập nhật Jira ngay sau Review sẽ giúp Sprint 2 thực tế hơn.")
    add_heading(doc, "7.4. Hướng phát triển", 2)
    add_bullets(doc, [
        "Hoàn thiện Jira Sprint 3 theo thứ tự ưu tiên, bổ sung test và bằng chứng Review.",
        "Đưa Jira Sprint 4 về Product Backlog nếu chưa đủ năng lực; chỉ chọn lại sau khi ước lượng.",
        "Tối ưu frontend bằng route-based code splitting; tăng observability và kiểm thử phân quyền.",
        "Chuẩn hóa CI/CD, migration, backup và hướng dẫn triển khai để tiến tới môi trường production.",
    ])

    add_heading(doc, "PHỤ LỤC A. MA TRẬN ĐỐI CHIẾU SPRINT", 1, page_break=True)
    rows = [
        ("Sprint 1 báo cáo", "Jira Sprint 1", 8, 39, 39, "Done"),
        ("Sprint 1 báo cáo", "Jira Sprint 2", 9, 60, 60, "Done"),
        ("Sprint 2 báo cáo", "Jira Sprint 3", 16, 59, 0, "To Do"),
        ("Sprint 2 báo cáo", "Jira Sprint 4", 15, "Chưa ước lượng", 0, "To Do / mở rộng"),
    ]
    add_table(doc, ["Sprint báo cáo", "Jira Sprint", "Issue", "Tổng SP", "SP Done", "Trạng thái"], rows, widths=[3.3, 3.0, 2.0, 3.0, 2.5, 3.0], font_size=9)
    add_heading(doc, "PHỤ LỤC B. LIÊN KẾT BÀN GIAO", 1, page_break=True)
    links = [
        ("GitHub develop", "https://github.com/huynhlethanhnhan/TaskSyncEnterprise/tree/develop"),
        ("Jira Backlog", "https://task-snycs-enterprise.atlassian.net/jira/software/projects/JD/boards/1/backlog"),
        ("Google Drive nguồn", "https://drive.google.com/drive/u/0/folders/1DRviUSV70Enlq5uNAF0L59GLsUsliNWT"),
        ("Bảng tính Agile đã đối chiếu", "https://docs.google.com/spreadsheets/d/1pBk8gqkxyRddmChIlGe78xFSgaf-65pRsKoh2rqwkx4/edit"),
    ]
    for label, url in links:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(label + ": "), 11, bold=True)
        add_hyperlink(p, url, url)
    add_para(doc, "Ngày chốt dữ liệu: 10/08/2026. Khi Jira thay đổi sau mốc này, nhóm cần cập nhật lại dashboard bảng tính và các bảng Sprint trong báo cáo.", italic=True, first_line=False, color=GRAY)

    doc.core_properties.title = "Báo cáo đồ án Agile - TaskSyncEnterprise"
    doc.core_properties.subject = "Báo cáo Sprint, kiến trúc, kiểm thử và đối chiếu Jira"
    doc.core_properties.author = "Nhóm TaskSyncEnterprise"
    doc.core_properties.keywords = "Agile, Scrum, TaskSyncEnterprise, Jira, Sprint"
    doc.save(str(OUT_DOCX))
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
